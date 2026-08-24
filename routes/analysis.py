"""数据分析路由 Blueprint — 测试报告分析、Excel CR分析、Excel智能整理、PDF生成

包含：
- 测试报告分析（/api/test-report-*）
- Excel CR问题分析（/api/excel-analyze*）
- Excel智能整理与PDF导出（/excel-*, /preview, /convert, /upload）
- AI增强分析（/api/excel-analyze-ai*）
"""
import os
import re
import gc
import json
import time
import hashlib
import logging
import traceback
import threading
from datetime import datetime
from functools import wraps
from html import escape as _html_escape

from flask import Blueprint, request, jsonify, Response, stream_with_context, send_file, g, current_app

import auth
import ai_utils
from routes.common import (
    ExcelReader, validate_file_id, render_pdf, MD2PDF_PREVIEW_CSS,
    background_tasks, load_task_meta, save_task_meta, _CST,
)
from report_builders import _build_cr_analysis_report_html

logger = logging.getLogger(__name__)

get_ai_config = ai_utils.get_ai_config
_call_ai = ai_utils.call_ai
_call_ai_stream = ai_utils.call_ai_stream

# Excel 分析器导入
from excel_analyzers import (
    _analyze_issue_sheet, _analyze_issue_sheet_fast, _detect_issue_columns,
    _log_mem,
)

# 分析结果缓存
_analysis_cache = {}


def _escape_html(text):
    """HTML转义辅助函数"""
    if text is None:
        return ''
    return _html_escape(str(text))


# ==================== 认证装饰器 ====================

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        user = auth.get_current_user()
        if not user:
            return jsonify({'error': '请先登录'}), 401
        g.user = user
        return f(*args, **kwargs)
    return decorated


def login_required_or_guest(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        user = auth.get_current_user()
        if not user and not auth.ALLOW_GUEST:
            return jsonify({'error': '请先登录'}), 401
        g.user = user
        return f(*args, **kwargs)
    return decorated


def create_analysis_blueprint():
    """创建数据分析路由 Blueprint"""
    bp = Blueprint('analysis', __name__)

    # ==================== 测试报告分析 ====================

    @bp.route('/api/test-report-analyze', methods=['POST'])
    def api_test_report_analyze():
        """上传测试报告Excel，返回sheet列表和文件ID"""
        if 'file' not in request.files:
            return jsonify({'error': '请选择文件'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择文件'}), 400
        filename_lower = file.filename.lower()
        if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
            return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        if file_size > 200 * 1024 * 1024:
            return jsonify({'error': f'文件过大({file_size // 1024 // 1024}MB)，最大支持200MB'}), 413
        orig_ext = os.path.splitext(file.filename)[1].lower()
        if orig_ext not in ('.xlsx', '.xls', '.csv'):
            orig_ext = '.xlsx'
        try:
            file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"test_{file_id}{orig_ext}")
            file.save(file_path)
            reader = ExcelReader(file_path)
            reader.open()
            sheet_names = reader.get_sheet_names()
            reader.close()
            return jsonify({'status': 'success', 'data': {'file_id': file_id, 'file_name': file.filename, 'sheet_names': sheet_names}})
        except Exception as e:
            logger.error(f"测试报告上传失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/api/test-report-analyze-sheet', methods=['POST'])
    def api_test_report_analyze_sheet():
        """分析指定sheet的测试报告数据"""
        data = request.get_json(silent=True) or {}
        file_id = data.get('file_id', '')
        sheet_name = data.get('sheet_name', '')
        if not file_id or not sheet_name:
            return jsonify({'error': '缺少参数: file_id, sheet_name'}), 400
        if not validate_file_id(file_id):
            return jsonify({'error': '无效的文件ID'}), 400
        file_path = None
        for ext in ['.xlsx', '.xls', '.csv']:
            candidate = os.path.join(current_app.config['UPLOAD_FOLDER'], f"test_{file_id}{ext}")
            if os.path.exists(candidate):
                file_path = candidate
                break
        if not file_path:
            return jsonify({'error': f'文件不存在: {file_id}'}), 404
        try:
            reader = ExcelReader(file_path)
            reader.open()
            rows = reader.get_sheet_data(sheet_name)
            reader.close()
            if not rows or len(rows) < 2:
                return jsonify({'status': 'done', 'data': {'headers': [], 'summary': {}, 'modules': [], 'developers': [], 'severity': {}, 'status': {}, 'trend': []}})
            headers = [str(c).strip() if c else '' for c in rows[0]]
            data_rows = rows[1:]
            result = _analyze_test_report(headers, data_rows)
            del rows, data_rows
            gc.collect()
            return jsonify({'status': 'done', 'data': result})
        except Exception as e:
            logger.error(f"测试报告分析失败: {traceback.format_exc()}")
            return jsonify({'status': 'error', 'error': str(e)}), 500

    @bp.route('/api/test-report-debug', methods=['POST'])
    def api_test_report_debug():
        """调试接口：返回原始数据前几行"""
        data = request.get_json(silent=True) or {}
        file_id = data.get('file_id', '')
        sheet_name = data.get('sheet_name', '')
        if not file_id or not sheet_name:
            return jsonify({'error': '缺少参数'}), 400
        file_path = None
        for ext in ['.xlsx', '.xls', '.csv']:
            candidate = os.path.join(current_app.config['UPLOAD_FOLDER'], f"test_{file_id}{ext}")
            if os.path.exists(candidate):
                file_path = candidate
                break
        if not file_path:
            return jsonify({'error': '文件不存在'}), 404
        try:
            reader = ExcelReader(file_path)
            reader.open()
            rows = reader.get_sheet_data(sheet_name)
            reader.close()
            return jsonify({'headers': rows[0] if rows else [], 'sample_rows': rows[1:6] if len(rows) > 1 else [], 'total_rows': len(rows) - 1 if rows else 0})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @bp.route('/api/test-report-pdf', methods=['POST'])
    def api_test_report_pdf():
        """生成测试报告分析PDF"""
        data = request.get_json(silent=True) or {}
        analysis_data = data.get('analysis_data', {})
        watermark = data.get('watermark', '')
        custom_title = data.get('custom_title', '').strip()
        file_name = data.get('file_name', '')
        if not analysis_data:
            return jsonify({'error': '缺少分析数据'}), 400
        try:
            html_content = _build_test_report_html(analysis_data, watermark, file_name, custom_title)
            import tempfile as tf
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            if custom_title:
                safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
                pdf_filename = f"{safe_title}_{datetime.now(_CST).strftime('%Y%m%d_%H%M%S')}.pdf"
                download_name = f"{safe_title}.pdf"
            else:
                pdf_filename = f"test_report_{int(time.time())}.pdf"
                download_name = pdf_filename
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            render_pdf(html_path, pdf_path, margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'}, extra_wait_ms=2000)
            try:
                os.unlink(html_path)
            except Exception:
                pass
            return jsonify({'status': 'success', 'filename': pdf_filename, 'download_name': download_name})
        except Exception as e:
            logger.error(f"测试报告PDF生成失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/api/test-report-download/<filename>')
    def api_test_report_download(filename):
        from werkzeug.utils import secure_filename
        safe_name = secure_filename(filename)
        if not safe_name or safe_name != filename:
            return jsonify({'error': '无效的文件名'}), 400
        filepath = os.path.join(current_app.config['PDF_FOLDER'], safe_name)
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
        return jsonify({'error': '文件不存在'}), 404

    @bp.route('/api/test-report-ai-analysis', methods=['POST'])
    @login_required_or_guest
    def api_test_report_ai_analysis():
        """AI分析测试报告"""
        ai_config = get_ai_config()
        if not ai_config.get('enabled'):
            return jsonify({'error': 'AI功能未配置'}), 503
        data = request.get_json(silent=True) or {}
        analysis = data.get('analysis', {})
        if not analysis:
            return jsonify({'error': '缺少分析数据'}), 400
        summary = analysis.get('summary', {})
        modules = analysis.get('modules', [])
        prompt = f"""你是一位资深测试质量专家，请基于以下测试报告分析数据生成专业的AI分析报告。

## 测试统计
- 总用例数: {summary.get('total', 0)}
- 通过: {summary.get('passed', 0)}
- 失败: {summary.get('failed', 0)}
- 阻塞: {summary.get('blocked', 0)}
- 跳过: {summary.get('skipped', 0)}
- 通过率: {summary.get('pass_rate', 'N/A')}

## 模块统计 (Top 10)
{chr(10).join([f"- {m.get('name', 'N/A')}: {m.get('total', 0)}个用例, 失败{m.get('failed', 0)}个" for m in (modules[:10] if isinstance(modules, list) else [])]) or 'N/A'}

请按以下格式输出分析：
### 📊 总体评估
### 🔍 失败根因分析
### ⚠️ 高风险模块
### 💡 改进建议
请使用简洁专业的中文。"""
        try:
            messages = [{'role': 'user', 'content': prompt}]
            reply = _call_ai(messages, max_tokens=1500, temperature=0.3, timeout=60)
            return jsonify({'status': 'success', 'analysis': reply})
        except Exception as e:
            logger.error(f'AI测试报告分析失败: {e}')
            return jsonify({'error': f'AI分析失败: {str(e)}'}), 502

    @bp.route('/api/test-report-ai-stream', methods=['POST'])
    @login_required_or_guest
    def api_test_report_ai_stream():
        """SSE流式版：AI分析测试报告"""
        ai_config = get_ai_config()
        if not ai_config.get('enabled'):
            return jsonify({'error': 'AI功能未配置'}), 503
        data = request.get_json(silent=True) or {}
        analysis = data.get('analysis', {})
        if not analysis:
            return jsonify({'error': '缺少分析数据'}), 400
        summary = analysis.get('summary', {})
        modules = analysis.get('modules', [])
        prompt = f"""你是一位资深测试质量专家，请基于以下测试报告分析数据生成专业的AI分析报告。

## 测试统计
- 总用例数: {summary.get('total', 0)}
- 通过: {summary.get('passed', 0)}
- 失败: {summary.get('failed', 0)}
- 阻塞: {summary.get('blocked', 0)}
- 跳过: {summary.get('skipped', 0)}
- 通过率: {summary.get('pass_rate', 'N/A')}

## 模块统计 (Top 10)
{chr(10).join([f"- {m.get('name', 'N/A')}: {m.get('total', 0)}个用例, 失败{m.get('failed', 0)}个" for m in (modules[:10] if isinstance(modules, list) else [])]) or 'N/A'}

请按以下格式输出分析：
### 📊 总体评估
### 🔍 失败根因分析
### ⚠️ 高风险模块
### 💡 改进建议
请使用简洁专业的中文。"""
        messages = [{'role': 'user', 'content': prompt}]
        return Response(
            stream_with_context(_call_ai_stream(messages, max_tokens=1500, temperature=0.3)),
            content_type='text/event-stream',
            headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no', 'Connection': 'keep-alive'}
        )

    # ==================== Excel CR 分析 ====================

    @bp.route('/api/excel-analyze', methods=['POST'])
    def api_excel_analyze():
        """上传Excel文件，返回sheet列表"""
        if 'file' not in request.files:
            return jsonify({'error': '请选择文件'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择文件'}), 400
        filename_lower = file.filename.lower()
        if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
            return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        if file_size > 200 * 1024 * 1024:
            return jsonify({'error': f'文件过大({file_size // 1024 // 1024}MB)，云端最大支持200MB。超大文件建议使用分片上传'}), 413
        orig_ext = os.path.splitext(file.filename)[1].lower()
        if orig_ext not in ('.xlsx', '.xls', '.csv'):
            orig_ext = '.xlsx'
        try:
            file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"excel_{file_id}{orig_ext}")
            file.save(file_path)
            logger.info(f"========== 收到问题分析文件: {file.filename} ==========")
            reader = ExcelReader(file_path)
            reader.open()
            sheet_names = reader.get_sheet_names()
            reader.close()
            return jsonify({'status': 'success', 'data': {'file_id': file_id, 'file_name': file.filename, 'sheet_names': sheet_names}})
        except Exception as e:
            logger.error(f"文件上传失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/api/excel-analyze-fields', methods=['POST'])
    def api_excel_analyze_fields():
        """轻量级字段映射接口：同步执行（仅需读取表头）"""
        data = request.get_json(silent=True) or {}
        file_id = data.get('file_id', '')
        sheet_name = data.get('sheet_name', '')
        logger.info(f"excel-analyze-fields 请求: file_id={file_id}, sheet_name={sheet_name}")
        if not file_id or not sheet_name:
            return jsonify({'error': '缺少参数: file_id, sheet_name'}), 400
        if not validate_file_id(file_id):
            return jsonify({'error': '无效的文件ID'}), 400
        file_path = None
        for ext in ['.xlsx', '.xls', '.csv']:
            candidate = os.path.join(current_app.config['UPLOAD_FOLDER'], f"excel_{file_id}{ext}")
            if os.path.exists(candidate):
                file_path = candidate
                break
        if not file_path:
            return jsonify({'error': f'文件不存在: {file_id}'}), 404
        try:
            reader = ExcelReader(file_path)
            reader.open()
            if reader._is_html:
                headers, first_data_row, data_row_count = reader._parse_html_headers_only()
                reader.close()
                if not headers:
                    return jsonify({'status': 'done', 'data': {'headers': [], 'detected_columns': {}, 'detected_fields_count': 0, 'current_sheet': sheet_name, 'summary': {'total_issues': 0}, 'sample_data': []}})
                headers = [str(c).strip() if c else '' for c in headers]
                col_map = _detect_issue_columns(headers)
                raw_detected = {'issue_id': col_map.get('id', -1), 'title': col_map.get('title', -1), 'module': col_map.get('module', -1), 'severity': col_map.get('severity', -1), 'status': col_map.get('status', -1), 'developer': col_map.get('developer', -1), 'create_date': col_map.get('created_date', -1), 'resolve_date': col_map.get('resolved_date', -1), 'fixed_date': col_map.get('closed_date', -1), 'fixed_version': col_map.get('fix_version', -1)}
                detected_columns = {k: v for k, v in raw_detected.items() if v >= 0}
                sample_data = [first_data_row] if first_data_row else []
                total_issues = data_row_count
                logger.info(f"字段映射完成(轻量级): detected_columns={detected_columns}, total_issues≈{total_issues}")
                result = {'headers': headers, 'detected_columns': detected_columns, 'detected_fields_count': len(detected_columns), 'current_sheet': sheet_name, 'summary': {'total_issues': total_issues}, 'sample_data': sample_data}
                gc.collect()
                return jsonify({'status': 'done', 'data': result})
            else:
                # v6.0 优化：只读表头 + 流式统计行数，不加载全部数据
                try:
                    from excel_analyzers import _read_excel_headers_only
                    headers = _read_excel_headers_only(file_path, sheet_name)
                except Exception:
                    headers = None

                if not headers:
                    # 回退到原方法
                    rows = reader.get_sheet_data(sheet_name)
                    reader.close()
                    if not rows or len(rows) < 1:
                        return jsonify({'status': 'done', 'data': {'headers': [], 'detected_columns': {}, 'detected_fields_count': 0, 'current_sheet': sheet_name, 'summary': {'total_issues': 0}, 'sample_data': []}})
                    headers = [str(c).strip() if c else '' for c in rows[0]]
                    col_map = _detect_issue_columns(headers)
                    raw_detected = {'issue_id': col_map.get('id', -1), 'title': col_map.get('title', -1), 'module': col_map.get('module', -1), 'severity': col_map.get('severity', -1), 'status': col_map.get('status', -1), 'developer': col_map.get('developer', -1), 'create_date': col_map.get('created_date', -1), 'resolve_date': col_map.get('resolved_date', -1), 'fixed_date': col_map.get('closed_date', -1), 'fixed_version': col_map.get('fix_version', -1)}
                    detected_columns = {k: v for k, v in raw_detected.items() if v >= 0}
                    data_rows = rows[1:]
                    total_issues = sum(1 for row in data_rows if any(str(c).strip() for c in row))
                    sample_data = data_rows[:3] if data_rows else []
                    logger.info(f"字段映射完成(回退): detected_columns={detected_columns}, total_issues={total_issues}")
                    result = {'headers': headers, 'detected_columns': detected_columns, 'detected_fields_count': len(detected_columns), 'current_sheet': sheet_name, 'summary': {'total_issues': total_issues}, 'sample_data': sample_data}
                    del rows, data_rows
                    gc.collect()
                    return jsonify({'status': 'done', 'data': result})

                reader.close()
                headers = [str(c).strip() if c else '' for c in headers]
                col_map = _detect_issue_columns(headers)
                raw_detected = {'issue_id': col_map.get('id', -1), 'title': col_map.get('title', -1), 'module': col_map.get('module', -1), 'severity': col_map.get('severity', -1), 'status': col_map.get('status', -1), 'developer': col_map.get('developer', -1), 'create_date': col_map.get('created_date', -1), 'resolve_date': col_map.get('resolved_date', -1), 'fixed_date': col_map.get('closed_date', -1), 'fixed_version': col_map.get('fix_version', -1)}
                detected_columns = {k: v for k, v in raw_detected.items() if v >= 0}

                # 流式统计行数 + 取前3行示例
                total_issues = 0
                sample_data = []
                try:
                    from openpyxl import load_workbook
                    wb = load_workbook(file_path, read_only=True, data_only=True)
                    if sheet_name not in wb.sheetnames:
                        sheet_name = wb.sheetnames[0]
                    ws = wb[sheet_name]
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        if i == 0:
                            continue  # 跳过表头
                        row_vals = [str(c).strip() if c is not None else '' for c in row]
                        if any(row_vals):
                            total_issues += 1
                            if len(sample_data) < 3:
                                sample_data.append(row_vals)
                    wb.close()
                except Exception as e:
                    logger.warning(f"流式统计行数失败，使用估算值: {e}")
                    total_issues = 0

                logger.info(f"字段映射完成(优化): detected_columns={detected_columns}, total_issues={total_issues}")
                result = {'headers': headers, 'detected_columns': detected_columns, 'detected_fields_count': len(detected_columns), 'current_sheet': sheet_name, 'summary': {'total_issues': total_issues}, 'sample_data': sample_data}
                gc.collect()
                return jsonify({'status': 'done', 'data': result})
        except Exception as e:
            logger.error(f"字段映射失败: {traceback.format_exc()}")
            return jsonify({'status': 'error', 'error': str(e)}), 500

    @bp.route('/api/task-status', methods=['POST'])
    def api_task_status():
        """查询后台任务状态"""
        data = request.get_json(silent=True) or {}
        task_id = data.get('task_id', '')
        if not task_id:
            return jsonify({'error': '无效的 task_id'}), 400
        task = background_tasks.get(task_id)
        if task is None:
            task = load_task_meta(task_id)
            if task is None:
                return jsonify({'status': 'error', 'error': '任务不存在或已过期，请重新上传文件'}), 400
        if task['status'] == 'processing':
            created_at = task.get('created_at', 0)
            if created_at and (time.time() - created_at) > 900:
                task['status'] = 'error'
                task['error'] = '分析超时（超过15分钟），可能是文件过大或格式异常，请尝试减少数据量或转换为 .xlsx 格式'
                save_task_meta(task_id, task)
        resp = {'status': task['status']}
        if task['status'] == 'processing':
            resp['progress'] = task.get('progress', 0)
            resp['progress_msg'] = task.get('progress_msg', '正在分析...')
        elif task['status'] == 'done':
            resp['data'] = task['result']
        elif task['status'] == 'error':
            resp['error'] = task['error']
        return jsonify(resp)

    @bp.route('/api/task-cancel', methods=['POST'])
    def api_task_cancel():
        """取消正在进行的分析任务

        v6.0 新增：支持取消异步分析任务
        """
        data = request.get_json(silent=True) or {}
        task_id = data.get('task_id', '')
        if not task_id:
            return jsonify({'error': '无效的 task_id'}), 400

        task_info = background_tasks.get(task_id)
        if not task_info:
            task_info = load_task_meta(task_id)
            if not task_info:
                return jsonify({'status': 'error', 'error': '任务不存在或已过期'}), 404

        current_status = task_info.get('status', 'unknown')
        if current_status in ('done', 'error', 'cancelled'):
            return jsonify({
                'status': 'success',
                'message': f'任务已处于 {current_status} 状态，无需取消',
                'task_status': current_status
            })

        # 标记取消
        task_info['cancelled'] = True
        task_info['status'] = 'cancelled'
        task_info['error'] = '用户取消了分析任务'
        task_info['completed_at'] = time.time()
        background_tasks[task_id] = task_info
        save_task_meta(task_id, task_info)

        logger.info(f"任务 {task_id} 被用户取消")
        return jsonify({
            'status': 'success',
            'message': '取消请求已发送',
            'task_status': 'cancelled'
        })

    @bp.route('/api/excel-analyze-sheet', methods=['POST'])
    def api_excel_analyze_sheet():
        """完整分析接口：启动后台分析，立即返回 task_id"""
        data = request.get_json(silent=True) or {}
        file_id = data.get('file_id', '')
        sheet_name = data.get('sheet_name', '')
        logger.info(f"excel-analyze-sheet 请求: file_id={file_id}, sheet_name={sheet_name}")
        if not file_id or not sheet_name:
            return jsonify({'error': f'缺少参数: file_id={repr(file_id)}, sheet_name={repr(sheet_name)}'}), 400
        if not validate_file_id(file_id):
            return jsonify({'error': '无效的文件ID'}), 400
        file_path = None
        for ext in ['.xlsx', '.xls', '.csv']:
            candidate = os.path.join(current_app.config['UPLOAD_FOLDER'], f"excel_{file_id}{ext}")
            if os.path.exists(candidate):
                file_path = candidate
                break
        if not file_path:
            return jsonify({'error': f'文件不存在: {file_id}'}), 404
        task_id = hashlib.md5(f"sheet_{file_id}_{sheet_name}_{time.time()}".encode()).hexdigest()[:16]
        task_data = {'status': 'processing', 'result': None, 'error': None, 'created_at': time.time()}
        background_tasks[task_id] = task_data
        save_task_meta(task_id, task_data)

        def _do_full_analysis():
            try:
                # 检查是否已取消
                task_meta = background_tasks.get(task_id)
                if task_meta and task_meta.get('cancelled'):
                    background_tasks[task_id]['status'] = 'cancelled'
                    background_tasks[task_id]['error'] = '任务已取消'
                    save_task_meta(task_id, background_tasks[task_id])
                    return

                gc.collect()
                _log_mem("开始分析前")
                t0 = time.time()
                file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
                use_fast = file_size > 10 * 1024 * 1024
                logger.info(f"分析模式: {'fast(pandas)' if use_fast else 'standard(openpyxl)'}, 文件大小: {file_size/1024/1024:.1f}MB")
                if use_fast:
                    def progress_cb(percent, message):
                        # 检查取消
                        task_meta = background_tasks.get(task_id)
                        if task_meta and task_meta.get('cancelled'):
                            return
                        background_tasks[task_id]['progress'] = percent
                        background_tasks[task_id]['progress_msg'] = message
                        save_task_meta(task_id, background_tasks[task_id])
                    result = _analyze_issue_sheet_fast(file_path, sheet_name, progress_cb=progress_cb)
                else:
                    result = _analyze_issue_sheet(file_path, sheet_name)

                # 再次检查取消
                task_meta = background_tasks.get(task_id)
                if task_meta and task_meta.get('cancelled'):
                    background_tasks[task_id]['status'] = 'cancelled'
                    background_tasks[task_id]['error'] = '任务已取消'
                    save_task_meta(task_id, background_tasks[task_id])
                    return

                elapsed = time.time() - t0
                _log_mem(f"分析完成，耗时 {elapsed:.1f}s")
                gc.collect()
                background_tasks[task_id]['result'] = result
                background_tasks[task_id]['status'] = 'done'
                background_tasks[task_id]['progress'] = 100
                background_tasks[task_id]['progress_msg'] = '分析完成'
                background_tasks[task_id]['completed_at'] = time.time()
                save_task_meta(task_id, background_tasks[task_id])
            except Exception as e:
                error_detail = str(e) if str(e) else f'{type(e).__name__} (无详细错误信息)'
                logger.error(f"分析失败: {traceback.format_exc()}")
                background_tasks[task_id]['error'] = error_detail
                background_tasks[task_id]['status'] = 'error'
                background_tasks[task_id]['completed_at'] = time.time()
                save_task_meta(task_id, background_tasks[task_id])

        thread = threading.Thread(target=_do_full_analysis, daemon=True)
        thread.start()
        return jsonify({'status': 'success', 'data': {'task_id': task_id}})

    # ==================== MD2PDF / 旧Excel工具 ====================

    @bp.route('/preview', methods=['POST'])
    def api_preview():
        data = request.json or {}
        markdown_content = data.get('content', '')
        watermark = data.get('watermark', '')
        if not markdown_content:
            return jsonify({'html': ''})
        try:
            import markdown
            import re as _re
            # 检测并转换 Mermaid 代码块为 <div class="mermaid">
            has_mermaid = bool(_re.search(r'```mermaid', markdown_content, _re.IGNORECASE))
            if has_mermaid:
                def mermaid_replacer(match):
                    code = match.group(1).strip()
                    return f'<div class="mermaid">{code}</div>'
                markdown_content = _re.sub(
                    r'```mermaid\s*\n(.*?)```',
                    mermaid_replacer,
                    markdown_content,
                    flags=_re.DOTALL | _re.IGNORECASE
                )
            html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite', 'tables', 'fenced_code'])
            if watermark:
                html_content += f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">{watermark}</div>'''
            return jsonify({'html': html_content, 'has_mermaid': has_mermaid})
        except Exception as e:
            logger.error(f"预览失败: {traceback.format_exc()}")
            return jsonify({'html': '', 'error': str(e)}), 500

    @bp.route('/convert', methods=['POST'])
    def api_convert():
        data = request.json or {}
        markdown_content = data.get('content', '')
        watermark = data.get('watermark', '')
        filename = data.get('filename', '')
        if not markdown_content:
            return jsonify({'error': '内容不能为空'}), 400
        try:
            import markdown
            import tempfile as tf
            import re as _re
            # 检测并转换 Mermaid 代码块
            has_mermaid = bool(_re.search(r'```mermaid', markdown_content, _re.IGNORECASE))
            if has_mermaid:
                def mermaid_replacer(match):
                    code = match.group(1).strip()
                    return f'<div class="mermaid">{code}</div>'
                markdown_content = _re.sub(
                    r'```mermaid\s*\n(.*?)```',
                    mermaid_replacer,
                    markdown_content,
                    flags=_re.DOTALL | _re.IGNORECASE
                )
            html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite', 'tables', 'fenced_code'])
            watermark_html = ''
            if watermark:
                watermark_html = f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">{watermark}</div>'''
            # Mermaid 脚本
            mermaid_script = ''
            wait_selector = None
            extra_wait = 0
            if has_mermaid:
                mermaid_script = '''
                <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
                <script>
                    mermaid.initialize({ startOnLoad: true, theme: 'default', securityLevel: 'loose',
                        flowchart: { useMaxWidth: true, htmlLabels: true },
                        sequence: { useMaxWidth: true }, gantt: { useMaxWidth: true } });
                    window.addEventListener('load', function() {
                        setTimeout(function() { document.body.setAttribute('data-mermaid-done', 'true'); }, 2000);
                    });
                </script>
                <style>.mermaid { text-align: center; margin: 20px 0; page-break-inside: avoid; }
                .mermaid svg { max-width: 100% !important; height: auto !important; }</style>
                '''
                wait_selector = 'body[data-mermaid-done="true"]'
                extra_wait = 3000
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>{MD2PDF_PREVIEW_CSS}</style>{mermaid_script}</head><body>{html_content}{watermark_html}</body></html>''')
                html_path = f.name
            if filename:
                safe_filename = re.sub(r'[^\w\s-]', '', filename).strip() or 'document'
                pdf_filename = f"{safe_filename}_{int(time.time())}.pdf"
            else:
                pdf_filename = f"md2pdf_{int(time.time())}.pdf"
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            task_id = hashlib.md5(f"convert_{time.time()}".encode()).hexdigest()[:16]
            task_data = {'status': 'processing', 'result': None, 'error': None, 'created_at': time.time()}
            background_tasks[task_id] = task_data
            save_task_meta(task_id, task_data)

            def _do_convert_pdf():
                try:
                    render_pdf(html_path, pdf_path, wait_selector=wait_selector, extra_wait_ms=extra_wait)
                    background_tasks[task_id]['result'] = {'filename': pdf_filename}
                    background_tasks[task_id]['status'] = 'done'
                    save_task_meta(task_id, background_tasks[task_id])
                except Exception as e:
                    error_detail = str(e) if str(e) else f'{type(e).__name__}'
                    logger.error(f"PDF转换失败: {traceback.format_exc()}")
                    background_tasks[task_id]['error'] = error_detail
                    background_tasks[task_id]['status'] = 'error'
                    save_task_meta(task_id, background_tasks[task_id])

            thread = threading.Thread(target=_do_convert_pdf, daemon=True)
            thread.start()
            return jsonify({'status': 'success', 'data': {'task_id': task_id}})
        except Exception as e:
            logger.error(f"PDF生成失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/upload', methods=['POST'])
    def api_upload():
        if 'file' not in request.files:
            return jsonify({'error': '请选择文件'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择文件'}), 400
        watermark = request.form.get('watermark', '')
        orig_name = os.path.splitext(file.filename)[0]
        if not file.filename.lower().endswith('.docx'):
            return jsonify({'error': '只支持Word文件(.docx)'}), 400
        try:
            file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"word_{file_id}.docx")
            file.save(file_path)
            from docx2pdf import convert
            pdf_filename = f"{orig_name}_{int(time.time())}.pdf"
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            convert(file_path, pdf_path)
            return jsonify({'filename': pdf_filename, 'original_name': orig_name})
        except ImportError:
            try:
                from docx import Document
                doc = Document(file_path)
                html_content = ''
                for para in doc.paragraphs:
                    html_content += f'<p>{para.text}</p>'
                for table in doc.tables:
                    html_content += '<table border="1">'
                    for row in table.rows:
                        html_content += '<tr>'
                        for cell in row.cells:
                            html_content += f'<td>{cell.text}</td>'
                        html_content += '</tr>'
                    html_content += '</table>'
                watermark_html = ''
                if watermark:
                    watermark_html = f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">{watermark}</div>'''
                import tempfile as tf
                with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                    f.write(f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>body {{ font-family: -apple-system, "Segoe UI", Roboto, "PingFang SC", "Microsoft YaHei", sans-serif; padding: 40px; line-height: 1.8; }} table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }} td {{ border: 1px solid #ddd; padding: 8px 12px; }}</style></head><body>{html_content}{watermark_html}</body></html>''')
                    html_path = f.name
                pdf_filename = f"{orig_name}_{int(time.time())}.pdf"
                pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
                render_pdf(html_path, pdf_path)
                return jsonify({'filename': pdf_filename, 'original_name': orig_name})
            except Exception as e2:
                logger.error(f"Word转PDF失败: {traceback.format_exc()}")
                return jsonify({'error': f'Word转PDF失败: {str(e2)}'}), 500
        except Exception as e:
            logger.error(f"Word上传失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/excel-parse', methods=['POST'])
    def api_excel_parse():
        if 'file' not in request.files:
            return jsonify({'error': '请选择文件'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择文件'}), 400
        filename_lower = file.filename.lower()
        if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
            return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400
        try:
            file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
            orig_ext = os.path.splitext(file.filename)[1].lower()
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"excel_parse_{file_id}{orig_ext}")
            file.save(file_path)
            reader = ExcelReader(file_path)
            reader.open()
            sheet_names = reader.get_sheet_names()
            sheets_data = []
            for sheet_name in sheet_names:
                try:
                    rows = reader.get_sheet_data(sheet_name)
                    if not rows:
                        sheets_data.append({'name': sheet_name, 'row_count': 0, 'column_count': 0, 'headers': [], 'data_preview': [], 'rows': [], 'summary': '空工作表', 'categories': {}})
                        continue
                    headers = [str(c).strip() if c else '' for c in rows[0]]
                    data_rows = rows[1:] if len(rows) > 1 else []
                    col_count = len(headers)
                    row_count = len(data_rows)
                    data_preview = []
                    for row in data_rows[:5]:
                        cells = [str(c).strip() if c else '' for c in row]
                        while len(cells) < col_count:
                            cells.append('')
                        data_preview.append(cells[:col_count])
                    all_rows = []
                    for row in data_rows:
                        cells = [str(c).strip() if c else '' for c in row]
                        while len(cells) < col_count:
                            cells.append('')
                        all_rows.append(cells[:col_count])
                    categories = _categorize_headers(headers)
                    summary_parts = []
                    if row_count > 0:
                        summary_parts.append(f'{row_count}行数据')
                    non_empty_headers = [h for h in headers if h]
                    if non_empty_headers:
                        summary_parts.append(f'{len(non_empty_headers)}列字段')
                    date_cols = [h for h in headers if any(kw in h.lower() for kw in ['date', '日期', '时间', 'created', 'updated', 'resolved'])]
                    if date_cols:
                        summary_parts.append('含时间字段')
                    numeric_count = 0
                    for col_idx in range(min(col_count, 20)):
                        for row in data_rows[:10]:
                            cells = [str(c).strip() if c else '' for c in row]
                            if col_idx < len(cells) and cells[col_idx]:
                                try:
                                    float(cells[col_idx].replace(',', ''))
                                    numeric_count += 1
                                    break
                                except (ValueError, IndexError):
                                    pass
                    if numeric_count > col_count * 0.3:
                        summary_parts.append('数值型数据')
                    status_cols = [h for h in headers if any(kw in h.lower() for kw in ['status', '状态', 'type', '类型', 'category', '分类'])]
                    if status_cols:
                        summary_parts.append('含状态字段')
                    summary = ' · '.join(summary_parts) if summary_parts else f'{row_count}行 x {col_count}列'
                    max_cols = min(col_count, 100)
                    sheets_data.append({'name': sheet_name, 'row_count': row_count, 'column_count': max_cols, 'headers': headers[:max_cols], 'data_preview': [row[:max_cols] for row in data_preview], 'rows': [row[:max_cols] for row in all_rows], 'summary': summary, 'categories': categories})
                except Exception as e:
                    logger.warning(f"解析sheet {sheet_name} 失败: {e}")
                    sheets_data.append({'name': sheet_name, 'row_count': 0, 'column_count': 0, 'headers': [], 'data_preview': [], 'summary': f'解析失败: {str(e)}', 'categories': {}})
            reader.close()
            return jsonify({'status': 'success', 'data': {'file_name': file.filename, 'total_sheets': len(sheet_names), 'sheets': sheets_data}})
        except Exception as e:
            logger.error(f"Excel解析失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/excel-organize', methods=['POST'])
    def api_excel_organize():
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({'error': '请求数据格式错误'}), 400
        structured_data = data.get('structured_data', {})
        user_request = data.get('user_request', '')
        selected_sheets = data.get('selected_sheets', [])
        if not structured_data or not user_request:
            return jsonify({'error': '缺少必要参数'}), 400
        try:
            ai_config = get_ai_config()
            sheets = structured_data.get('sheets', [])
            if selected_sheets:
                sheets = [s for s in sheets if s.get('name') in selected_sheets]
            if not sheets:
                return jsonify({'error': '没有选中的工作表数据'}), 400
            data_summary = _prepare_excel_summary(structured_data, sheets)
            if ai_config.get('enabled'):
                organized = _ai_organize_excel(data_summary, user_request, ai_config)
            else:
                organized = _local_organize_excel(data_summary, user_request)
            return jsonify({'status': 'success', 'data': organized})
        except Exception as e:
            logger.error(f"Excel整理失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/excel-organize-pdf', methods=['POST'])
    def api_excel_organize_pdf():
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({'error': '请求数据格式错误'}), 400
        organized_data = data.get('organized_data', {})
        watermark = data.get('watermark', '')
        if not organized_data:
            return jsonify({'error': '缺少整理后的数据'}), 400
        try:
            sections = organized_data.get('sections', [])
            summary = organized_data.get('summary', '')
            user_request = organized_data.get('user_request', '')
            html_content = _build_excel_report_html(sections, summary, user_request, watermark)
            import tempfile as tf
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            pdf_filename = f"excel_report_{int(time.time())}.pdf"
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            render_pdf(html_path, pdf_path, margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'}, extra_wait_ms=1000)
            try:
                os.unlink(html_path)
            except Exception:
                pass
            return jsonify({'status': 'success', 'filename': pdf_filename})
        except Exception as e:
            logger.error(f"Excel PDF生成失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/excel-pdf', methods=['POST'])
    def api_excel_pdf():
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({'error': '请求数据格式错误'}), 400
        structured_data = data.get('structured_data', {})
        selected_sheets = data.get('selected_sheets', [])
        watermark = data.get('watermark', '')
        custom_title = data.get('custom_title', '').strip()
        if not structured_data:
            return jsonify({'error': '缺少结构化数据'}), 400
        try:
            html_content = _build_excel_structured_report_html(structured_data, selected_sheets, watermark, custom_title)
            import tempfile as tf
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            if custom_title:
                safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
                pdf_filename = f"{safe_title}_{int(time.time())}.pdf"
                download_name = f"{safe_title}.pdf"
            else:
                pdf_filename = f"excel_pdf_{int(time.time())}.pdf"
                download_name = pdf_filename
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            render_pdf(html_path, pdf_path, margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'}, extra_wait_ms=1500)
            try:
                os.unlink(html_path)
            except Exception:
                pass
            return jsonify({'status': 'success', 'filename': pdf_filename, 'download_name': download_name})
        except Exception as e:
            logger.error(f"Excel PDF生成失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    @bp.route('/excel-select-pdf', methods=['POST'])
    def api_excel_select_pdf():
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({'error': '请求数据格式错误'}), 400
        structured_data = data.get('structured_data', {})
        selected_data = data.get('selected_data', {})
        selected_columns = data.get('selected_columns', {})
        watermark = data.get('watermark', '')
        custom_title = data.get('custom_title', '').strip()
        if not selected_data:
            return jsonify({'error': '缺少选中数据'}), 400
        try:
            html_content = _build_excel_selected_report_html(structured_data, selected_data, selected_columns, watermark, custom_title)
            import tempfile as tf
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            if custom_title:
                safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
                pdf_filename = f"{safe_title}_{int(time.time())}.pdf"
                download_name = f"{safe_title}.pdf"
            else:
                pdf_filename = f"excel_select_pdf_{int(time.time())}.pdf"
                download_name = pdf_filename
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            render_pdf(html_path, pdf_path, margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'}, extra_wait_ms=1500)
            try:
                os.unlink(html_path)
            except Exception:
                pass
            return jsonify({'status': 'success', 'filename': pdf_filename, 'download_name': download_name})
        except Exception as e:
            logger.error(f"Excel select PDF生成失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    # ==================== Excel AI 分析 ====================

    @bp.route('/api/excel-analyze-ai', methods=['POST'])
    @login_required_or_guest
    def api_excel_analyze_ai():
        """v3.0: AI 增强 CR 问题分析"""
        ai_config = get_ai_config()
        if not ai_config.get('enabled'):
            return jsonify({'error': 'AI功能未配置，请在设置页面配置 API Key'}), 503
        data = request.get_json(silent=True) or {}
        analysis = data.get('analysis', {})
        if not analysis:
            return jsonify({'error': '缺少分析数据'}), 400
        summary = analysis.get('summary', {})
        module_stats = analysis.get('module_stats', {})
        dev_stats = analysis.get('dev_stats', {})
        daily_stats = analysis.get('daily_stats', [])

        # 模块统计
        modules_list = []
        if isinstance(module_stats, dict):
            for mod, stats in module_stats.items():
                if isinstance(stats, dict):
                    total = stats.get('total', 0)
                    resolved = stats.get('resolved', 0)
                    unresolved = stats.get('unresolved', 0)
                    rate = f"{(resolved/total*100):.1f}%" if total > 0 else "N/A"
                    modules_list.append({'name': mod, 'total': total, 'resolved': resolved, 'unresolved': unresolved, 'rate': rate})
        modules_list.sort(key=lambda x: x['total'], reverse=True)
        modules_text = ''
        for m in modules_list[:10]:
            modules_text += f"- {m['name']}: {m['total']}个问题 (已解决{m['resolved']}, 未解决{m['unresolved']}, 解决率{m['rate']})\n"

        # 研发统计
        devs_list = []
        if isinstance(dev_stats, dict):
            for dev, stats in dev_stats.items():
                if isinstance(stats, dict):
                    total = stats.get('total', 0)
                    resolved = stats.get('resolved', 0)
                    unresolved = stats.get('unresolved', 0)
                    rate = (resolved/total*100) if total > 0 else 0
                    devs_list.append({'name': dev, 'total': total, 'resolved': resolved, 'unresolved': unresolved, 'rate': rate})
        devs_list.sort(key=lambda x: x['total'], reverse=True)
        devs_text = ''
        for d in devs_list[:10]:
            efficiency = "高效" if d['rate'] >= 80 else ("中等" if d['rate'] >= 50 else "待提升")
            devs_text += f"- {d['name']}: {d['total']}个问题 (已解决{d['resolved']}, 未解决{d['unresolved']}, 解决率{d['rate']:.1f}%, 效率:{efficiency})\n"

        devs_by_rate = sorted(devs_list, key=lambda x: x['rate'], reverse=True)
        high_eff = [d['name'] for d in devs_by_rate[:3] if d['total'] >= 3]
        low_eff = [d['name'] for d in devs_by_rate[-3:] if d['total'] >= 3]

        # 负责人Top
        dev_top_text = ''
        for d in devs_list[:5]:
            dev_top_text += f"{d['name']} {d['total']}、"
        dev_top_text = dev_top_text.rstrip('、')

        # 重灾模块
        modules_by_unresolved = sorted(modules_list, key=lambda x: x['unresolved'], reverse=True)
        hot_modules_text = ''
        for m in modules_by_unresolved[:5]:
            hot_modules_text += f"{m['name']} {m['unresolved']}、"
        hot_modules_text = hot_modules_text.rstrip('、')

        # Blocker统计
        blocker_total = summary.get('blocker_total', 0)
        blocker_unresolved = summary.get('blocker_unresolved', 0)
        blocker_unresolved_rate = summary.get('blocker_unresolved_rate', 0)

        # 状态分布
        blocker_unresolved_status = summary.get('blocker_unresolved_status_dist', {})
        status_text = ''
        if isinstance(blocker_unresolved_status, dict) and blocker_unresolved_status:
            for status, count in sorted(blocker_unresolved_status.items(), key=lambda x: x[1], reverse=True):
                status_text += f"{status} {count} / "
            status_text = status_text.rstrip(' / ')

        prompt = f"""你是一位资深智能硬件质量管理专家，请基于以下CR问题分析数据生成专业分析报告。

## 整体概览
- 总问题数: {summary.get('total_issues', 0)}
- 已解决: {summary.get('total_resolved', 0)}
- 未解决: {summary.get('total_unresolved', 0)}
- 解决率: {summary.get('resolution_rate', 'N/A')}%

## Blocker 重点关注
- Blocker 总计: {blocker_total}
- Blocker 未关闭: {blocker_unresolved}（占比 {blocker_unresolved_rate}%）
- 未关闭状态分布: {status_text or '无数据'}

## 重灾模块（按未解决数 Top5）
{hot_modules_text or '无数据'}

## 负责人 Top（按问题数 Top5）
{dev_top_text or '无数据'}

## 模块问题分布 (Top 10)
{modules_text or '无数据'}

## 研发人员效率 (Top 10)
{devs_text or '无数据'}
高效研发: {', '.join(high_eff) if high_eff else '无'}
待提升研发: {', '.join(low_eff) if low_eff else '无'}

请输出：整体质量评估、Blocker深度分析、模块问题分析、研发效率分析、高风险影响评估、具体改进建议。必须基于真实数据，引用具体数字，禁止说数据不足。"""
        try:
            messages = [{'role': 'user', 'content': prompt}]
            reply = _call_ai(messages, max_tokens=1500, temperature=0.3, timeout=60)
            return jsonify({'status': 'success', 'analysis': reply})
        except Exception as e:
            logger.error(f'AI CR分析失败: {e}')
            return jsonify({'error': f'AI分析失败: {str(e)}'}), 502

    @bp.route('/api/excel-analyze-ai-stream', methods=['POST'])
    @login_required_or_guest
    def api_excel_analyze_ai_stream():
        """SSE 流式版：AI 增强 CR 问题分析"""
        ai_config = get_ai_config()
        if not ai_config.get('enabled'):
            return jsonify({'error': 'AI功能未配置，请在设置页面配置 API Key'}), 503
        data = request.get_json(silent=True) or {}
        analysis = data.get('analysis', {})
        if not analysis:
            return jsonify({'error': '缺少分析数据'}), 400
        summary = analysis.get('summary', {})
        module_stats = analysis.get('module_stats', {})
        dev_stats = analysis.get('dev_stats', {})
        daily_stats = analysis.get('daily_stats', [])

        # 模块统计：转列表按问题数排序取Top10
        modules_list = []
        if isinstance(module_stats, dict):
            for mod, stats in module_stats.items():
                if isinstance(stats, dict):
                    total = stats.get('total', 0)
                    resolved = stats.get('resolved', 0)
                    unresolved = stats.get('unresolved', 0)
                    rate = f"{(resolved/total*100):.1f}%" if total > 0 else "N/A"
                    modules_list.append({'name': mod, 'total': total, 'resolved': resolved, 'unresolved': unresolved, 'rate': rate})
        modules_list.sort(key=lambda x: x['total'], reverse=True)
        modules_text = ''
        for m in modules_list[:10]:
            modules_text += f"- {m['name']}: {m['total']}个问题 (已解决{m['resolved']}, 未解决{m['unresolved']}, 解决率{m['rate']})\n"

        # 研发统计：转列表按问题数排序取Top10，计算效率
        devs_list = []
        if isinstance(dev_stats, dict):
            for dev, stats in dev_stats.items():
                if isinstance(stats, dict):
                    total = stats.get('total', 0)
                    resolved = stats.get('resolved', 0)
                    unresolved = stats.get('unresolved', 0)
                    rate = (resolved/total*100) if total > 0 else 0
                    devs_list.append({'name': dev, 'total': total, 'resolved': resolved, 'unresolved': unresolved, 'rate': rate})
        devs_list.sort(key=lambda x: x['total'], reverse=True)
        devs_text = ''
        for d in devs_list[:10]:
            efficiency = "高效" if d['rate'] >= 80 else ("中等" if d['rate'] >= 50 else "待提升")
            devs_text += f"- {d['name']}: {d['total']}个问题 (已解决{d['resolved']}, 未解决{d['unresolved']}, 解决率{d['rate']:.1f}%, 效率:{efficiency})\n"

        # 研发效率排名
        devs_by_rate = sorted(devs_list, key=lambda x: x['rate'], reverse=True)
        high_eff = [d['name'] for d in devs_by_rate[:3] if d['total'] >= 3]
        low_eff = [d['name'] for d in devs_by_rate[-3:] if d['total'] >= 3]

        # 负责人Top（按问题数排序）
        dev_top_text = ''
        for d in devs_list[:5]:
            dev_top_text += f"{d['name']} {d['total']}、"
        dev_top_text = dev_top_text.rstrip('、')

        # 重灾模块（按未解决数排序Top5）
        modules_by_unresolved = sorted(modules_list, key=lambda x: x['unresolved'], reverse=True)
        hot_modules_text = ''
        for m in modules_by_unresolved[:5]:
            hot_modules_text += f"{m['name']} {m['unresolved']}、"
        hot_modules_text = hot_modules_text.rstrip('、')

        # Blocker统计
        blocker_total = summary.get('blocker_total', 0)
        blocker_unresolved = summary.get('blocker_unresolved', 0)
        blocker_unresolved_rate = summary.get('blocker_unresolved_rate', 0)

        # 未关闭状态分布
        unresolved_status = summary.get('unresolved_status_dist', {})
        blocker_unresolved_status = summary.get('blocker_unresolved_status_dist', {})
        status_text = ''
        if isinstance(blocker_unresolved_status, dict) and blocker_unresolved_status:
            for status, count in sorted(blocker_unresolved_status.items(), key=lambda x: x[1], reverse=True):
                status_text += f"{status} {count} / "
            status_text = status_text.rstrip(' / ')
        elif isinstance(unresolved_status, dict) and unresolved_status:
            for status, count in sorted(unresolved_status.items(), key=lambda x: x[1], reverse=True)[:6]:
                status_text += f"{status} {count} / "
            status_text = status_text.rstrip(' / ')

        # 每日趋势（最近14天）
        daily_text = ''
        if isinstance(daily_stats, list) and len(daily_stats) > 0:
            recent = daily_stats[-14:]
            daily_text = f"最近{len(recent)}天趋势:\n"
            for d in recent:
                daily_text += f"- {d.get('date','')}: 新增{d.get('new_count',0)}, 解决{d.get('resolved_count',0)}\n"

        prompt = f"""你是一位资深智能硬件质量管理专家，请基于以下CR问题分析数据生成专业、有深度的分析报告。

## 一、整体概览
- 总问题数: {summary.get('total_issues', 0)}
- 已解决: {summary.get('total_resolved', 0)}
- 未解决: {summary.get('total_unresolved', 0)}
- 解决率: {summary.get('resolution_rate', 'N/A')}%

## 二、Blocker 重点关注
- Blocker 总计: {blocker_total}
- Blocker 未关闭: {blocker_unresolved}（占比 {blocker_unresolved_rate}%）
- 未关闭状态分布: {status_text or '无数据'}

## 三、重灾模块（按未解决数排序 Top5）
{hot_modules_text or '无数据'}

## 四、负责人 Top（按问题数排序 Top5）
{dev_top_text or '无数据'}

## 五、模块问题分布 (Top 10，按问题数排序)
{modules_text or '无数据'}

## 六、研发人员效率分析 (Top 10，按问题数排序)
{devs_text or '无数据'}

**研发效率排名:**
- 高效研发(解决率≥80%且问题数≥3): {', '.join(high_eff) if high_eff else '无'}
- 待提升研发(解决率低且问题数≥3): {', '.join(low_eff) if low_eff else '无'}

## 七、每日趋势
{daily_text or '无数据'}

---

请按以下结构输出深度分析，**必须基于上述真实数据，禁止编造数据或说数据缺失**：

### 📊 一、整体质量评估
基于总问题数、解决率、Blocker占比，评估当前项目整体质量状态。

### 🔥 二、Blocker 深度分析
- Blocker未关闭数量和占比意味着什么风险？
- 未关闭状态分布（New/Reopened等）反映了什么问题？
- 哪些模块的Blocker最集中？需要优先投入哪些资源？

### 📦 三、模块问题分析
- 重灾模块有哪些？未解决问题集中在哪些模块？
- 模块之间是否存在关联性问题？
- 哪些模块需要重点关注和资源倾斜？

### 👥 四、研发效率分析
- 负责人Top的问题分布是否均衡？是否存在某些研发负担过重？
- 高效研发有哪些值得总结的经验？
- 待提升研发可能的原因是什么？有哪些改进建议？

### ⚠️ 五、高风险与影响评估
- Blocker未关闭对版本发布、用户体验、稳定性的影响？
- 如果不及时解决，可能导致什么后果？

### 💡 六、具体改进建议
给出可执行的改进措施，包括：
- 模块层面：哪些模块需要重点投入资源？
- 研发层面：如何优化问题分配和提升解决效率？
- 流程层面：是否需要优化问题跟踪、验证、回归流程？
- 优先级：接下来一周应该优先解决哪些Blocker？

要求：
1. 所有分析必须基于上面提供的真实数据，引用具体数字
2. 禁止使用"数据不足"、"无法判断"等推脱性表述
3. 语言专业、简洁，避免空话套话
4. 重点突出Blocker风险、模块问题和研发效率"""
        messages = [{'role': 'user', 'content': prompt}]
        return Response(
            stream_with_context(_call_ai_stream(messages, max_tokens=2500, temperature=0.3)),
            content_type='text/event-stream',
            headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no', 'Connection': 'keep-alive'}
        )

    @bp.route('/api/excel-analyze-pdf', methods=['POST'])
    def api_excel_analyze_pdf():
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({'error': '请求数据格式错误'}), 400
        analysis_data = data.get('analysis_data', {})
        watermark = data.get('watermark', '')
        custom_title = data.get('custom_title', '').strip()
        file_name = data.get('file_name', '')
        ai_analysis = data.get('ai_analysis', '')
        all_issues = data.get('all_issues', [])
        if not analysis_data:
            return jsonify({'error': '缺少分析数据'}), 400
        try:
            html_content = _build_cr_analysis_report_html(analysis_data, watermark, file_name, custom_title, ai_analysis, all_issues)
            import tempfile as tf
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            if custom_title:
                safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
                pdf_filename = f"{safe_title}_{datetime.now(_CST).strftime('%Y%m%d_%H%M%S')}.pdf"
                download_name = f"{safe_title}.pdf"
            else:
                pdf_filename = f"cr_analysis_{int(time.time())}.pdf"
                download_name = pdf_filename
            pdf_path = os.path.join(current_app.config['PDF_FOLDER'], pdf_filename)
            render_pdf(html_path, pdf_path, margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'}, extra_wait_ms=3000, wait_selector='canvas')
            try:
                os.unlink(html_path)
            except Exception:
                pass
            return jsonify({'status': 'success', 'filename': pdf_filename, 'download_name': download_name})
        except Exception as e:
            logger.error(f"CR分析PDF生成失败: {traceback.format_exc()}")
            return jsonify({'error': str(e)}), 500

    return bp


# ==================== 辅助函数（模块级） ====================

def _analyze_test_report(headers, data_rows):
    """分析测试报告数据"""
    col_map = {}
    for i, h in enumerate(headers):
        h_lower = h.lower()
        if any(kw in h_lower for kw in ['模块', 'module', 'component']):
            col_map['module'] = i
        elif any(kw in h_lower for kw in ['状态', 'status', 'result']):
            col_map['status'] = i
        elif any(kw in h_lower for kw in ['用例', 'case', '名称', 'title', 'name']):
            col_map['name'] = i
        elif any(kw in h_lower for kw in ['优先级', 'priority', '严重', 'severity']):
            col_map['priority'] = i
        elif any(kw in h_lower for kw in ['执行人', 'tester', 'owner']):
            col_map['tester'] = i
        elif any(kw in h_lower for kw in ['日期', 'date', '时间', 'time']):
            col_map['date'] = i

    total = len(data_rows)
    passed = failed = blocked = skipped = 0
    module_stats = {}
    status_stats = {}

    for row in data_rows:
        status = ''
        if 'status' in col_map and col_map['status'] < len(row):
            status = str(row[col_map['status']]).strip().lower()
        module = '未分类'
        if 'module' in col_map and col_map['module'] < len(row):
            module = str(row[col_map['module']]).strip() or '未分类'

        if any(kw in status for kw in ['pass', '通过', '成功', 'ok']):
            passed += 1
        elif any(kw in status for kw in ['fail', '失败', 'error', '错误']):
            failed += 1
        elif any(kw in status for kw in ['block', '阻塞', 'blocked']):
            blocked += 1
        elif any(kw in status for kw in ['skip', '跳过', '未执行']):
            skipped += 1
        else:
            passed += 1

        if module not in module_stats:
            module_stats[module] = {'name': module, 'total': 0, 'passed': 0, 'failed': 0, 'blocked': 0, 'skipped': 0}
        module_stats[module]['total'] += 1
        if any(kw in status for kw in ['pass', '通过', '成功', 'ok']):
            module_stats[module]['passed'] += 1
        elif any(kw in status for kw in ['fail', '失败', 'error', '错误']):
            module_stats[module]['failed'] += 1
        elif any(kw in status for kw in ['block', '阻塞', 'blocked']):
            module_stats[module]['blocked'] += 1
        elif any(kw in status for kw in ['skip', '跳过', '未执行']):
            module_stats[module]['skipped'] += 1

        if status:
            status_stats[status] = status_stats.get(status, 0) + 1

    pass_rate = f"{(passed / total * 100):.1f}%" if total > 0 else 'N/A'
    modules_sorted = sorted(module_stats.values(), key=lambda x: x['failed'], reverse=True)

    return {
        'headers': headers,
        'summary': {'total': total, 'passed': passed, 'failed': failed, 'blocked': blocked, 'skipped': skipped, 'pass_rate': pass_rate},
        'modules': modules_sorted[:20],
        'status': status_stats,
        'trend': [],
    }


def _build_test_report_html(analysis_data, watermark, file_name, custom_title=''):
    """构建测试报告HTML"""
    summary = analysis_data.get('summary', {})
    modules = analysis_data.get('modules', [])
    title = _escape_html(custom_title) if custom_title else '📊 测试报告分析'
    now = datetime.now(_CST).strftime('%Y-%m-%d %H:%M:%S')

    watermark_html = ''
    if watermark:
        watermark_html = f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">{_escape_html(watermark)}</div>'''

    modules_rows = ''
    for m in modules[:30]:
        modules_rows += f'''<tr><td>{_escape_html(m.get('name', ''))}</td><td>{m.get('total', 0)}</td><td>{m.get('passed', 0)}</td><td style="color:#e74c3c;">{m.get('failed', 0)}</td><td>{m.get('blocked', 0)}</td><td>{m.get('skipped', 0)}</td></tr>'''

    return f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>
body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif; padding: 40px; color: #1d1d1f; }}
h1 {{ font-size: 24px; font-weight: 700; text-align: center; margin-bottom: 8px; }}
.meta {{ text-align: center; color: #6e6e73; font-size: 12px; margin-bottom: 24px; }}
.summary-box {{ display: flex; justify-content: center; gap: 24px; margin-bottom: 32px; flex-wrap: wrap; }}
.stat-card {{ background: #f5f5f7; border-radius: 12px; padding: 16px 24px; text-align: center; min-width: 100px; }}
.stat-num {{ font-size: 28px; font-weight: 700; }}
.stat-label {{ font-size: 12px; color: #6e6e73; margin-top: 4px; }}
table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }}
th {{ background: #1d1d1f; color: white; padding: 10px 12px; text-align: left; }}
td {{ border: 1px solid #e5e5ea; padding: 8px 12px; }}
h2 {{ font-size: 18px; margin-top: 32px; border-bottom: 2px solid #0071e3; padding-bottom: 8px; }}
</style></head><body>
{watermark_html}
<h1>{title}</h1>
<div class="meta">文件: {_escape_html(file_name)} | 生成时间: {now}</div>
<div class="summary-box">
<div class="stat-card"><div class="stat-num">{summary.get('total', 0)}</div><div class="stat-label">总用例</div></div>
<div class="stat-card"><div class="stat-num" style="color:#34c759;">{summary.get('passed', 0)}</div><div class="stat-label">通过</div></div>
<div class="stat-card"><div class="stat-num" style="color:#e74c3c;">{summary.get('failed', 0)}</div><div class="stat-label">失败</div></div>
<div class="stat-card"><div class="stat-num" style="color:#ff9500;">{summary.get('blocked', 0)}</div><div class="stat-label">阻塞</div></div>
<div class="stat-card"><div class="stat-num" style="color:#8e8e93;">{summary.get('skipped', 0)}</div><div class="stat-label">跳过</div></div>
<div class="stat-card"><div class="stat-num">{summary.get('pass_rate', 'N/A')}</div><div class="stat-label">通过率</div></div>
</div>
<h2>📋 模块统计</h2>
<table><thead><tr><th>模块</th><th>总数</th><th>通过</th><th>失败</th><th>阻塞</th><th>跳过</th></tr></thead><tbody>{modules_rows}</tbody></table>
</body></html>'''


def _categorize_headers(headers):
    """将表头分类"""
    categories = {}
    category_rules = {
        '基本信息': ['id', 'key', '编号', '名称', 'name', 'title', '标题', 'summary', '描述', 'description'],
        '状态信息': ['status', '状态', 'state', 'resolved', '解决', 'closed', '关闭'],
        '时间信息': ['date', '日期', '时间', 'created', 'updated', 'resolved', 'due'],
        '人员信息': ['assignee', 'developer', '负责人', '研发', 'reporter', '报告人', 'creator', '创建人'],
        '优先级': ['priority', '优先级', 'severity', '严重性', 'criticality'],
        '模块/组件': ['component', 'module', '模块', '组件', 'project', '项目'],
        '版本信息': ['version', '版本', 'fix version', 'affected version'],
        '数值指标': ['count', 'total', '金额', 'cost', 'price', 'rate', '比率', 'percentage'],
    }
    for header in headers:
        h_lower = header.lower().strip()
        if not h_lower:
            continue
        for category, keywords in category_rules.items():
            if any(kw in h_lower for kw in keywords):
                if category not in categories:
                    categories[category] = []
                categories[category].append({'name': header})
                break
    return categories


def _prepare_excel_summary(structured_data, sheets):
    """准备Excel数据摘要"""
    summary = {'file_name': structured_data.get('file_name', ''), 'sheets': []}
    for sheet in sheets:
        sheet_info = {'name': sheet.get('name', ''), 'row_count': sheet.get('row_count', 0), 'column_count': sheet.get('column_count', 0), 'headers': sheet.get('headers', []), 'data_preview': sheet.get('data_preview', [])[:10]}
        summary['sheets'].append(sheet_info)
    return summary


def _ai_organize_excel(data_summary, user_request, ai_config):
    """使用AI整理Excel数据"""
    try:
        import requests as req
        sheets_text = ''
        for s in data_summary.get('sheets', []):
            headers_str = ', '.join(s.get('headers', [])[:20])
            preview_str = ''
            for row in s.get('data_preview', [])[:3]:
                preview_str += ' | '.join([str(c)[:30] for c in row[:10]]) + '\n'
            sheets_text += f"\n工作表「{s['name']}」({s['row_count']}行x{s['column_count']}列):\n列: {headers_str}\n数据预览:\n{preview_str}"
        prompt = f"""你是一个数据分析师。请根据用户需求整理以下Excel数据。

文件: {data_summary.get('file_name', '')}
{sheets_text}

用户需求: {user_request}

请以JSON格式返回:
{{"summary": "数据总览总结（2-3句话）", "sections": [{{"title": "章节标题", "content": "详细内容，使用HTML格式", "table": [["表头1", "表头2"], ["数据1", "数据2"]]}}]}}

要求: 根据用户需求提取关键信息，使用HTML格式化输出，表格数据以二维数组形式提供，内容要简洁明了"""
        response = req.post(
            f"{ai_config.get('base_url', 'https://dashscope.aliyuncs.com/api/v1')}/services/aigc/text-generation/generation",
            headers={'Authorization': f'Bearer {ai_config.get("api_key", "")}', 'Content-Type': 'application/json'},
            json={'model': ai_config.get('model', 'qwen-turbo'), 'input': {'messages': [{'role': 'system', 'content': '你是一个专业的数据分析师。'}, {'role': 'user', 'content': prompt}]}, 'parameters': {'result_format': 'message', 'max_tokens': 2000}},
            timeout=30
        )
        if response.status_code == 200:
            result = response.json()
            output_text = result.get('output', {}).get('choices', [{}])[0].get('message', {}).get('content', '')
            try:
                json_str = output_text
                if '```json' in json_str:
                    json_str = json_str.split('```json')[1].split('```')[0]
                elif '```' in json_str:
                    json_str = json_str.split('```')[1].split('```')[0]
                json_str = json_str.strip()
                parsed = json.loads(json_str)
                return _build_organized_result(parsed)
            except (json.JSONDecodeError, IndexError):
                return _build_text_result(output_text, user_request)
        else:
            logger.warning(f"AI请求失败: {response.status_code}")
            return _local_organize_excel(data_summary, user_request)
    except Exception as e:
        logger.warning(f"AI整理失败，使用本地整理: {e}")
        return _local_organize_excel(data_summary, user_request)


def _local_organize_excel(data_summary, user_request):
    """本地整理Excel数据"""
    sheets = data_summary.get('sheets', [])
    sections = []
    summary_text = f"文件「{data_summary.get('file_name', '')}」共包含 {len(sheets)} 个工作表。"
    for sheet in sheets:
        headers = sheet.get('headers', [])
        data_preview = sheet.get('data_preview', [])
        row_count = sheet.get('row_count', 0)
        table_data = [headers[:10]]
        for row in data_preview[:10]:
            table_data.append([str(c)[:50] for c in row[:10]])
        numeric_cols = []
        for col_idx in range(min(len(headers), 10)):
            numeric_count = 0
            for row in data_preview:
                if col_idx < len(row) and row[col_idx]:
                    try:
                        float(str(row[col_idx]).replace(',', ''))
                        numeric_count += 1
                    except ValueError:
                        pass
            if numeric_count > len(data_preview) * 0.3:
                numeric_cols.append(headers[col_idx])
        section_content = f"<p>工作表「{sheet.get('name', '')}」包含 <strong>{row_count}</strong> 行数据，<strong>{len(headers)}</strong> 列字段。</p>"
        if numeric_cols:
            section_content += f"<p>主要数值字段: {', '.join(numeric_cols)}</p>"
        section_content += "<h4>数据预览</h4>"
        sections.append({'title': f"📋 {sheet.get('name', '')} ({row_count}行)", 'content': section_content, 'table': table_data})
    sections.insert(0, {'title': '📊 数据总览', 'content': f"<p>{summary_text}</p><p><strong>用户需求:</strong> {user_request}</p><p>请查看以下各工作表的详细数据预览。</p>", 'table': []})
    return {'summary': summary_text, 'sections': sections}


def _build_organized_result(parsed):
    """构建整理后的结果"""
    sections = parsed.get('sections', [])
    html_sections = []
    for section in sections:
        content = section.get('content', '')
        table = section.get('table', [])
        if table and len(table) > 1:
            headers = table[0] if table else []
            rows = table[1:] if len(table) > 1 else []
            table_html = '<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px;"><thead><tr>'
            for h in headers:
                table_html += f'<th style="background:#1d1d1f;color:white;padding:10px 12px;text-align:left;font-weight:600;font-size:12px;">{h}</th>'
            table_html += '</tr></thead><tbody>'
            for row in rows:
                table_html += '<tr>'
                for cell in row:
                    table_html += f'<td style="padding:8px 12px;border-bottom:1px solid #e5e5ea;">{cell}</td>'
                table_html += '</tr>'
            table_html += '</tbody></table>'
            content += table_html
        html_sections.append({'title': section.get('title', ''), 'content': content, 'table': table})
    return {'summary': parsed.get('summary', ''), 'sections': html_sections}


def _build_text_result(text, user_request):
    """从纯文本构建结果"""
    sections = [{'title': '📊 AI分析结果', 'content': text.replace('\n', '<br>'), 'table': []}]
    return {'summary': f'根据需求「{user_request}」生成的分析结果', 'sections': sections}


def _build_excel_report_html(sections, summary, user_request, watermark):
    """构建Excel报告HTML"""
    sections_html = ''
    for section in sections:
        title = section.get('title', '')
        content = section.get('content', '')
        sections_html += f'''<div style="margin-bottom:32px;"><h2 style="font-size:18px;font-weight:700;margin-bottom:16px;color:#1d1d1f;border-bottom:2px solid #0071e3;padding-bottom:8px;">{title}</h2><div style="font-size:13px;line-height:1.8;color:#3c3c43;">{content}</div></div>'''
    watermark_html = ''
    if watermark:
        watermark_html = f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">{watermark}</div>'''
    return f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>
body {{ font-family: -apple-system, "SF Pro Display", "PingFang SC", "Microsoft YaHei", "Segoe UI", Roboto, sans-serif; padding: 40px; line-height: 1.6; max-width: 900px; margin: 0 auto; }}
h1 {{ font-size: 24px; font-weight: 700; margin-bottom: 8px; color: #1d1d1f; }}
table {{ border-collapse: collapse; width: 100%; margin: 1em 0; font-size: 13px; }}
th {{ background: #1d1d1f; color: white; padding: 10px 12px; text-align: left; font-weight: 600; font-size: 12px; }}
td {{ border: 1px solid #e5e5ea; padding: 8px 12px; }}
tr:nth-child(even) td {{ background: #f5f5f7; }}
.header {{ text-align: center; margin-bottom: 32px; padding-bottom: 20px; border-bottom: 1px solid #e5e5ea; }}
.summary-box {{ background: linear-gradient(135deg, #f0f7ff, #e8f1ff); border: 1px solid #bae0ff; border-radius: 12px; padding: 16px 20px; margin-bottom: 24px; font-size: 13px; color: #3c3c43; }}
.user-request {{ font-size: 12px; color: #6e6e73; margin-top: 8px; font-style: italic; }}
</style></head><body>
{watermark_html}
<div class="header"><h1>📊 数据分析报告</h1></div>
<div class="summary-box"><strong>📋 分析摘要:</strong> {summary}{f'<div class="user-request">💡 需求: {user_request}</div>' if user_request else ''}</div>
{sections_html}
</body></html>'''


def _build_excel_structured_report_html(structured_data, selected_sheets, watermark, custom_title=''):
    """构建结构化数据报告HTML"""
    watermark_html = ''
    if watermark:
        watermark_items = ''.join([f'<div style="position:absolute;left:{x}px;top:{y}px;transform:rotate(-30deg);font-size:12px;font-weight:400;color:#0071e3;white-space:nowrap;">{_escape_html(watermark)}</div>' for x in range(80, 1200, 200) for y in range(150, 1000, 200)])
        watermark_html = f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.08;overflow:hidden;">{watermark_items}</div>'''
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    title = _escape_html(custom_title) if custom_title else '📊 数据分析报告'
    sheets_data = structured_data.get('sheets', []) if isinstance(structured_data, dict) else []
    sheet_map = {}
    max_cols = 0
    for sheet in sheets_data:
        if isinstance(sheet, dict):
            name = sheet.get('name', '')
            if name:
                sheet_map[name] = sheet
                col_count = len(sheet.get('headers', []))
                if col_count > max_cols:
                    max_cols = col_count
    use_landscape = max_cols > 15
    page_size = 'A4 landscape' if use_landscape else 'A4'
    sheets_to_show = selected_sheets if selected_sheets else list(sheet_map.keys())
    sheets_html = ''
    for sheet_name in sheets_to_show:
        sheet_data = sheet_map.get(sheet_name, {})
        if not sheet_data:
            continue
        headers = sheet_data.get('headers', [])
        rows = sheet_data.get('rows', [])
        if not headers:
            continue
        if not rows:
            rows = sheet_data.get('data_preview', [])
        if not rows:
            continue
        num_cols = len(headers)
        aligned_headers = [_escape_html(h) for h in headers]
        if num_cols > 50:
            font_size, padding = '6px', '2px 3px'
        elif num_cols > 30:
            font_size, padding = '8px', '3px 4px'
        elif num_cols > 15:
            font_size, padding = '9px', '4px 5px'
        else:
            font_size, padding = '10px', '5px 8px'
        header_html = ''.join([f'<th style="padding:{padding};font-size:{font_size};">{h}</th>' for h in aligned_headers])
        rows_html = ''
        display_rows = rows[:50]
        for row in display_rows:
            aligned_row = list(row[:num_cols]) + [''] * max(0, num_cols - len(row))
            cells = ''.join([f'<td style="border:1px solid #e5e5ea;padding:{padding};font-size:{font_size};">{_escape_html(cell)}</td>' for cell in aligned_row])
            rows_html += f'<tr>{cells}</tr>'
        total_rows = len(rows)
        sheets_html += f'''<div style="margin-bottom:20px;break-inside:avoid;"><h2 style="font-size:14px;font-weight:600;margin-bottom:10px;color:#1d1d1f;padding:6px 10px;background:#f5f5f7;border-radius:6px;border-left:3px solid #0071e3;">📋 {_escape_html(sheet_name)} ({num_cols}列)</h2><div style="overflow-x:auto;border:1px solid #e5e5ea;border-radius:6px;"><table style="border-collapse:collapse;width:100%;"><thead><tr style="background:#1d1d1f;color:white;">{header_html}</tr></thead><tbody>{rows_html}</tbody></table></div>{f'<p style="font-size:9px;color:#6e6e73;margin-top:4px;text-align:right;">共 {total_rows} 行数据，仅显示前50行</p>' if total_rows > 50 else ''}</div>'''
    if not sheets_html:
        sheets_html = '<p style="text-align:center;color:#6e6e73;padding:40px;">暂无数据</p>'
    return f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>
@page {{ size: {page_size}; margin: 10mm 8mm; }}
body {{ font-family: -apple-system, "SF Pro Display", "PingFang SC", "Microsoft YaHei", "Segoe UI", Roboto, sans-serif; padding: 5px; line-height: 1.3; max-width: 100%; margin: 0 auto; color: #1d1d1f; }}
table {{ border-collapse: collapse; width: 100%; margin: 0; }}
th {{ background: #1d1d1f; color: white; text-align: left; font-weight: 600; }}
td {{ border: 1px solid #e5e5ea; }}
h1 {{ font-size: 20px; font-weight: 700; margin: 0 0 6px 0; color: #1d1d1f; }}
h2 {{ font-size: 14px; font-weight: 600; margin-bottom: 10px; }}
.header-box {{ text-align:center; margin-bottom:16px; padding-bottom:12px; border-bottom:2px solid #e5e5ea; }}
.meta-info {{ font-size: 9px; color:#6e6e73; margin-top:4px; }}
</style></head><body>
{watermark_html}
<div class="header-box"><h1>{title}</h1><div class="meta-info">生成时间: {now}</div></div>
{sheets_html}
</body></html>'''


def _build_excel_selected_report_html(structured_data, selected_data, selected_columns, watermark, custom_title=''):
    """构建选中数据报告HTML"""
    watermark_html = ''
    if watermark:
        watermark_items = ''.join([f'<div style="position:absolute;left:{x}px;top:{y}px;transform:rotate(-30deg);font-size:12px;font-weight:400;color:#0071e3;white-space:nowrap;">{_escape_html(watermark)}</div>' for x in range(80, 1200, 200) for y in range(150, 1000, 200)])
        watermark_html = f'''<div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.08;overflow:hidden;">{watermark_items}</div>'''
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    title = _escape_html(custom_title) if custom_title else '📊 数据分析报告'
    sheets_data = structured_data.get('sheets', []) if isinstance(structured_data, dict) else []
    sheet_map = {}
    for sheet in sheets_data:
        if isinstance(sheet, dict):
            name = sheet.get('name', '')
            if name:
                sheet_map[name] = sheet
    max_cols = 0
    for sheet_name, row_indices in selected_data.items():
        col_indices = selected_columns.get(sheet_name, [])
        if len(col_indices) > max_cols:
            max_cols = len(col_indices)
    use_landscape = max_cols > 15
    page_size = 'A4 landscape' if use_landscape else 'A4'
    tables_html = ''
    for sheet_name, row_indices in selected_data.items():
        sheet_data = sheet_map.get(sheet_name, {})
        if not sheet_data:
            continue
        headers = sheet_data.get('headers', [])
        rows = sheet_data.get('rows', [])
        if not rows:
            rows = sheet_data.get('data_preview', [])
        if not headers or not rows:
            continue
        col_indices = selected_columns.get(sheet_name, list(range(len(headers))))
        filtered_headers = [_escape_html(headers[i]) if i < len(headers) else '' for i in col_indices]
        num_filtered_cols = len(filtered_headers)
        filtered_rows = []
        for row_idx in row_indices:
            if row_idx < len(rows):
                row = rows[row_idx]
                filtered_row = [_escape_html(row[i]) if i < len(row) else '' for i in col_indices]
                filtered_row = filtered_row[:num_filtered_cols] + [''] * max(0, num_filtered_cols - len(filtered_row))
                filtered_rows.append(filtered_row)
        if not filtered_rows:
            continue
        if num_filtered_cols > 50:
            font_size, padding = '6px', '2px 3px'
        elif num_filtered_cols > 30:
            font_size, padding = '8px', '3px 4px'
        elif num_filtered_cols > 15:
            font_size, padding = '9px', '4px 5px'
        else:
            font_size, padding = '10px', '5px 8px'
        header_html = ''.join([f'<th style="padding:{padding};font-size:{font_size};">{h}</th>' for h in filtered_headers])
        rows_html = ''
        display_rows = filtered_rows[:50]
        for row in display_rows:
            cells = ''.join([f'<td style="border:1px solid #e5e5ea;padding:{padding};font-size:{font_size};">{c}</td>' for c in row])
            rows_html += f'<tr>{cells}</tr>'
        total_filtered = len(filtered_rows)
        tables_html += f'''<div style="margin-bottom:20px;break-inside:avoid;"><h2 style="font-size:14px;font-weight:600;margin-bottom:10px;color:#1d1d1f;padding:6px 10px;background:#f5f5f7;border-radius:6px;border-left:3px solid #0071e3;">📋 {_escape_html(sheet_name)} ({num_filtered_cols}列)</h2><div style="overflow-x:auto;border:1px solid #e5e5ea;border-radius:6px;"><table style="border-collapse:collapse;width:100%;"><thead><tr style="background:#1d1d1f;color:white;">{header_html}</tr></thead><tbody>{rows_html}</tbody></table></div>{f'<p style="font-size:9px;color:#6e6e73;margin-top:4px;text-align:right;">共 {total_filtered} 行数据，仅显示前50行</p>' if total_filtered > 50 else ''}</div>'''
    if not tables_html:
        tables_html = '<p style="text-align:center;color:#6e6e73;padding:40px;">暂无选中数据</p>'
    return f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>
@page {{ size: {page_size}; margin: 10mm 8mm; }}
body {{ font-family: -apple-system, "SF Pro Display", "PingFang SC", "Microsoft YaHei", "Segoe UI", Roboto, sans-serif; padding: 5px; line-height: 1.3; max-width: 100%; margin: 0 auto; color: #1d1d1f; }}
table {{ border-collapse: collapse; width: 100%; margin: 0; }}
th {{ background: #1d1d1f; color: white; text-align: left; font-weight: 600; }}
td {{ border: 1px solid #e5e5ea; }}
h1 {{ font-size: 20px; font-weight: 700; margin: 0 0 6px 0; color: #1d1d1f; }}
h2 {{ font-size: 14px; font-weight: 600; margin-bottom: 10px; }}
.header-box {{ text-align:center; margin-bottom:16px; padding-bottom:12px; border-bottom:2px solid #e5e5ea; }}
.meta-info {{ font-size: 9px; color:#6e6e73; margin-top:4px; }}
</style></head><body>
{watermark_html}
<div class="header-box"><h1>{title}</h1><div class="meta-info">生成时间: {now}</div></div>
{tables_html}
</body></html>'''
