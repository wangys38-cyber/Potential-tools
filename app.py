from flask import Flask, render_template, request, send_file, send_from_directory, jsonify, redirect, session, url_for, make_response, Response, stream_with_context, g
import os
import sys
import re
import logging
import traceback
import json
import tempfile
import time
import hashlib
import secrets
import gc
import requests
from datetime import datetime, timezone, timedelta
from functools import wraps
from jinja2 import BytecodeCache

# 认证模块
import auth
import db
import ai_utils
import feishu_push
import report_builders
from report_builders import _analyze_sheet_detail, _build_cr_analysis_report_html, _build_test_report_pdf_html
from routes.pages import create_pages_blueprint

# 性能优化：Whingoise直接服务静态文件，Flask-Compress启用gzip
from whitenoise import WhiteNoise
from flask_compress import Compress
# 实时语音识别：flask-sock 提供 WebSocket 支持
from flask_sock import Sock
import websocket as _ws_client  # websocket-client，连接 DashScope

import datetime as _dt

def _get_static_version():
    """静态资源版本号：优先用 git commit hash，其次用 app.py mtime，避免每分钟变化导致缓存失效"""
    _base = os.path.abspath(os.path.dirname(__file__))
    # 1. Railway 注入的 git commit
    sha = os.environ.get('RAILWAY_GIT_COMMIT_SHA', '')
    if sha:
        return sha[:8]
    # 2. 尝试读取本地 git HEAD
    try:
        import subprocess
        result = subprocess.run(['git', 'rev-parse', '--short', 'HEAD'],
                              capture_output=True, text=True, timeout=2, cwd=_base)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception:
        pass
    # 3. 回退到 app.py 修改时间（仅代码变更时才变）
    try:
        mtime = os.path.getmtime(os.path.join(_base, 'app.py'))
        return _dt.datetime.fromtimestamp(mtime).strftime('%Y%m%d%H%M')
    except Exception:
        return _dt.datetime.now().strftime('%Y%m%d%H%M')

_STATIC_VERSION = _get_static_version()

_CST = timezone(timedelta(hours=8))

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

if getattr(sys, 'frozen', False):
    base_dir = sys._MEIPASS
else:
    base_dir = os.path.abspath(os.path.dirname(__file__))

template_dir = os.path.join(base_dir, 'templates')
static_dir = os.path.join(base_dir, 'static')
app = Flask(__name__, template_folder=template_dir, static_folder=static_dir)

# 实时语音识别 WebSocket 支持
sock = Sock(app)

# 启用 gzip 压缩（HTML/JSON/CSS/JS 响应自动压缩，减少传输量 60-80%）
Compress(app)
app.config['COMPRESS_MIMETYPES'] = [
    'text/html', 'text/css', 'text/xml',
    'application/json', 'application/javascript',
    'application/xml', 'image/svg+xml',
]
app.config['COMPRESS_LEVEL'] = 6
app.config['COMPRESS_MIN_SIZE'] = 500  # 仅压缩大于500B的响应，避免小响应压缩开销

# 生产环境检测（需在 WhiteNoise 配置前定义）
_is_production = bool(os.environ.get('RAILWAY_STATIC_URL') or os.environ.get('PORT'))

# Whitenoise: 直接服务静态文件，跳过Flask请求处理（性能提升10倍+）
# 静态资源都带 ?v=版本号 做 cache busting，生产环境可安全长期缓存
app.wsgi_app = WhiteNoise(
    app.wsgi_app,
    root=static_dir,
    prefix='/static/',
    max_age=31536000 if _is_production else 0,  # 生产环境缓存1年，开发环境不缓存
)

# 生产环境优化：关闭模板自动重载（避免每次请求检查文件修改时间）
app.config['TEMPLATES_AUTO_RELOAD'] = not _is_production
app.config['DEBUG'] = not _is_production
app.secret_key = auth.SESSION_SECRET

# Session 配置 — 确保登录状态持久化、跨页面共享
app.config['SESSION_PERMANENT'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = _is_production  # HTTPS环境下启用Secure

# 静态文件缓存 — 生产环境长期缓存（带版本号cache busting），开发环境不缓存
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 31536000 if _is_production else 0

# Jinja2 字节码缓存 — 避免每次请求重新解析模板文件（解析速度提升5-10倍）
_jinja_cache_dir = '/dev/shm/jinja_cache' if _is_production else os.path.join(base_dir, '.jinja_cache')
os.makedirs(_jinja_cache_dir, exist_ok=True)

class _ShmBytecodeCache(BytecodeCache):
    """基于内存文件系统的Jinja2字节码缓存"""
    def __init__(self, directory):
        self.directory = directory
    def load_bytecode(self, bucket):
        f = os.path.join(self.directory, bucket.key)
        if os.path.exists(f):
            with open(f, 'rb') as fp:
                bucket.load_bytecode(fp)
    def dump_bytecode(self, bucket):
        f = os.path.join(self.directory, bucket.key)
        with open(f, 'wb') as fp:
            bucket.write_bytecode(fp)

app.jinja_env.bytecode_cache = _ShmBytecodeCache(_jinja_cache_dir)

# 配置 - Railway等云平台使用 /tmp 作为可写目录
if os.environ.get('RAILWAY_STATIC_URL') or os.environ.get('PORT'):
    _runtime_dir = '/tmp/toolbox'
else:
    _runtime_dir = base_dir

app.config['UPLOAD_FOLDER'] = os.path.join(_runtime_dir, 'uploads')
app.config['PDF_FOLDER'] = os.path.join(_runtime_dir, 'pdfs')
app.config['AI_CONFIG_FILE'] = os.path.join(_runtime_dir, 'ai_config.json')
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200MB - 音频文件上传限制

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['PDF_FOLDER'], exist_ok=True)

# 注入模板上下文：当前用户信息（仅对模板渲染生效）
@app.context_processor
def inject_user():
    # 快速检查 session，避免无谓的字典构造
    if not session.get('user_id'):
        return dict(current_user=None, is_logged_in=False, STATIC_VERSION=_STATIC_VERSION)
    return dict(
        current_user={
            'id': session.get('user_id'),
            'name': session.get('user_name', ''),
            'email': session.get('user_email', ''),
            'avatar': session.get('user_avatar', ''),
            'provider': session.get('user_provider', ''),
        },
        is_logged_in=True,
        STATIC_VERSION=_STATIC_VERSION,
    )


# ==================== 登录拦截 ====================
# 允许无需登录即可访问的路径前缀（按频率排序，命中即返回）
_PUBLIC_PATHS = (
    '/static/',
    '/assets/',
    '/login',
    '/auth/',
    '/health',
    '/favicon.ico',
    '/api/merit',          # v2.0: 功德查询允许匿名访问
    '/api/user/preferences', # v2.0: 偏好查询允许匿名访问
    '/api/upload-audio',     # API端点自行检查认证，避免302重定向导致JSON解析失败
    '/api/transcription-status', # 同上
    '/api/ai-models',        # 模型列表允许匿名查看（端点内部检查认证）
    '/api/ai-config',        # 同上
    '/api/ai-test',          # 同上
    '/api/ai-chat',          # AI 对话 SSE 自行检查认证
    '/api/test-report-ai-stream',  # 测试报告 AI 流式分析自行检查认证
    '/api/excel-analyze-ai-stream', # CR 分析 AI 流式自行检查认证
    '/api/generate-minutes-stream', # 会议纪要 AI 流式自行检查认证
    '/api/weekly-report-stream',    # 周报 AI 流式自行检查认证
    '/api/notes/sync',       # 笔记同步API自行检查认证
    '/api/docs',             # 文档仓库API自行检查认证
    '/api/docs/<int:doc_id>', # 文档详情API自行检查认证
    '/api/upload-init',      # 上传API自行检查认证，返回JSON 401
    '/api/upload-chunk',     # 同上
    '/api/upload-complete',  # 同上
    '/ws/',                  # WebSocket端点自行检查认证
)

@app.before_request
def require_login():
    """全局登录拦截：未登录用户自动跳转到登录页"""
    path = request.path

    # 定期清理过期任务（非阻塞，不影响请求处理）
    _maybe_cleanup()

    # 公开路径 — 最先检查，快速放行（静态文件、登录、OAuth回调等）
    for prefix in _PUBLIC_PATHS:
        if path.startswith(prefix):
            return None

    # 已登录 — 放行（仅检查 session，不查数据库）
    if session.get('user_id'):
        return None

    # 如果允许游客模式 — 放行
    if auth.ALLOW_GUEST:
        return None

    # 记录用户原始请求路径，登录后跳回
    if path != '/' and not path.startswith('/api/'):
        session['next_url'] = path
    else:
        session['next_url'] = '/'

    # API 请求返回 401，页面请求 302 跳转
    if path.startswith('/api/'):
        return jsonify({'error': '请先登录', 'need_login': True}), 401
    return redirect(url_for('login_page'))


# ==================== 定期清理 ====================
_last_cleanup_time = 0
_CLEANUP_INTERVAL = 3600  # 1小时清理一次

def _maybe_cleanup():
    """定期清理过期的后台任务和上传会话"""
    global _last_cleanup_time
    now = time.time()
    if now - _last_cleanup_time < _CLEANUP_INTERVAL:
        return
    _last_cleanup_time = now
    try:
        db.cleanup_old_tasks(max_age_hours=6)
        logger.info("清理过期任务完成")
    except Exception as e:
        logger.error(f"清理过期任务失败: {e}")


# ==================== 全局错误处理 ====================
@app.errorhandler(413)
def request_entity_too_large(error):
    """文件超过 MAX_CONTENT_LENGTH 时返回 JSON 而非默认 HTML 页面"""
    return jsonify({'error': f'文件过大，最大支持 {app.config["MAX_CONTENT_LENGTH"] // 1024 // 1024}MB'}), 413

@app.errorhandler(500)
def internal_server_error(error):
    """服务器内部错误返回 JSON 而非默认 HTML 页面"""
    logger.error(f"500错误: {traceback.format_exc()}")
    return jsonify({'error': '服务器内部错误，请稍后重试'}), 500

@app.errorhandler(Exception)
def handle_exception(error):
    """捕获所有未处理异常，返回 JSON（不泄露内部错误详情）"""
    logger.error(f"未处理异常: {traceback.format_exc()}")
    # 安全：生产环境不返回详细错误信息，仅记录日志
    return jsonify({'error': '服务器内部错误，请稍后重试'}), 500


@app.after_request
def add_cache_headers(response):
    """为静态资源添加缓存头，减少重复下载，并添加安全响应头"""
    path = request.path
    # JS/CSS 文件：带版本号查询参数时缓存1小时，否则不缓存
    if path.startswith('/static/') and (path.endswith('.js') or path.endswith('.css')):
        response.headers['Cache-Control'] = 'public, max-age=3600, must-revalidate'
    # 其他静态文件缓存1天
    elif path.startswith('/static/') or path.startswith('/assets/'):
        response.headers['Cache-Control'] = 'public, max-age=3600'
    # API 响应不缓存
    elif path.startswith('/api/'):
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    # HTML页面 — 确保ETag与压缩正确配合
    elif response.headers.get('ETag'):
        response.headers['Vary'] = 'Accept-Encoding'

    # 安全响应头（适用于所有响应）
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    # HSTS — 仅在 HTTPS 环境下生效
    if request.is_secure:
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'

    return response


# Register Blueprints
try:
    from bp_ai import register as register_ai_bp
    register_ai_bp(app)
except ImportError as e:
    logger.warning(f"bp_ai Blueprint 加载失败: {e}")

try:
    from bp_user import register as register_user_bp
    register_user_bp(app)
except ImportError as e:
    logger.warning(f"bp_user Blueprint 加载失败: {e}")
from date_utils import normalize_date

class ExcelReader:
    """统一的 Excel 读取器，支持 .xls、.xlsx、.csv 格式，以及 HTML 格式的 Excel 文件"""
    
    def __init__(self, file_path):
        self.file_path = file_path
        self.ext = os.path.splitext(file_path)[1].lower()
        self._wb = None
        self._is_xls = self.ext == '.xls'
        self._is_csv = self.ext == '.csv'
        self._is_html = False
        self._read_only = False
        self._csv_data = None  # CSV 数据缓存
        
    def open(self):
        # CSV 格式：用 pandas 读取（自动检测编码）
        if self._is_csv:
            import pandas as pd
            import csv as csv_module
            # 自动检测分隔符
            try:
                with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    sample = f.read(8192)
                try:
                    dialect = csv_module.Sniffer().sniff(sample, delimiters=',;\t|')
                    delimiter = dialect.delimiter
                except csv_module.Error:
                    delimiter = ','
            except Exception:
                delimiter = ','

            # 尝试多种编码 + 自动检测的分隔符
            for encoding in ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'latin1']:
                try:
                    df = pd.read_csv(
                        self.file_path,
                        encoding=encoding,
                        dtype=str,
                        keep_default_na=False,
                        sep=delimiter,
                        engine='python',
                        on_bad_lines='skip'
                    )
                    if len(df.columns) > 1 or (len(df.columns) == 1 and delimiter != ','):
                        self._csv_data = df
                        return self
                    # 如果只有1列且分隔符是逗号，可能是分隔符不对，尝试其他分隔符
                    for alt_delim in [';', '\t', '|']:
                        try:
                            df2 = pd.read_csv(
                                self.file_path,
                                encoding=encoding,
                                dtype=str,
                                keep_default_na=False,
                                sep=alt_delim,
                                engine='python',
                                on_bad_lines='skip'
                            )
                            if len(df2.columns) > 1:
                                self._csv_data = df2
                                return self
                        except Exception:
                            continue
                except Exception:
                    continue

            # 全部失败，用 csv 模块兜底
            try:
                for encoding in ['utf-8-sig', 'utf-8', 'gbk', 'latin1']:
                    try:
                        rows = []
                        with open(self.file_path, 'r', encoding=encoding, errors='replace', newline='') as f:
                            reader = csv_module.reader(f)
                            for row in reader:
                                rows.append(row)
                        if rows and len(rows) > 0:
                            import pandas as pd
                            df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(columns=rows[0])
                            self._csv_data = df
                            return self
                    except Exception:
                        continue
            except Exception:
                pass

            raise ValueError('无法解析CSV文件，请检查文件格式和编码')

        # 检测是否是 HTML 格式的假 Excel 文件
        if self._is_xls:
            with open(self.file_path, 'rb') as f:
                header = f.read(100)
            if b'<html' in header.lower() or b'<!doctype' in header.lower():
                self._is_html = True
        
        if self._is_html:
            return self  # HTML 格式不需要打开
        
        if self._is_xls:
            import xlrd
            try:
                self._wb = xlrd.open_workbook(self.file_path)
            except Exception as e:
                raise ValueError(f'无法读取 .xls 文件: {str(e)}。如果文件是从网页下载的，请转换为 .xlsx 格式后再上传。')
        else:
            from openpyxl import load_workbook
            # 先尝试 read_only=True（内存效率高），如果读到的行太少则回退到 read_only=False
            self._wb = load_workbook(self.file_path, data_only=True, read_only=True)
            self._read_only = True
        return self
    
    def close(self):
        if self._is_csv:
            self._csv_data = None
            return
        if self._wb and not self._is_xls and not self._is_html:
            self._wb.close()
    
    @property
    def sheetnames(self):
        if self._is_csv:
            return ['Sheet1']
        if self._is_html:
            return ['Sheet1']  # HTML 格式默认返回一个 sheet
        if self._is_xls:
            return self._wb.sheet_names()
        return self._wb.sheetnames
    
    def get_sheet_data(self, sheet_name):
        """获取指定 sheet 的所有行数据，返回 list of lists"""
        if self._is_csv:
            if self._csv_data is None:
                self.open()
            df = self._csv_data
            try:
                # 表头 + 数据行（使用 values 提高性能）
                headers = [str(h).strip() if h is not None else '' for h in df.columns.tolist()]
                rows = [headers]
                # 使用 itertuples 比 iterrows 更快
                for row in df.itertuples(index=False, name=None):
                    rows.append([str(v).strip() if v is not None else '' for v in row])
                return rows
            except Exception as e:
                logger.warning(f"CSV get_sheet_data 失败，尝试备选方式: {e}")
                # 备选：直接转 list
                headers = [str(h) for h in df.columns.tolist()]
                rows = [headers] + df.values.tolist()
                return [[str(v).strip() if v is not None else '' for v in row] for row in rows]

        if self._is_html:
            return self._parse_html_excel()
        
        if self._is_xls:
            sheet = self._wb.sheet_by_name(sheet_name)
            rows = []
            for row_idx in range(sheet.nrows):
                row = [str(sheet.cell_value(row_idx, col_idx)).strip() if sheet.cell_value(row_idx, col_idx) != '' else '' 
                       for col_idx in range(sheet.ncols)]
                rows.append(row)
            return rows
        else:
            ws = self._wb[sheet_name]
            rows = [[str(cell).strip() if cell is not None else '' for cell in row] for row in ws.iter_rows(values_only=True)]
            
            # read_only=True 模式下某些 Excel 文件只能读到1-2行（openpyxl已知bug）
            # 检测到行数异常少时，回退到 read_only=False 重新读取
            if self._read_only and len(rows) <= 2:
                try:
                    self._wb.close()
                except Exception:
                    pass
                from openpyxl import load_workbook
                self._wb = load_workbook(self.file_path, data_only=True, read_only=False)
                self._read_only = False
                ws = self._wb[sheet_name]
                rows = [[str(cell).strip() if cell is not None else '' for cell in row] for row in ws.iter_rows(values_only=True)]
            
            return rows
    
    def get_headers_only(self, sheet_name=None):
        """轻量级方法：只读取表头行，不加载全部数据（避免 OOM）"""
        if self._is_html:
            return self._parse_html_headers_only()
        
        # 非 HTML 格式：打开后只读第一行
        if not self._wb:
            self.open()
        rows = self.get_sheet_data(sheet_name or (self.sheetnames[0] if self.sheetnames else None))
        return rows[0] if rows else []
    
    def _parse_html_headers_only(self):
        """只解析 HTML 表头，不加载整个文件到内存（流式读取前 N KB）"""
        try:
            # 只读取文件前 500KB 来提取表头和第一行数据
            # 增大到 500KB 以防 thead 超过 200KB（多语言/长列名场景）
            READ_SIZE = 500 * 1024
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                chunk = f.read(READ_SIZE)
            
            # 用正则提取 <th> 标签内容作为表头
            th_pattern = re.compile(r'<th[^>]*>(.*?)</th>', re.IGNORECASE | re.DOTALL)
            th_matches = th_pattern.findall(chunk)
            
            # 清理 HTML 标签和实体
            def clean_html_text(text):
                # 移除嵌套标签
                text = re.sub(r'<[^>]+>', '', text)
                # HTML 实体
                text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
                text = text.replace('&quot;', '"').replace('&#39;', "'").replace('&nbsp;', ' ')
                # 处理数字 HTML 实体
                text = re.sub(r'&#(\d+);', lambda m: chr(int(m.group(1))) if int(m.group(1)) < 65536 else '', text)
                text = re.sub(r'&#x([0-9a-fA-F]+);', lambda m: chr(int(m.group(1), 16)) if int(m.group(1), 16) < 65536 else '', text)
                return text.strip()
            
            all_headers = [clean_html_text(m) for m in th_matches if clean_html_text(m)]
            
            if not all_headers:
                # 尝试从第一个 <tr> 中的 <td> 提取
                tr_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.IGNORECASE | re.DOTALL)
                tr_matches = tr_pattern.findall(chunk)
                if tr_matches:
                    td_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.IGNORECASE | re.DOTALL)
                    td_matches = td_pattern.findall(tr_matches[0])
                    all_headers = [clean_html_text(m) for m in td_matches if clean_html_text(m)]
            
            if not all_headers:
                all_headers = ['Project', 'Key', 'Summary', 'Issue Type', 'Status',
                              'Priority', 'Resolution', 'Assignee', 'Reporter', 'Creator',
                              'Created', 'Last Viewed', 'Updated', 'Resolved', 'Affects Version/s']
            
            # 提取第一行数据（thead 之后的第一个 tr 中的 td）
            first_data_row = []
            # 找到 </thead> 后的内容
            thead_end = chunk.lower().find('</thead>')
            if thead_end >= 0:
                after_thead = chunk[thead_end:]
                # 找第一个 <tr>...</tr>
                tr_match = re.search(r'<tr[^>]*>(.*?)</tr>', after_thead, re.IGNORECASE | re.DOTALL)
                if tr_match:
                    td_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.IGNORECASE | re.DOTALL)
                    td_matches = td_pattern.findall(tr_match.group(1))
                    first_data_row = [clean_html_text(m) for m in td_matches]
            
            # 估算数据行数：统计文件中 <tr 标签出现次数（不加载整个文件）
            data_row_count = 0
            try:
                with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    # 分块统计 <tr 出现次数
                    for piece in iter(lambda: f.read(1024 * 1024), ''):
                        data_row_count += piece.count('<tr')
            except Exception:
                pass
            # 减去表头行
            data_row_count = max(0, data_row_count - 1)
            
            del chunk
            gc.collect()
            
            logger.info(f"HTML表头流式解析完成: {len(all_headers)}列, 约{data_row_count}行")
            
            # 返回表头 + 第一行数据 + 行数估算
            return all_headers, first_data_row, data_row_count
                
        except Exception as e:
            raise ValueError(f'解析 HTML 格式 Excel 文件失败: {str(e)}')
    
    def _parse_html_excel(self):
        """使用 lxml 高性能解析 HTML 格式 Excel 文件（替代逐行正则，速度提升 10-50 倍）"""
        try:
            from lxml import html as lxml_html
            import html as html_module

            # === 第一步：用 lxml 解析整个 HTML（C 级解析器，极快） ===
            with open(self.file_path, 'rb') as f:
                tree = lxml_html.fromstring(f.read())

            # 找到主 table（取第一个含 thead 或行数最多的 table）
            # 直接使用 xpath（cssselect 可能因缺少依赖而抛异常）
            tables = tree.xpath('//table')
            if not tables:
                # 没有 table，回退到正则方式
                return self._parse_html_excel_regex()

            table = tables[0]
            # 如果有多个 table，选行数最多的
            if len(tables) > 1:
                max_rows = 0
                for t in tables:
                    cnt = len(t.xpath('.//tr'))
                    if cnt > max_rows:
                        max_rows = cnt
                        table = t

            # === 第二步：提取表头 ===
            all_headers = []
            # 优先从 thead > tr > th 提取
            thead = table.find('thead')
            if thead is not None:
                # 处理多行表头：取最后一行（通常是实际列头），并展开 colspan
                thead_rows = thead.findall('tr')
                if thead_rows:
                    # 选择 th 最多的一行作为表头行（通常是最后一行）
                    best_row = max(thead_rows, key=lambda r: len(r.findall('th')))
                    for th in best_row.findall('th'):
                        text = html_module.unescape(lxml_html.tostring(th, encoding='unicode', method='text').strip())
                        # 处理 colspan：重复表头文本填满列数
                        colspan = th.get('colspan', '1')
                        try:
                            span = int(colspan)
                        except (ValueError, TypeError):
                            span = 1
                        for _ in range(span):
                            all_headers.append(text if text else '')
                else:
                    for th in thead.iter('th'):
                        text = html_module.unescape(lxml_html.tostring(th, encoding='unicode', method='text').strip())
                        if text:
                            all_headers.append(text)
            # 回退：从第一行 td 提取
            if not all_headers:
                first_tr = table.find('.//tr')
                if first_tr is not None:
                    for td in first_tr.iter('td'):
                        text = html_module.unescape(lxml_html.tostring(td, encoding='unicode', method='text').strip())
                        if text:
                            all_headers.append(text)

            if not all_headers:
                all_headers = ['Project', 'Key', 'Summary', 'Issue Type', 'Status',
                              'Priority', 'Resolution', 'Assignee', 'Reporter', 'Creator',
                              'Created', 'Last Viewed', 'Updated', 'Resolved', 'Affects Version/s']

            # 扫描关键特殊列：Severity / Component / Fix Version
            thead_col_map = {}
            for idx, h in enumerate(all_headers):
                h_lower = h.lower()
                if 'severity' in h_lower and 'severity_col' not in thead_col_map:
                    thead_col_map['severity_col'] = idx
                if 'component' in h_lower and 'component_col' not in thead_col_map:
                    thead_col_map['component_col'] = idx
                if ('fix version' in h_lower or 'fixversion' in h_lower) and 'fix_version_col' not in thead_col_map:
                    thead_col_map['fix_version_col'] = idx

            # 保留所有列（之前只保留前15列+3特殊列，导致列错位时severity数据丢失）
            full_headers = list(all_headers)
            keep_cols_sorted = list(range(len(all_headers)))
            _log_mem(f"HTML lxml解析：保留全部{len(full_headers)}列")

            # === 第三步：提取数据行（tbody 中的 tr > td） ===
            result_rows = [full_headers]
            row_count = 0

            # 获取 tbody，如果没有就用 table 本身
            tbody = table.find('tbody')
            tr_source = tbody if tbody is not None else table

            for tr in tr_source.iter('tr'):
                # 跳过表头行
                if tr.find('th') is not None:
                    continue

                tds = tr.findall('td')
                if not tds:
                    continue

                # 提取每个 td 的文本（text_content 是 C 级方法，极快）
                td_texts = [html_module.unescape(td.text_content().strip()) for td in tds]

                row = []
                for col_idx in keep_cols_sorted:
                    if col_idx < len(td_texts):
                        row.append(td_texts[col_idx])
                    else:
                        row.append('')

                if any(c.strip() for c in row):
                    result_rows.append(row)
                    row_count += 1

            del tree, table
            gc.collect()

            logger.info(f"HTML lxml解析完成：{row_count}行 x {len(full_headers)}列")
            _log_mem(f"HTML lxml解析完成：{row_count}行 x {len(full_headers)}列")
            return result_rows

        except Exception as e:
            error_msg = f'解析 HTML 格式 Excel 文件失败: {type(e).__name__}: {str(e)}'
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            # 回退到正则方式
            logger.warning("lxml 解析失败，回退到正则解析")
            return self._parse_html_excel_regex()

    def _parse_html_excel_regex(self):
        """正则方式解析 HTML 格式（回退方案）"""
        try:
            # HTML 实体和标签清理
            def clean_html_text(text):
                text = re.sub(r'<[^>]+>', '', text)
                text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
                text = text.replace('&quot;', '"').replace('&#39;', "'").replace('&nbsp;', ' ')
                text = re.sub(r'&#(\d+);', lambda m: chr(int(m.group(1))) if int(m.group(1)) < 65536 else '', text)
                text = re.sub(r'&#x([0-9a-fA-F]+);', lambda m: chr(int(m.group(1), 16)) if int(m.group(1), 16) < 65536 else '', text)
                return text.strip()

            READ_SIZE = 1024 * 1024  # 1MB — 确保读取所有 <th> 标签（Jira导出156列约需500KB+）
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                head_chunk = f.read(READ_SIZE)

            th_pattern = re.compile(r'<th[^>]*>(.*?)</th>', re.IGNORECASE | re.DOTALL)
            th_matches = th_pattern.findall(head_chunk)
            # 处理 colspan 属性
            th_full_pattern = re.compile(r'<th([^>]*)>(.*?)</th>', re.IGNORECASE | re.DOTALL)
            th_full_matches = th_full_pattern.findall(head_chunk)
            all_headers = []
            for attrs_str, content in th_full_matches:
                text = clean_html_text(content)
                # 解析 colspan
                colspan_match = re.search(r'colspan\s*=\s*["\']?(\d+)', attrs_str, re.IGNORECASE)
                span = int(colspan_match.group(1)) if colspan_match else 1
                for _ in range(span):
                    all_headers.append(text)

            if not all_headers:
                tr_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.IGNORECASE | re.DOTALL)
                tr_match = tr_pattern.search(head_chunk)
                if tr_match:
                    td_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.IGNORECASE | re.DOTALL)
                    td_matches = td_pattern.findall(tr_match.group(1))
                    all_headers = [clean_html_text(m) for m in td_matches if clean_html_text(m)]

            if not all_headers:
                all_headers = ['Project', 'Key', 'Summary', 'Issue Type', 'Status',
                              'Priority', 'Resolution', 'Assignee', 'Reporter', 'Creator',
                              'Created', 'Last Viewed', 'Updated', 'Resolved', 'Affects Version/s']

            thead_col_map = {}
            for idx, h in enumerate(all_headers):
                h_lower = h.lower()
                if 'severity' in h_lower and 'severity_col' not in thead_col_map:
                    thead_col_map['severity_col'] = idx
                if 'component' in h_lower and 'component_col' not in thead_col_map:
                    thead_col_map['component_col'] = idx
                if ('fix version' in h_lower or 'fixversion' in h_lower) and 'fix_version_col' not in thead_col_map:
                    thead_col_map['fix_version_col'] = idx

            # 保留所有列（与 lxml 路径保持一致）
            full_headers = list(all_headers)
            keep_cols_sorted = list(range(len(all_headers)))
            _log_mem(f"HTML正则解析：保留全部{len(full_headers)}列")

            result_rows = [full_headers]
            row_count = 0
            td_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.IGNORECASE | re.DOTALL)

            thead_end_pos = head_chunk.lower().find('</thead>')
            skip_header = thead_end_pos >= 0
            table_start_pos = head_chunk.lower().find('<table')
            if not skip_header and table_start_pos >= 0:
                start_pos = table_start_pos
            elif skip_header:
                start_pos = thead_end_pos + 8
            else:
                start_pos = 0

            table_ended = False
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                f.seek(start_pos)
                buffer = ''
                CHUNK_SIZE = 512 * 1024

                while True and not table_ended:
                    piece = f.read(CHUNK_SIZE)
                    if not piece:
                        break
                    buffer += piece

                    while True:
                        tr_start = buffer.find('<tr')
                        if tr_start == -1:
                            buffer = ''
                            break
                        tr_end = buffer.find('</tr>', tr_start)
                        if tr_end == -1:
                            buffer = buffer[tr_start:]
                            if len(buffer) > 1024 * 1024:
                                buffer = ''
                            break

                        tr_content = buffer[tr_start:tr_end + 5]
                        buffer = buffer[tr_end + 5:]

                        if '<th' in tr_content.lower():
                            next_tr = buffer.find('<tr')
                            between = buffer[:next_tr] if next_tr >= 0 else buffer
                            if '</table>' in between.lower():
                                table_ended = True
                                break
                            continue

                        td_matches = td_pattern.findall(tr_content)
                        if not td_matches:
                            continue

                        row = []
                        for col_idx in keep_cols_sorted:
                            if col_idx < len(td_matches):
                                row.append(clean_html_text(td_matches[col_idx]))
                            else:
                                row.append('')

                        if any(c.strip() for c in row):
                            result_rows.append(row)
                            row_count += 1

                        next_tr = buffer.find('<tr')
                        between = buffer[:next_tr] if next_tr >= 0 else buffer
                        if '</table>' in between.lower():
                            table_ended = True
                            break

                    del piece

            del buffer, head_chunk
            gc.collect()

            logger.info(f"HTML正则解析完成：{row_count}行 x {len(full_headers)}列, table_ended={table_ended}")
            _log_mem(f"HTML正则解析完成：{row_count}行 x {len(full_headers)}列")
            return result_rows

        except Exception as e:
            error_msg = f'解析 HTML 格式 Excel 文件失败: {type(e).__name__}: {str(e)}'
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            raise ValueError(error_msg)
    
    def get_sheet_names(self):
        """获取所有 sheet 名称"""
        return self.sheetnames


def read_excel_file(file_path, sheet_name=None):
    """读取 Excel 文件，返回 (sheet_names, sheet_data) 或 (sheet_names, None)"""
    reader = ExcelReader(file_path)
    reader.open()
    
    sheet_names = reader.get_sheet_names()
    sheet_data = None
    
    if sheet_name:
        sheet_data = reader.get_sheet_data(sheet_name)
    
    reader.close()
    return sheet_names, sheet_data

# === 安全工具函数 ===
import re as _re
def _validate_file_id(file_id):
    """验证 file_id 格式，防止路径遍历攻击
    合法格式：16位十六进制字符（md5 hexdigest[:16]）
    """
    if not file_id or not isinstance(file_id, str):
        return False
    return bool(_re.match(r'^[a-f0-9]{16}$', file_id))

# === AI 配置管理 — 委托给 ai_utils 共享模块 ===
# 保留向后兼容别名，供 app.py 内其他路由调用
get_ai_config = ai_utils.get_ai_config
_call_ai = ai_utils.call_ai
_call_ai_stream = ai_utils.call_ai_stream


# === AI 路由已迁移到 bp_ai.py Blueprint ===
# 以下路由由 bp_ai Blueprint 提供：/api/ai-chat, /api/ai-models,
# /api/ai-config (GET/POST), /api/ai-test
# 保留 get_ai_config/_call_ai/_call_ai_stream 别名供 app.py 其他路由调用

# ==================== 认证装饰器 ====================
def login_required(f):
    """严格登录装饰器：未登录返回401，登录用户存入 g.user"""
    @wraps(f)
    def decorated(*args, **kwargs):
        user = auth.get_current_user()
        if not user:
            return jsonify({'error': '请先登录'}), 401
        g.user = user
        return f(*args, **kwargs)
    return decorated


def login_required_or_guest(f):
    """登录或访客装饰器：ALLOW_GUEST=true 时允许访客访问，否则要求登录"""
    @wraps(f)
    def decorated(*args, **kwargs):
        user = auth.get_current_user()
        if not user and not auth.ALLOW_GUEST:
            return jsonify({'error': '请先登录'}), 401
        g.user = user
        return f(*args, **kwargs)
    return decorated


# === 页面路由 ===
# ==================== 认证路由 ====================

@app.route('/login')
def login_page():
    """登录页面"""
    error = request.args.get('error', '')
    return render_template('login.html',
                           error=error,
                           feishu_configured=auth.is_configured('feishu'),
                           google_configured=auth.is_configured('google'),
                           allow_guest=auth.ALLOW_GUEST)


@app.route('/auth/feishu')
def auth_feishu():
    """发起飞书OAuth登录"""
    return auth.feishu_login()


@app.route('/auth/feishu/callback')
def auth_feishu_callback():
    """飞书OAuth回调"""
    return auth.feishu_callback()


@app.route('/auth/google')
def auth_google():
    """发起Google OAuth登录"""
    return auth.google_login()


@app.route('/auth/google/callback')
def auth_google_callback():
    """Google OAuth回调"""
    return auth.google_callback()


@app.route('/auth/logout')
def auth_logout():
    """退出登录"""
    return auth.logout()


# ==================== 模板渲染缓存 + ETag ====================
# 内存缓存已渲染的模板，配合ETag实现304 Not Modified
# 静态模板（不含current_user）全量缓存；含current_user的按用户缓存
_template_cache = {}

# 不含动态用户信息的模板 — 可全局缓存
_STATIC_TEMPLATES = frozenset({
    'excel_analysis.html', 'md2pdf.html',
    'plan_generator.html',
})

def _get_template_mtime(template_name):
    """获取模板文件的修改时间，用于缓存失效检测"""
    try:
        tpl_path = os.path.join(app.template_folder, template_name)
        if os.path.isfile(tpl_path):
            return int(os.path.getmtime(tpl_path))
    except Exception:
        pass
    return 0

def cached_render(template_name, **context):
    """渲染模板并缓存结果，支持ETag/304。
    
    - 静态模板：全局缓存，首次渲染后后续请求直接返回304或缓存内容
    - 动态模板：按用户缓存，同一用户重复访问直接返回304
    - 模板文件修改后自动失效缓存
    """
    if template_name in _STATIC_TEMPLATES:
        cache_key = template_name
    else:
        uid = session.get('user_id', 0)
        cache_key = f'{template_name}:{uid}'

    mtime = _get_template_mtime(template_name)
    cached = _template_cache.get(cache_key)
    if cached is not None:
        cached_mtime, etag, html = cached
        # 模板文件未修改且缓存存在 — 使用缓存
        if cached_mtime == mtime:
            # 浏览器发送 If-None-Match — 内容未变，返回304（无body，瞬时响应）
            if request.headers.get('If-None-Match') == etag:
                resp = make_response('', 304)
                resp.headers['ETag'] = etag
                resp.headers['Cache-Control'] = 'no-cache'  # 必须验证，但304省带宽
                return resp
            resp = make_response(html)
            resp.headers['ETag'] = etag
            resp.headers['Cache-Control'] = 'no-cache'
            return resp
        # 模板文件已修改 — 清除旧缓存，重新渲染

    # 首次渲染或缓存失效后重新渲染
    html = render_template(template_name, **context)
    etag = hashlib.md5(html.encode('utf-8')).hexdigest()[:16]
    _template_cache[cache_key] = (mtime, etag, html)

    resp = make_response(html)
    resp.headers['ETag'] = etag
    resp.headers['Cache-Control'] = 'no-cache'
    return resp


# 注册页面路由 Blueprint（15个简单模板渲染路由已移至 routes/pages.py）
app.register_blueprint(create_pages_blueprint(cached_render))
@app.route('/api/jira-search', methods=['POST'])
def api_jira_search():
    """Jira 搜索代理 — 通过 JQL 或 Jira 链接获取 CR 数据"""
    data = request.get_json(force=True)
    domain = data.get('domain', '').strip().rstrip('/')
    email = data.get('email', '').strip()
    api_token = data.get('api_token', '').strip()
    jql = data.get('jql', '').strip()
    jira_url = data.get('jira_url', '').strip()
    max_results = min(int(data.get('max_results', 500)), 1000)

    if not domain:
        return jsonify({'error': '请填写 Jira 域名'}), 400

    # 域名归一化：缺少 scheme 时自动补 https://
    if domain and not domain.startswith(('http://', 'https://')):
        domain = 'https://' + domain

    # 如果传入 Jira 链接，尝试从中提取 JQL
    if jira_url and not jql:
        try:
            from urllib.parse import urlparse, parse_qs, unquote
            parsed = urlparse(jira_url)
            params = parse_qs(parsed.query)
            if 'jql' in params:
                jql = unquote(params['jql'][0])
            elif 'filter' in params:
                # JQL 原生支持 filter=ID 语法，直接使用，无需额外请求 filter 接口
                filter_id = params['filter'][0]
                jql = f"filter={filter_id}"
            else:
                # 尝试从路径中提取 filter ID，如 /issues/?filter=86482 或 /browse/PROJ-123
                import re
                m = re.search(r'filter[=/](\d+)', jira_url)
                if m:
                    jql = f"filter={m.group(1)}"
        except Exception:
            pass

    if not jql:
        return jsonify({'error': '请提供 JQL 或 Jira 搜索链接'}), 400

    headers = {'Accept': 'application/json', 'Content-Type': 'application/json'}
    payload = {
        'jql': jql,
        'maxResults': max_results,
        'fields': ['summary', 'status', 'created', 'resolutiondate',
                   'fixVersions', 'issuetype', 'priority', 'assignee', 'reporter']
    }

    # 判断 Jira Cloud 还是 Server：Cloud 域名含 .atlassian.net
    is_cloud = '.atlassian.net' in domain
    # API 版本：Cloud 先试 api/3 再 api/2；Server 直接用 api/2（减少请求数，避免限流）
    api_versions = ['/rest/api/3/search', '/rest/api/2/search'] if is_cloud else ['/rest/api/2/search']

    # 构建鉴权尝试方案（先邮箱，再短用户名，最后 Bearer Token）
    auth_attempts = []
    if email and api_token:
        auth_attempts.append(('basic', (email, api_token)))
        if '@' in email:
            auth_attempts.append(('basic', (email.split('@')[0], api_token)))
        auth_attempts.append(('bearer', api_token))
    else:
        auth_attempts.append(('none', None))

    resp = None
    last_status = None
    try:
        for auth_type, auth_val in auth_attempts:
            req_headers = dict(headers)
            if auth_type == 'bearer':
                req_headers['Authorization'] = f'Bearer {auth_val}'
                auth_val = None
            for api_path in api_versions:
                search_url = f"{domain}{api_path}"
                resp = requests.post(search_url, json=payload, auth=auth_val,
                                     headers=req_headers, timeout=30)
                last_status = resp.status_code
                if resp.status_code == 200:
                    break
                if resp.status_code == 429:
                    # 限流：立即停止所有尝试，避免加剧限流
                    retry_after = resp.headers.get('Retry-After', '30')
                    return jsonify({'error': f'Jira 请求过于频繁(429)：服务器要求等待 {retry_after} 秒后重试。请稍后再试，不要连续点击导入'}), 429
                if resp.status_code == 404:
                    continue  # api/3 不存在，试 api/2
                break  # 401/403/其他：换下一种鉴权
            if resp.status_code == 200:
                break
            if resp.status_code == 404:
                continue
            break  # 非 404：换下一种鉴权

        if not resp or resp.status_code != 200:
            if last_status == 401:
                return jsonify({'error': 'Jira 鉴权失败(401)：已尝试邮箱/短用户名/Bearer三种方式。Jira Server 请确认用户名(短账号)和密码；或生成 Personal Access Token 填入 Token 栏'}), 401
            if last_status == 403:
                return jsonify({'error': 'Jira 无权限(403)：当前账号无权访问该 filter 或项目'}), 403
            if last_status == 404:
                return jsonify({'error': f'Jira API 未找到(404)：请确认域名 {domain} 是否正确'}), 404
            return jsonify({'error': f'Jira API 返回 {last_status}: {resp.text[:200] if resp else "无响应"}'}), last_status or 502

        result = resp.json()
        issues = result.get('issues', [])
        total = result.get('total', len(issues))

        # 转换为统一格式
        cr_data = []
        for issue in issues:
            fields = issue.get('fields', {})
            status = fields.get('status', {}).get('name', 'Unknown')
            created = fields.get('created', '')[:10] if fields.get('created') else ''
            resolved = fields.get('resolutiondate', '')[:10] if fields.get('resolutiondate') else ''
            fix_versions = [v.get('name', '') for v in fields.get('fixVersions', [])]
            version = fix_versions[0] if fix_versions else '未指定'
            issuetype = fields.get('issuetype', {}).get('name', '')
            priority = fields.get('priority', {}).get('name', '')
            assignee = fields.get('assignee', {}).get('displayName', '未分配') if fields.get('assignee') else '未分配'

            cr_data.append({
                'key': issue.get('key', ''),
                'summary': fields.get('summary', ''),
                'status': status,
                'created': created,
                'resolved': resolved,
                'version': version,
                'issuetype': issuetype,
                'priority': priority,
                'assignee': assignee,
                'is_resolved': 1 if resolved else 0,
            })

        return jsonify({
            'total': total,
            'returned': len(cr_data),
            'issues': cr_data,
            'jql': jql,
        })

    except requests.exceptions.Timeout:
        return jsonify({'error': 'Jira 请求超时，请检查网络或代理设置'}), 504
    except requests.exceptions.ConnectionError:
        return jsonify({'error': '无法连接到 Jira 服务器，请检查域名和网络'}), 502
    except Exception as e:
        return jsonify({'error': f'请求失败: {str(e)}'}), 500
@app.route('/api/generate-minutes', methods=['POST'])
@login_required
def api_generate_minutes():
    """AI生成会议纪要"""

    data = request.get_json(silent=True) or {}
    transcript = data.get('transcript', '').strip()
    title = data.get('title', '未命名会议')
    attendees = data.get('attendees', '未填写')
    date = data.get('date', '')
    model = data.get('model')  # v3.0: 支持前端指定模型

    if not transcript or len(transcript) < 10:
        return jsonify({'error': '转写内容太少，无法生成纪要'}), 400

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请联系管理员设置API Key'}), 503

    try:
        # 构建提示词
        prompt = f"""请根据以下会议语音转写内容，生成一份结构化的会议纪要。

会议信息：
- 主题：{title}
- 日期：{date}
- 参会人员：{attendees}

会议转写内容：
{transcript[:16000]}

请生成会议纪要，使用HTML格式，包含以下部分：
1. <h2>会议概要</h2> - 简要说明会议目的和主要内容（2-3句话）
2. <h2>讨论要点</h2> - 列出讨论的主要议题和关键观点，使用<ul><li>格式
3. <h2>决议事项</h2> - 列出会议达成的决定，使用<ol><li>格式
4. <h2>待办事项</h2> - 列出需要后续跟进的任务，如有责任人请标注，使用<table>格式（列：序号、任务、责任人、截止时间）

要求：
- 直接输出HTML内容，不要包含```html标记
- 语言简洁专业
- 如果转写内容不清晰，合理推断并标注
- 只输出HTML，不要其他文字"""

        minutes_html = _call_ai(
            messages=[
                {'role': 'system', 'content': '你是一个专业的会议纪要撰写助手。请根据会议转写内容生成结构清晰、内容准确的会议纪要。'},
                {'role': 'user', 'content': prompt}
            ],
            model=model,
            max_tokens=3000,
            temperature=0.3,
            timeout=90
        )

        # 清理可能的markdown标记
        if '```html' in minutes_html:
            minutes_html = minutes_html.split('```html')[1].split('```')[0]
        elif '```' in minutes_html:
            minutes_html = minutes_html.split('```')[1].split('```')[0]
        minutes_html = minutes_html.strip()

        return jsonify({
            'status': 'success',
            'minutes': minutes_html
        })

    except requests.exceptions.Timeout:
        return jsonify({'error': 'AI服务响应超时，请稍后重试'}), 504
    except Exception as e:
        logger.error(f"生成会议纪要失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-minutes-stream', methods=['POST'])
@login_required
def api_generate_minutes_stream():
    """SSE 流式版：AI生成会议纪要 — 边生成边输出"""

    data = request.get_json(silent=True) or {}
    transcript = data.get('transcript', '').strip()
    title = data.get('title', '未命名会议')
    attendees = data.get('attendees', '未填写')
    date = data.get('date', '')
    model = data.get('model')

    if not transcript or len(transcript) < 10:
        return jsonify({'error': '转写内容太少，无法生成纪要'}), 400

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请联系管理员设置API Key'}), 503

    prompt = f"""请根据以下会议语音转写内容，生成一份结构化的会议纪要。

会议信息：
- 主题：{title}
- 日期：{date}
- 参会人员：{attendees}

会议转写内容：
{transcript[:16000]}

请生成会议纪要，使用Markdown格式，包含以下部分：

### 📋 会议概要
简要说明会议目的和主要内容（2-3句话）

### 💬 讨论要点
列出讨论的主要议题和关键观点

### ✅ 决议事项
列出会议达成的决定

### 📝 待办事项
列出需要后续跟进的任务，如有责任人请标注

要求：
- 语言简洁专业
- 如果转写内容不清晰，合理推断并标注
- 使用 Markdown 格式输出"""

    messages = [
        {'role': 'system', 'content': '你是一个专业的会议纪要撰写助手。请根据会议转写内容生成结构清晰、内容准确的会议纪要。'},
        {'role': 'user', 'content': prompt}
    ]

    return Response(
        stream_with_context(_call_ai_stream(messages, model=model, max_tokens=3000, temperature=0.3)),
        content_type='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no', 'Connection': 'keep-alive'}
    )


# === v3.0 OCR 图片文字识别 ===
@app.route('/api/ocr', methods=['POST'])
@login_required_or_guest
def api_ocr():
    """OCR 图片文字识别 — 使用 qwen-vl 多模态模型"""

    # 支持 base64 图片或文件上传
    import base64
    import requests as req

    image_data_url = None
    prompt = '请识别图片中的所有文字内容，保持原有格式和排版。如果图片中包含表格，请用 Markdown 表格格式输出。只输出识别到的文字，不要其他说明。'

    if 'file' in request.files:
        file = request.files['file']
        img_bytes = file.read()
        if len(img_bytes) > 10 * 1024 * 1024:
            return jsonify({'error': '图片不能超过 10MB'}), 400
        b64 = base64.b64encode(img_bytes).decode('utf-8')
        mime = file.content_type or 'image/png'
        image_data_url = f'data:{mime};base64,{b64}'
    elif request.is_json:
        data = request.get_json(silent=True) or {}
        image_data_url = data.get('image')
        custom_prompt = data.get('prompt')
        if custom_prompt:
            prompt = custom_prompt
    else:
        return jsonify({'error': '请上传图片或提供 base64 图片数据'}), 400

    if not image_data_url:
        return jsonify({'error': '图片数据为空'}), 400

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置'}), 503

    try:
        is_openai = 'dashscope' not in ai_config.get('base_url', '')
        if is_openai:
            # OpenAI 兼容格式（火山引擎/豆包视觉模型）
            response = req.post(
                f"{ai_config.get('base_url', '').rstrip('/')}/chat/completions",
                headers={
                    'Authorization': f'Bearer {ai_config.get("api_key", "")}',
                    'Content-Type': 'application/json'
                },
                json={
                    'model': ai_config.get('vision_model', 'doubao-seed-1-6-250615'),
                    'messages': [
                        {
                            'role': 'user',
                            'content': [
                                {'type': 'image_url', 'image_url': {'url': image_data_url}},
                                {'type': 'text', 'text': prompt}
                            ]
                        }
                    ],
                    'max_tokens': 2000
                },
                timeout=60
            )
            if response.status_code == 200:
                result = response.json()
                text = result.get('choices', [{}])[0].get('message', {}).get('content', '')
                return jsonify({'status': 'success', 'text': text.strip()})
            else:
                logger.error(f"OCR失败: {response.status_code} - {response.text[:200]}")
                return jsonify({'error': f'OCR服务错误({response.status_code}): {response.text[:100]}'}), 502
        else:
            # DashScope 格式（qwen-vl 多模态）
            response = req.post(
                f"{ai_config.get('base_url', 'https://dashscope.aliyuncs.com/api/v1')}/services/aigc/multimodal-generation/generation",
                headers={
                    'Authorization': f'Bearer {ai_config.get("api_key", "")}',
                    'Content-Type': 'application/json'
                },
                json={
                    'model': 'qwen-vl-plus',
                    'input': {
                        'messages': [
                            {
                                'role': 'user',
                                'content': [
                                    {'image': image_data_url},
                                    {'text': prompt}
                                ]
                            }
                        ]
                    }
                },
                timeout=60
            )
            if response.status_code == 200:
                result = response.json()
                text = result.get('output', {}).get('choices', [{}])[0].get('message', {}).get('content', [{}])
                if isinstance(text, list):
                    text = ''.join(item.get('text', '') if isinstance(item, dict) else str(item) for item in text)
                return jsonify({'status': 'success', 'text': text.strip()})
            else:
                logger.error(f"OCR失败: {response.status_code} - {response.text[:200]}")
                return jsonify({'error': f'OCR服务错误({response.status_code})'}), 502

    except req.exceptions.Timeout:
        return jsonify({'error': 'OCR服务响应超时'}), 504
    except Exception as e:
        logger.error(f"OCR异常: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# === v3.0 智能周报生成 ===
@app.route('/api/weekly-report', methods=['POST'])
@login_required_or_guest
def api_weekly_report():
    """AI 智能周报生成"""

    data = request.get_json(silent=True) or {}
    notes = data.get('notes', '').strip()
    meetings = data.get('meetings', '').strip()
    cr_issues = data.get('cr_issues', '').strip()
    extra = data.get('extra', '').strip()
    model = data.get('model')
    name = data.get('name', '')
    week_range = data.get('week_range', '')

    # 至少需要一项内容
    if not any([notes, meetings, cr_issues, extra]):
        return jsonify({'error': '请至少填写一项本周工作内容'}), 400

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置'}), 503

    sections = []
    if name:
        sections.append(f'汇报人：{name}')
    if week_range:
        sections.append(f'汇报周期：{week_range}')
    if notes:
        sections.append(f'【工作笔记/记录】\n{notes[:3000]}')
    if meetings:
        sections.append(f'【会议内容摘要】\n{meetings[:3000]}')
    if cr_issues:
        sections.append(f'【问题/CR记录】\n{cr_issues[:3000]}')
    if extra:
        sections.append(f'【其他补充】\n{extra[:2000]}')

    content_block = '\n\n'.join(sections)

    prompt = f"""请根据以下本周工作素材，生成一份结构化的周报。

素材内容：
{content_block}

请生成周报，使用 HTML 格式，包含以下部分：
1. <h2>本周工作总结</h2> — 概括本周主要工作内容（3-5 条要点）
2. <h2>关键成果</h2> — 本周取得的关键成果或进展
3. <h2>问题与风险</h2> — 遇到的问题、风险及应对措施
4. <h2>下周计划</h2> — 下周工作重点和计划安排

要求：
- 直接输出 HTML，不要 ```html 标记
- 语言简洁专业
- 合理归纳整理，不要简单罗列原文
- 只输出 HTML 内容"""

    try:
        result = _call_ai(
            messages=[
                {'role': 'system', 'content': '你是一个专业的项目汇报助手。请根据用户提供的素材，生成结构清晰、内容准确的周报。'},
                {'role': 'user', 'content': prompt}
            ],
            model=model,
            max_tokens=3000,
            temperature=0.3,
            timeout=90
        )

        # 清理 markdown 标记
        if '```html' in result:
            result = result.split('```html')[1].split('```')[0]
        elif '```' in result:
            result = result.split('```')[1].split('```')[0]
        result = result.strip()

        return jsonify({'status': 'success', 'report': result})

    except Exception as e:
        logger.error(f"周报生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/weekly-report-stream', methods=['POST'])
@login_required
def api_weekly_report_stream():
    """SSE 流式版：AI 智能周报生成 — 边生成边输出"""

    data = request.get_json(silent=True) or {}
    notes = data.get('notes', '').strip()
    meetings = data.get('meetings', '').strip()
    cr_issues = data.get('cr_issues', '').strip()
    extra = data.get('extra', '').strip()
    model = data.get('model')
    name = data.get('name', '')
    week_range = data.get('week_range', '')

    if not any([notes, meetings, cr_issues, extra]):
        return jsonify({'error': '请至少填写一项本周工作内容'}), 400

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置'}), 503

    sections = []
    if name:
        sections.append(f'汇报人：{name}')
    if week_range:
        sections.append(f'汇报周期：{week_range}')
    if notes:
        sections.append(f'【工作笔记/记录】\n{notes[:3000]}')
    if meetings:
        sections.append(f'【会议内容摘要】\n{meetings[:3000]}')
    if cr_issues:
        sections.append(f'【问题/CR记录】\n{cr_issues[:3000]}')
    if extra:
        sections.append(f'【其他补充】\n{extra[:2000]}')

    content_block = '\n\n'.join(sections)

    prompt = f"""请根据以下本周工作素材，生成一份结构化的周报。

素材内容：
{content_block}

请生成周报，使用 Markdown 格式，包含以下部分：

## 本周工作总结
概括本周主要工作内容（3-5 条要点）

## 关键成果
本周取得的关键成果或进展

## 问题与风险
遇到的问题、风险及应对措施

## 下周计划
下周工作重点和计划安排

要求：
- 语言简洁专业
- 合理归纳整理，不要简单罗列原文
- 使用 Markdown 格式输出"""

    messages = [
        {'role': 'system', 'content': '你是一个专业的项目汇报助手。请根据用户提供的素材，生成结构清晰、内容准确的周报。'},
        {'role': 'user', 'content': prompt}
    ]

    return Response(
        stream_with_context(_call_ai_stream(messages, model=model, max_tokens=3000, temperature=0.3)),
        content_type='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no', 'Connection': 'keep-alive'}
    )


# ==================== v3.0 设置面板 ====================
@app.route('/api/system-info', methods=['GET'])
@login_required_or_guest
def api_system_info():
    """系统诊断信息"""

    import psutil
    p = psutil.Process()
    mem_info = p.memory_info()

    # 数据库状态
    db_status = db.check_db()

    # AI 配置状态
    ai_config = get_ai_config()

    # 上传目录
    upload_dir = app.config.get('UPLOAD_FOLDER', '/tmp/toolbox/uploads')
    upload_size = 0
    upload_count = 0
    if os.path.exists(upload_dir):
        for fname in os.listdir(upload_dir):
            fpath = os.path.join(upload_dir, fname)
            if os.path.isfile(fpath):
                upload_count += 1
                upload_size += os.path.getsize(fpath)

    return jsonify({
        'status': 'ok',
        'version': '3.0',
        'python': sys.version.split()[0],
        'memory_mb': round(mem_info.rss / 1024 / 1024, 1),
        'cpu_percent': p.cpu_percent(interval=0.1),
        'threads': p.num_threads(),
        'db': db_status,
        'ai': {
            'enabled': ai_config.get('enabled', False),
            'base_url': ai_config.get('base_url', ''),
            'model': ai_config.get('model', ''),
            'has_key': bool(ai_config.get('api_key', '').strip()),
        },
        'uploads': {
            'count': upload_count,
            'size_mb': round(upload_size / 1024 / 1024, 1),
        },
        'uptime': time.time() - p.create_time(),
    })


@app.route('/api/settings/theme', methods=['POST'])
@login_required_or_guest
def api_save_theme():
    """保存用户主题偏好（主题色、模式）"""
    user = g.user
    data = request.get_json(silent=True) or {}
    theme_mode = data.get('theme_mode', 'auto')  # light / dark / auto
    accent_color = data.get('accent_color', '')   # hex color

    # 保存到用户偏好
    if user:
        prefs = db.get_user_preferences(user['id']) or {}
        db.set_user_preferences(user['id'], theme=theme_mode,
                                accent_color=accent_color if accent_color else None)

    return jsonify({'status': 'success', 'theme_mode': theme_mode, 'accent_color': accent_color})


@app.route('/api/settings/theme', methods=['GET'])
def api_get_theme():
    """获取用户主题偏好"""
    user = auth.get_current_user()
    theme_mode = 'auto'
    accent_color = '#0071e3'

    if user:
        prefs = db.get_user_preferences(user['id']) or {}
        theme_mode = prefs.get('theme', 'auto')
        accent_color = prefs.get('accent_color', '') or '#0071e3'

    return jsonify({
        'status': 'success',
        'theme_mode': theme_mode,
        'accent_color': accent_color
    })


# ==================== v5.0 飞书推送 API ====================

@app.route('/api/settings/feishu', methods=['GET'])
def api_get_feishu():
    """获取用户飞书 Webhook 配置"""
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    webhook = db.get_feishu_webhook(user['id'])
    secret = db.get_feishu_secret(user['id'])
    return jsonify({
        'status': 'success',
        'webhook': webhook,
        'secret': secret,
        'configured': bool(webhook)
    })


@app.route('/api/settings/feishu', methods=['POST'])
def api_set_feishu():
    """保存用户飞书 Webhook 配置"""
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    data = request.get_json(silent=True) or {}
    webhook = (data.get('webhook') or '').strip()
    secret = (data.get('secret') or '').strip()
    if webhook and not webhook.startswith('https://open.feishu.cn/open-apis/bot/v2/hook/'):
        return jsonify({'status': 'error', 'error': 'Webhook URL 格式不正确，应以 https://open.feishu.cn/open-apis/bot/v2/hook/ 开头'}), 400
    db.set_feishu_webhook(user['id'], webhook)
    db.set_feishu_secret(user['id'], secret)
    return jsonify({'status': 'success', 'configured': bool(webhook)})


@app.route('/api/feishu/test', methods=['POST'])
def api_feishu_test():
    """测试飞书 Webhook 连接"""
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    webhook = db.get_feishu_webhook(user['id'])
    secret = db.get_feishu_secret(user['id'])
    if not webhook:
        return jsonify({'status': 'error', 'error': '请先配置飞书 Webhook'}), 400
    result = feishu_push.send_feishu_text(webhook, '✅ 工具集 v5.0 飞书推送测试成功！', secret=secret)
    if result['ok']:
        return jsonify({'status': 'success', 'message': '测试消息已发送'})
    else:
        return jsonify({'status': 'error', 'error': result.get('error', '发送失败')}), 502


@app.route('/api/feishu/push', methods=['POST'])
def api_feishu_push():
    """
    通用飞书推送接口
    Body: {
        "type": "text" | "card" | "weekly" | "meeting",
        "title": "标题",
        "content": "正文内容（text类型为纯文本，card类型为Markdown）",
        "summary": "概要（weekly/meeting类型）",
        "highlights": "重点（weekly类型）",
        "plans": "计划（weekly类型）",
        "decisions": "决议（meeting类型）",
        "todos": "待办（meeting类型）",
        "url": "源链接（可选）"
    }
    """
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    webhook = db.get_feishu_webhook(user['id'])
    secret = db.get_feishu_secret(user['id'])
    if not webhook:
        return jsonify({'status': 'error', 'error': '请先在设置中配置飞书 Webhook'}), 400

    data = request.get_json(silent=True) or {}
    msg_type = data.get('type', 'text')
    title = data.get('title', '工具集推送')
    source_url = data.get('url')

    try:
        if msg_type == 'text':
            content = data.get('content', '')
            result = feishu_push.send_feishu_text(webhook, content, secret=secret)
        elif msg_type == 'card':
            content = data.get('content', '')
            result = feishu_push.send_feishu_card(webhook, title, content, header_color='purple', link_url=source_url, secret=secret)
        elif msg_type == 'weekly':
            result = feishu_push.send_weekly_report(
                webhook, title,
                summary=data.get('summary', ''),
                highlights=data.get('highlights', ''),
                plans=data.get('plans', ''),
                source_url=source_url,
                secret=secret
            )
        elif msg_type == 'meeting':
            result = feishu_push.send_meeting_minutes(
                webhook, title,
                summary=data.get('summary', ''),
                decisions=data.get('decisions', ''),
                todos=data.get('todos', ''),
                source_url=source_url,
                secret=secret
            )
        else:
            return jsonify({'status': 'error', 'error': f'不支持的消息类型: {msg_type}'}), 400

        if result['ok']:
            return jsonify({'status': 'success', 'message': '推送成功'})
        else:
            return jsonify({'status': 'error', 'error': result.get('error', '推送失败')}), 502
    except Exception as e:
        logger.error(f'飞书推送异常: {e}')
        return jsonify({'status': 'error', 'error': f'推送异常: {str(e)}'}), 500


# ==================== v5.2 多设备数据同步 API ====================

@app.route('/api/sync/pull', methods=['GET'])
def api_sync_pull():
    """拉取云端同步数据"""
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    try:
        data = db.get_all_sync_states(user['id'])
        return jsonify({'status': 'success', 'data': data, 'server_time': time.time()})
    except Exception as e:
        logger.error(f'同步拉取异常: {e}')
        return jsonify({'status': 'error', 'error': str(e)}), 500


@app.route('/api/sync/push', methods=['POST'])
def api_sync_push():
    """推送本地数据到云端（支持多类型批量）"""
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    data = request.get_json(silent=True) or {}
    items = data.get('items', {})
    if not isinstance(items, dict) or not items:
        return jsonify({'status': 'error', 'error': '无效的同步数据'}), 400

    results = {}
    for dtype, content in items.items():
        if dtype in db.SYNC_TYPES:
            try:
                ts = db.set_sync_state(user['id'], dtype, content)
                results[dtype] = {'status': 'success', 'updated_at': ts}
            except Exception as e:
                results[dtype] = {'status': 'error', 'error': str(e)}
        else:
            results[dtype] = {'status': 'error', 'error': '不支持的同步类型'}
    return jsonify({'status': 'success', 'results': results, 'server_time': time.time()})


@app.route('/api/sync/status', methods=['GET'])
def api_sync_status():
    """获取同步状态（各类型最新更新时间）"""
    user = auth.get_current_user()
    if not user:
        return jsonify({'status': 'error', 'error': '请先登录'}), 401
    try:
        status = db.get_sync_status(user['id'])
        return jsonify({'status': 'success', 'sync_status': status, 'server_time': time.time()})
    except Exception as e:
        return jsonify({'status': 'error', 'error': str(e)}), 500


# ==================== 音频转写 API（DashScope Paraformer ASR） ====================

@app.route('/api/upload-audio', methods=['POST'])
@login_required
def api_upload_audio():
    """上传音频文件 — 异步处理：先保存文件，后台线程上传DashScope并提交转写"""

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请联系管理员设置API Key'}), 503

    # 获取上传的文件
    if 'audio' not in request.files:
        return jsonify({'error': '未收到音频文件'}), 400

    audio_file = request.files['audio']
    if not audio_file.filename:
        return jsonify({'error': '文件名为空'}), 400

    # 检查文件大小（限制200MB）
    audio_file.seek(0, 2)
    file_size = audio_file.tell()
    audio_file.seek(0)
    if file_size > 200 * 1024 * 1024:
        return jsonify({'error': '文件过大，最大支持200MB'}), 400

    # 检查文件类型
    allowed_extensions = {'.mp3', '.wav', '.m4a', '.aac', '.flac', '.ogg', '.wma', '.webm'}
    ext = os.path.splitext(audio_file.filename)[1].lower()
    if ext not in allowed_extensions:
        return jsonify({'error': f'不支持的文件格式: {ext}，支持: {", ".join(allowed_extensions)}'}), 400

    try:
        import uuid as _uuid

        # Step 1: 保存文件到磁盘（快速操作）
        task_id = f"asr_{_uuid.uuid4().hex[:16]}"
        saved_filename = f"{task_id}{ext}"
        saved_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)
        audio_file.save(saved_path)
        logger.info(f"音频文件已保存: {saved_filename}, {file_size} bytes, task={task_id}")

        # Step 2: 创建后台任务记录
        db.create_task(task_id, 'audio_transcription')
        db.update_task(task_id, status='uploading', progress=10)

        # Step 3: 启动后台线程处理 DashScope 上传 + 转写提交
        api_key = ai_config.get('api_key', '')
        orig_filename = audio_file.filename

        def _background_process():
            """后台线程：上传文件到DashScope → 提交转写任务"""
            try:
                import requests as req

                # Step A: 上传文件到 DashScope 文件服务
                db.update_task(task_id, status='uploading', progress=20)
                logger.info(f"[后台] 开始上传到DashScope: {task_id}")

                with open(saved_path, 'rb') as f:
                    upload_resp = req.post(
                        'https://dashscope.aliyuncs.com/api/v1/uploads',
                        headers={'Authorization': f'Bearer {api_key}'},
                        files={
                            'file': (orig_filename, f, 'application/octet-stream'),
                            'model': (None, 'paraformer-v2'),
                            'action': (None, 'put'),
                        },
                        timeout=300  # 5分钟超时，大文件需要更长时间
                    )

                if upload_resp.status_code != 200:
                    logger.error(f"[后台] DashScope上传失败: {upload_resp.status_code} - {upload_resp.text[:300]}")
                    db.update_task(task_id, status='failed', error=f'文件上传失败({upload_resp.status_code})')
                    return

                upload_data = upload_resp.json()
                file_url = upload_data.get('output', {}).get('upload_url', '')
                if not file_url:
                    file_url = upload_data.get('data', {}).get('url', '')

                if not file_url:
                    logger.error(f"[后台] DashScope上传返回异常: {upload_data}")
                    db.update_task(task_id, status='failed', error='文件上传返回异常')
                    return

                logger.info(f"[后台] 文件上传成功: {file_url[:80]}...")

                # Step B: 提交转写任务
                db.update_task(task_id, status='submitting', progress=50)

                transcription_resp = req.post(
                    'https://dashscope.aliyuncs.com/api/v1/services/audio/asr/transcription',
                    headers={
                        'Authorization': f'Bearer {api_key}',
                        'Content-Type': 'application/json',
                        'X-DashScope-Async': 'enable',
                    },
                    json={
                        'model': 'paraformer-v2',
                        'input': {'file_urls': [file_url]},
                        'parameters': {
                            'language_hints': ['zh', 'en'],
                            'disfluency_removal': False,
                            'paragraph': True,
                        }
                    },
                    timeout=30
                )

                if transcription_resp.status_code != 200:
                    logger.error(f"[后台] 转写提交失败: {transcription_resp.status_code} - {transcription_resp.text[:300]}")
                    db.update_task(task_id, status='failed', error=f'转写任务提交失败({transcription_resp.status_code})')
                    return

                task_data = transcription_resp.json()
                dashscope_task_id = task_data.get('output', {}).get('task_id', '')

                if not dashscope_task_id:
                    logger.error(f"[后台] 转写任务提交返回异常: {task_data}")
                    db.update_task(task_id, status='failed', error='转写任务提交返回异常')
                    return

                logger.info(f"[后台] 转写任务已提交: {dashscope_task_id}")
                db.update_task(task_id, status='transcribing', progress=60,
                               result={'dashscope_task_id': dashscope_task_id})

            except req.exceptions.Timeout:
                logger.error(f"[后台] DashScope上传超时: {task_id}")
                db.update_task(task_id, status='failed', error='文件上传超时，请尝试较小的文件')
            except Exception as e:
                logger.error(f"[后台] 音频处理失败: {traceback.format_exc()}")
                db.update_task(task_id, status='failed', error=str(e))
            finally:
                # 清理临时文件
                try:
                    if os.path.exists(saved_path):
                        os.remove(saved_path)
                except Exception:
                    pass

        # 启动后台线程
        thread = threading.Thread(target=_background_process, daemon=True)
        thread.start()

        # 立即返回任务ID
        return jsonify({
            'status': 'success',
            'task_id': task_id
        })

    except Exception as e:
        logger.error(f"音频上传失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/transcription-status/<task_id>')
@login_required
def api_transcription_status(task_id):
    """查询转写任务状态 — 支持后台任务 + DashScope双重状态查询"""

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置'}), 503

    try:
        # Step 1: 查询本地后台任务状态
        local_task = db.get_task(task_id)
        if not local_task:
            return jsonify({'error': '任务不存在'}), 404

        local_status = local_task.get('status', 'unknown')
        local_progress = local_task.get('progress', 0)
        local_error = local_task.get('error')

        # 如果任务失败，返回错误
        if local_status == 'failed':
            return jsonify({
                'status': 'FAILED',
                'task_id': task_id,
                'error': local_error or '处理失败'
            })

        # 如果还在上传/提交阶段，返回中间状态
        if local_status in ('uploading', 'submitting', 'pending'):
            status_map = {
                'pending': 'UPLOADING',
                'uploading': 'UPLOADING',
                'submitting': 'SUBMITTING',
            }
            return jsonify({
                'status': status_map.get(local_status, 'PENDING'),
                'task_id': task_id,
                'progress': local_progress
            })

        # 如果已经开始转写，查询 DashScope 状态
        if local_status == 'transcribing':
            result_data = local_task.get('result', {})
            if isinstance(result_data, str):
                try:
                    result_data = json.loads(result_data)
                except (json.JSONDecodeError, TypeError):
                    result_data = {}

            dashscope_task_id = result_data.get('dashscope_task_id', '')
            if not dashscope_task_id:
                return jsonify({'status': 'PENDING', 'task_id': task_id, 'progress': 60})

            # 查询 DashScope 转写状态
            import requests as req
            resp = req.get(
                f'https://dashscope.aliyuncs.com/api/v1/tasks/{dashscope_task_id}',
                headers={'Authorization': f'Bearer {ai_config.get("api_key", "")}'},
                timeout=15
            )

            if resp.status_code != 200:
                return jsonify({'error': f'查询失败({resp.status_code})'}), 502

            data = resp.json()
            task_status = data.get('output', {}).get('task_status', 'UNKNOWN')

            result = {
                'status': task_status,
                'task_id': task_id,
                'progress': 80
            }

            # 转写完成，获取结果
            if task_status == 'SUCCEEDED':
                results = data.get('output', {}).get('results', [])
                transcript_text = ''

                for r in results:
                    transcription_url = r.get('transcription_url', '')
                    if transcription_url:
                        try:
                            tr_resp = req.get(transcription_url, timeout=15)
                            if tr_resp.status_code == 200:
                                tr_data = tr_resp.json()
                                # 提取转写文本 — 尝试多种格式
                                transcripts = tr_data.get('transcripts', [])
                                for t in transcripts:
                                    transcript_text += t.get('text', '') + '\n'

                                if not transcript_text:
                                    sentences = tr_data.get('sentences', [])
                                    for s in sentences:
                                        transcript_text += s.get('text', '') + '\n'

                                if not transcript_text and tr_data.get('text'):
                                    transcript_text = tr_data['text']
                        except Exception as e:
                            logger.warning(f"获取转写结果失败: {e}")

                result['transcript'] = transcript_text.strip()
                result['progress'] = 100
                db.update_task(task_id, status='completed', progress=100,
                               result={'dashscope_task_id': dashscope_task_id, 'transcript': transcript_text.strip()})
                logger.info(f"转写完成，文本长度: {len(transcript_text)}")

            elif task_status == 'FAILED':
                result['error'] = data.get('output', {}).get('message', '转写失败')
                db.update_task(task_id, status='failed', error=result['error'])

            return jsonify(result)

        # 如果已完成，从数据库返回结果
        if local_status == 'completed':
            result_data = local_task.get('result', {})
            if isinstance(result_data, str):
                try:
                    result_data = json.loads(result_data)
                except (json.JSONDecodeError, TypeError):
                    result_data = {}
            return jsonify({
                'status': 'SUCCEEDED',
                'task_id': task_id,
                'progress': 100,
                'transcript': result_data.get('transcript', '')
            })

        return jsonify({'status': 'UNKNOWN', 'task_id': task_id})

    except Exception as e:
        logger.error(f"查询转写状态失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# ==================== 实时语音识别 WebSocket 代理 ====================

@sock.route('/ws/realtime-asr')
def realtime_asr_proxy(ws):
    """
    WebSocket 代理：浏览器 ↔ 后端 ↔ DashScope Paraformer 实时ASR
    流程：
    1. 浏览器连接 → 后端连接 DashScope WebSocket → 发送 run-task
    2. 浏览器发送 PCM 音频块 → 后端转发到 DashScope
    3. DashScope 返回识别结果 → 后端转发到浏览器
    4. 浏览器发送 stop → 后端发送 finish-task → 关闭连接
    """
    import threading
    import uuid as _uuid

    # Step 1: 认证
    user = auth.get_current_user()
    if not user:
        try:
            ws.send(json.dumps({'type': 'error', 'message': '请先登录'}))
        except Exception:
            pass
        ws.close()
        return

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        try:
            ws.send(json.dumps({'type': 'error', 'message': 'AI功能未配置，请联系管理员'}))
        except Exception:
            pass
        ws.close()
        return

    api_key = ai_config.get('api_key', '')
    task_id = str(_uuid.uuid4())

    # Step 2: 连接 DashScope WebSocket
    ds_url = 'wss://dashscope.aliyuncs.com/api-ws/v1/inference/'
    try:
        ds_ws = _ws_client.create_connection(
            ds_url,
            header=[f'Authorization: Bearer {api_key}'],
            timeout=15,
            enable_multithread=True,
        )
    except Exception as e:
        logger.error(f"连接DashScope WebSocket失败: {e}")
        try:
            ws.send(json.dumps({'type': 'error', 'message': f'连接AI服务失败: {str(e)}'}))
        except Exception:
            pass
        ws.close()
        return

    # Step 3: 发送 run-task 指令
    run_task_msg = {
        'header': {
            'action': 'run-task',
            'task_id': task_id,
            'streaming': 'duplex',
        },
        'payload': {
            'task_group': 'audio',
            'task': 'asr',
            'function': 'recognition',
            'model': 'paraformer-realtime-v2',
            'input': {},
            'parameters': {
                'format': 'pcm',
                'sample_rate': 16000,
                'language_hints': ['zh', 'en'],
                'semantic_punctuation_enabled': False,
                'disfluency_removal_enabled': False,
                'punctuation_prediction_enabled': True,
                'inverse_text_normalization_enabled': True,
                'heartbeat': True,
            }
        }
    }

    try:
        ds_ws.send(json.dumps(run_task_msg))
    except Exception as e:
        logger.error(f"发送run-task失败: {e}")
        try:
            ws.send(json.dumps({'type': 'error', 'message': '启动识别任务失败'}))
        except Exception:
            pass
        ds_ws.close()
        ws.close()
        return

    # Step 4: 等待 task-started 事件
    try:
        resp = ds_ws.recv()
        resp_data = json.loads(resp)
        event = resp_data.get('header', {}).get('event', '')
        if event != 'task-started':
            err_msg = resp_data.get('header', {}).get('error_message', '任务启动失败')
            logger.error(f"DashScope task-started 异常: {resp_data}")
            try:
                ws.send(json.dumps({'type': 'error', 'message': err_msg}))
            except Exception:
                pass
            ds_ws.close()
            ws.close()
            return
    except Exception as e:
        logger.error(f"等待task-started失败: {e}")
        try:
            ws.send(json.dumps({'type': 'error', 'message': '等待AI服务响应超时'}))
        except Exception:
            pass
        ds_ws.close()
        ws.close()
        return

    # 通知浏览器：可以开始发送音频
    try:
        ws.send(json.dumps({'type': 'ready', 'message': '实时识别已就绪'}))
    except Exception:
        ds_ws.close()
        return

    logger.info(f"[实时ASR] 任务已启动: {task_id}")

    # Step 5: 启动接收线程 — DashScope → 浏览器
    ws_closed = threading.Event()

    def _dashscope_receiver():
        """接收 DashScope 的识别结果，转发给浏览器"""
        try:
            while not ws_closed.is_set():
                try:
                    resp = ds_ws.recv()
                    # DashScope 返回速度受模型影响，recv 本身不设超时
                    # 靠 heartbeat 保活，task-failed/task-finished 自然退出
                except _ws_client.WebSocketTimeoutException:
                    continue
                except _ws_client.WebSocketConnectionClosedException:
                    break
                except Exception:
                    break

                if isinstance(resp, bytes):
                    continue

                try:
                    data = json.loads(resp)
                except (json.JSONDecodeError, TypeError):
                    continue

                event = data.get('header', {}).get('event', '')

                if event == 'result-generated':
                    sentence = data.get('payload', {}).get('output', {}).get('sentence', {})
                    text = sentence.get('text', '')
                    is_final = sentence.get('sentence_end', False)
                    # 跳过心跳包
                    if sentence.get('heartbeat'):
                        continue
                    try:
                        ws.send(json.dumps({
                            'type': 'transcript',
                            'text': text,
                            'is_final': is_final,
                            'begin_time': sentence.get('begin_time', 0),
                            'end_time': sentence.get('end_time', 0),
                        }))
                    except Exception:
                        break

                elif event == 'task-finished':
                    try:
                        ws.send(json.dumps({'type': 'finished'}))
                    except Exception:
                        pass
                    break

                elif event == 'task-failed':
                    err_msg = data.get('header', {}).get('error_message', '识别失败')
                    logger.error(f"[实时ASR] 任务失败: {err_msg}")
                    try:
                        ws.send(json.dumps({'type': 'error', 'message': err_msg}))
                    except Exception:
                        pass
                    break

        except Exception as e:
            logger.error(f"[实时ASR] 接收线程异常: {e}")
        finally:
            ws_closed.set()

    receiver_thread = threading.Thread(target=_dashscope_receiver, daemon=True)
    receiver_thread.start()

    # Step 6: 主循环 — 接收浏览器消息，转发音频到 DashScope
    try:
        while not ws_closed.is_set():
            try:
                message = ws.receive()
            except Exception:
                break

            if message is None:
                # 浏览器断开连接
                break

            if isinstance(message, bytes):
                # PCM 音频块 → 转发到 DashScope
                try:
                    ds_ws.send_binary(message)
                except Exception:
                    logger.warning("[实时ASR] 转发音频失败，连接可能已断开")
                    break

            elif isinstance(message, str):
                try:
                    data = json.loads(message)
                except (json.JSONDecodeError, TypeError):
                    continue

                if data.get('type') == 'stop':
                    # 发送 finish-task 指令
                    finish_msg = {
                        'header': {
                            'action': 'finish-task',
                            'task_id': task_id,
                            'streaming': 'duplex',
                        },
                        'payload': {
                            'input': {}
                        }
                    }
                    try:
                        ds_ws.send(json.dumps(finish_msg))
                    except Exception:
                        break
                    # 等待 task-finished 后由接收线程关闭
                    break

    except Exception as e:
        logger.error(f"[实时ASR] 主循环异常: {e}")
    finally:
        ws_closed.set()
        try:
            ds_ws.close()
        except Exception:
            pass
        logger.info(f"[实时ASR] 任务结束: {task_id}")


@app.route('/health')
def health():
    """健康检查端点，供 Railway / 负载均衡使用"""
    return jsonify({'status': 'ok', 'version': _STATIC_VERSION, 'pid': os.getpid()})


@app.route('/favicon.ico')
def favicon():
    """浏览器默认请求的favicon — 返回工具箱emoji SVG"""
    svg = "<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 100 100'><text y='.9em' font-size='90'>🧰</text></svg>"
    resp = make_response(svg, 200)
    resp.headers['Content-Type'] = 'image/svg+xml'
    resp.headers['Cache-Control'] = 'public, max-age=86400'
    return resp


@app.route('/api/debug')
@login_required
def api_debug():
    """诊断端点：返回部署版本、内存、任务状态等信息（需登录，生产环境禁用）"""
    if _is_production:
        return jsonify({'error': 'debug endpoint disabled in production'}), 403
    # 安全：要求登录才能访问
    import psutil
    p = psutil.Process()
    mem_info = p.memory_info()

    # 统计上传目录中的文件
    upload_dir = app.config['UPLOAD_FOLDER']
    uploaded_files = []
    total_upload_size = 0
    if os.path.exists(upload_dir):
        for fname in os.listdir(upload_dir):
            fpath = os.path.join(upload_dir, fname)
            if os.path.isfile(fpath):
                fsize = os.path.getsize(fpath)
                uploaded_files.append({'name': fname, 'size_kb': round(fsize / 1024, 1)})
                total_upload_size += fsize
            elif os.path.isdir(fpath):
                # 子目录（如 _task_meta, _chunk_meta）
                file_count = len(os.listdir(fpath))
                uploaded_files.append({'name': f'{fname}/', 'files': file_count})

    # 后台任务状态
    active_tasks = {
        tid: {'status': t.get('status'), 'age_seconds': round(time.time() - t.get('created_at', time.time()), 0)}
        for tid, t in _background_tasks.items()
    }

    return jsonify({
        'status': 'ok',
        'pid': os.getpid(),
        'memory': {
            'rss_mb': round(mem_info.rss / 1024 / 1024, 1),
            'vms_mb': round(mem_info.vms / 1024 / 1024, 1),
        },
        'background_tasks': {
            'total': len(_background_tasks),
            'active': active_tasks,
        },
        'uploads': {
            'files': uploaded_files[:20],
            'total_size_mb': round(total_upload_size / 1024 / 1024, 2),
        },
        'config': {
            'upload_folder': upload_dir,
            'pdf_folder': app.config['PDF_FOLDER'],
            'is_cloud': bool(os.environ.get('RAILWAY_STATIC_URL') or os.environ.get('PORT')),
        },
        'database': db.check_db(),
    })

# === 测试报告分析 API ===
@app.route('/api/test-report-analyze', methods=['POST'])
def api_test_report_upload():
    """第一阶段：上传文件，返回Sheet列表和文件ID"""
    if 'file' not in request.files:
        return jsonify({'error': '请选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '请选择文件'}), 400

    filename_lower = file.filename.lower()
    if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
        return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400

    orig_ext = os.path.splitext(file.filename)[1].lower()
    if orig_ext not in ('.xlsx', '.xls', '.csv'):
        orig_ext = '.xlsx'

    try:
        file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"test_report_{file_id}{orig_ext}")
        file.save(file_path)

        logger.info(f"========== 收到测试报告: {file.filename}, ID: {file_id} ==========")

        reader = ExcelReader(file_path)
        reader.open()
        sheet_names = reader.get_sheet_names()
        reader.close()

        return jsonify({
            'status': 'success',
            'data': {
                'file_id': file_id,
                'file_name': file.filename,
                'file_basename': os.path.splitext(file.filename)[0],
                'sheet_names': sheet_names,
                'sheet_count': len(sheet_names)
            }
        })

    except Exception as e:
        logger.error(f"测试报告上传失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/test-report-analyze-sheet', methods=['POST'])
def api_test_report_analyze_sheet():
    """第二阶段：分析指定Sheet的详细内容"""
    data = request.json or {}
    file_id = data.get('file_id', '')
    sheet_name = data.get('sheet_name', '')

    if not file_id:
        return jsonify({'error': '缺少file_id'}), 400
    if not _validate_file_id(file_id):
        return jsonify({'error': '无效的文件ID'}), 400
    if not sheet_name:
        return jsonify({'error': '缺少sheet_name'}), 400

    file_path = None
    for ext in ['.xlsx', '.xls', '.csv']:
        candidate = os.path.join(app.config['UPLOAD_FOLDER'], f"test_report_{file_id}{ext}")
        if os.path.exists(candidate):
            file_path = candidate
            break

    if not file_path:
        return jsonify({'error': '文件不存在，可能已过期'}), 404

    # 创建后台任务（避免同步分析超过 Railway 5 分钟 HTTP 超时）
    task_id = hashlib.md5(f"test_report_{file_id}_{sheet_name}_{time.time()}".encode()).hexdigest()[:16]
    task_data = {
        'status': 'processing',
        'result': None,
        'error': None,
        'created_at': time.time()
    }
    _background_tasks[task_id] = task_data
    _save_task_meta(task_id, task_data)

    def _do_test_report_analysis():
        try:
            logger.info(f"[分析任务 {task_id}] 开始分析: file={file_path}, sheet={sheet_name}")
            gc.collect()
            t0 = time.time()

            def _progress_cb(percent, msg):
                _background_tasks[task_id]['progress'] = percent
                _background_tasks[task_id]['progress_msg'] = msg
                _save_task_meta(task_id, _background_tasks[task_id])

            result = _analyze_sheet_detail(file_path, sheet_name, progress_cb=_progress_cb)
            elapsed = time.time() - t0
            logger.info(f"[分析任务 {task_id}] 分析完成，耗时 {elapsed:.1f}s, 测试项数={len(result.get('test_items', []))}")
            _background_tasks[task_id]['result'] = result
            _background_tasks[task_id]['status'] = 'done'
            _background_tasks[task_id]['progress'] = 100
            _save_task_meta(task_id, _background_tasks[task_id])
        except Exception as e:
            error_detail = str(e) if str(e) else f'{type(e).__name__} (无详细错误信息)'
            logger.error(f"[分析任务 {task_id}] 分析失败: {traceback.format_exc()}")
            _background_tasks[task_id]['error'] = error_detail
            _background_tasks[task_id]['status'] = 'error'
            _save_task_meta(task_id, _background_tasks[task_id])

    thread = threading.Thread(target=_do_test_report_analysis, daemon=True)
    thread.start()

    return jsonify({'status': 'success', 'data': {'task_id': task_id}})


@app.route('/api/test-report-debug', methods=['POST'])
@login_required
def api_test_report_debug():
    """诊断端点：返回Excel解析的详细信息（需登录，生产环境禁用）"""
    if _is_production:
        return jsonify({'error': 'debug endpoint disabled in production'}), 403
    data = request.json or {}
    file_id = data.get('file_id', '')
    sheet_name = data.get('sheet_name', '')

    if not file_id:
        return jsonify({'error': '缺少file_id'}), 400
    if not _validate_file_id(file_id):
        return jsonify({'error': '无效的文件ID'}), 400

    file_path = None
    for ext in ['.xlsx', '.xls', '.csv']:
        candidate = os.path.join(app.config['UPLOAD_FOLDER'], f"test_report_{file_id}{ext}")
        if os.path.exists(candidate):
            file_path = candidate
            break

    if not file_path:
        return jsonify({'error': '文件不存在，可能已过期'}), 404

    try:
        if not sheet_name:
            reader = ExcelReader(file_path)
            reader.open()
            sheet_name = reader.sheetnames[0] if reader.sheetnames else 'Sheet1'
            reader.close()

        result = _analyze_sheet_detail(file_path, sheet_name, return_debug=True)
        return jsonify({'status': 'success', 'data': result})
    except Exception as e:
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500


@app.route('/api/test-report-pdf', methods=['POST'])
def api_test_report_pdf():
    """将测试报告分析结果导出为PDF（含Motorola水印+日期）"""
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': '请求数据格式错误'}), 400

    analysis_data = data.get('analysis_data', {})
    file_name = data.get('file_name', '')
    sheet_name = data.get('sheet_name', '')
    ai_analysis = data.get('ai_analysis', '')

    if not analysis_data:
        return jsonify({'error': '缺少分析数据'}), 400

    try:
        html_content = _build_test_report_pdf_html(analysis_data, file_name, sheet_name, ai_analysis)

        import tempfile as tf
        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name

        safe_name = re.sub(r'[\\/:*?"<>|]', '_', file_name or 'test_report')
        pdf_filename = f"{safe_name}_{datetime.now(_CST).strftime('%Y%m%d_%H%M%S')}.pdf"
        download_name = f"{safe_name}.pdf"
        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

        _render_pdf(html_path, pdf_path,
                    margin={'top': '8mm', 'bottom': '8mm', 'left': '8mm', 'right': '8mm'},
                    extra_wait_ms=1000)

        try:
            os.unlink(html_path)
        except Exception:
            pass

        return jsonify({
            'status': 'success',
            'filename': pdf_filename,
            'download_name': download_name
        })
    except Exception as e:
        logger.error(f"测试报告PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/test-report-download/<filename>')
def api_test_report_download(filename):
    """下载生成的PDF文件"""
    safe_name = os.path.basename(filename)
    pdf_path = os.path.join(app.config['PDF_FOLDER'], safe_name)
    if not os.path.exists(pdf_path):
        return jsonify({'error': '文件不存在'}), 404
    download_name = request.args.get('download_name', safe_name)
    return send_file(pdf_path, as_attachment=True, download_name=download_name)


@app.route('/api/test-report-ai-analysis', methods=['POST'])
@login_required_or_guest
def api_test_report_ai_analysis():
    """v3.0: AI 增强测试报告分析 — 生成测试总结和风险洞察"""

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请在设置页面配置 API Key'}), 503

    data = request.get_json(silent=True) or {}
    analysis = data.get('analysis', {})
    if not analysis:
        return jsonify({'error': '缺少分析数据'}), 400

    # 构建AI提示词
    project_info = analysis.get('project_info', {})
    stats = analysis.get('stats', {})
    summary = analysis.get('summary', {})
    key_findings = analysis.get('key_findings', [])
    test_items = analysis.get('test_items', [])

    # 截取前30个测试项以控制token
    items_text = ''
    for item in test_items[:30]:
        items_text += f"- {item.get('name', 'N/A')}: {item.get('result', 'N/A')}"
        if item.get('remark'):
            items_text += f" ({item['remark']})"
        items_text += "\n"

    prompt = f"""你是一位资深软件测试专家，请基于以下测试报告数据生成专业的AI分析总结。

## 项目信息
- 项目名称: {project_info.get('name', 'N/A')}
- 软件版本: {project_info.get('version', 'N/A')}
- 测试日期: {project_info.get('date', 'N/A')}

## 测试统计
- 总测试项: {stats.get('total', 0)}
- 通过: {stats.get('pass', 0)}
- 失败: {stats.get('fail', 0)}
- 阻塞: {stats.get('block', 0)}
- 通过率: {stats.get('pass_rate', 'N/A')}

## 执行摘要
{summary.get('text', 'N/A')}

## 关键发现
{chr(10).join(f'- {f}' for f in key_findings) if key_findings else 'N/A'}

## 测试项详情（前30项）
{items_text or 'N/A'}

请按以下格式输出分析：

### 🎯 总体评估
（1-2段话总结测试整体状况，包括通过率评价和发布建议）

### ⚠️ 风险洞察
（列出3-5个关键风险点，每个风险用一句话描述）

### 📈 质量趋势
（分析测试质量趋势，指出薄弱环节）

### 💡 改进建议
（列出3-5条可操作的改进建议，每条用一句话描述）

请使用简洁专业的中文，避免空话套话。"""

    try:
        messages = [{'role': 'user', 'content': prompt}]
        reply = _call_ai(messages, max_tokens=1500, temperature=0.3, timeout=60)
        return jsonify({'status': 'success', 'analysis': reply})
    except Exception as e:
        logger.error(f'AI测试报告分析失败: {e}')
        return jsonify({'error': f'AI分析失败: {str(e)}'}), 502


@app.route('/api/test-report-ai-stream', methods=['POST'])
@login_required_or_guest
def api_test_report_ai_stream():
    """SSE 流式版：AI 增强测试报告分析 — 边生成边输出"""

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请在设置页面配置 API Key'}), 503

    data = request.get_json(silent=True) or {}
    analysis = data.get('analysis', {})
    if not analysis:
        return jsonify({'error': '缺少分析数据'}), 400

    project_info = analysis.get('project_info', {})
    stats = analysis.get('stats', {})
    summary = analysis.get('summary', {})
    key_findings = analysis.get('key_findings', [])
    test_items = analysis.get('test_items', [])

    items_text = ''
    for item in test_items[:30]:
        items_text += f"- {item.get('name', 'N/A')}: {item.get('result', 'N/A')}"
        if item.get('remark'):
            items_text += f" ({item['remark']})"
        items_text += "\n"

    prompt = f"""你是一位资深软件测试专家，请基于以下测试报告数据生成专业的AI分析总结。

## 项目信息
- 项目名称: {project_info.get('name', 'N/A')}
- 软件版本: {project_info.get('version', 'N/A')}
- 测试日期: {project_info.get('date', 'N/A')}

## 测试统计
- 总测试项: {stats.get('total', 0)}
- 通过: {stats.get('pass', 0)}
- 失败: {stats.get('fail', 0)}
- 阻塞: {stats.get('block', 0)}
- 通过率: {stats.get('pass_rate', 'N/A')}

## 执行摘要
{summary.get('text', 'N/A')}

## 关键发现
{chr(10).join(f'- {f}' for f in key_findings) if key_findings else 'N/A'}

## 测试项详情（前30项）
{items_text or 'N/A'}

请按以下格式输出分析：

### 🎯 总体评估
（1-2段话总结测试整体状况，包括通过率评价和发布建议）

### ⚠️ 风险洞察
（列出3-5个关键风险点，每个风险用一句话描述）

### 📈 质量趋势
（分析测试质量趋势，指出薄弱环节）

### 💡 改进建议
（列出3-5条可操作的改进建议，每条用一句话描述）

请使用简洁专业的中文，避免空话套话。"""

    messages = [{'role': 'user', 'content': prompt}]
    return Response(
        stream_with_context(_call_ai_stream(messages, max_tokens=1500, temperature=0.3)),
        content_type='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no', 'Connection': 'keep-alive'}
    )
    if not headers:
        return col_map
    # 清理表头：去掉换行符、@符号后面的内容、多余空格
    def clean_header(h):
        h = str(h).replace('\n', ' ').replace('\r', ' ')
        # 去掉 @xxx 部分
        if '@' in h:
            h = h.split('@')[0].strip()
        return h.lower().strip()
    headers_lower = [clean_header(h) for h in headers]

    for i, h in enumerate(headers_lower):
        if any(kw in h for kw in ['测试项', '测试内容', '测试用例', '名称', 'name', 'test item', 'test case', 'case', '指标', '项目', '检查项', 'test items']):
            col_map['name'] = i
        elif any(kw in h for kw in ['模块', 'module', 'component', 'commponent', '组件', '功能', '分类', 'category', '类型', '领域', '专项']):
            col_map['module'] = i
        elif any(kw in h for kw in ['severity', '严重程度', '严重性', '等级', 'level', '优先级', 'priority']):
            col_map['severity'] = i
        elif any(kw in h for kw in ['结果', 'result', 'pass/fail', '通过', '状态', 'status', '结论', '判定']):
            col_map['result'] = i
        elif any(kw in h for kw in ['key issue', '核心问题', '关键问题', '主要问题']):
            col_map['key_issue'] = i
        elif any(kw in h for kw in ['comment', '备注', 'remark', 'note', '说明', 'comment.', 'comments']):
            col_map['comment'] = i
        elif any(kw in h for kw in ['原因', 'reason', '描述', '问题', 'issue', '问题描述', '问题描述']):
            col_map['reason'] = i
        elif any(kw in h for kw in ['目标', 'target', '要求', '门槛', '标准值', 'sr6 target', 'sr5 target', 'sr4 target', 'sr target', '基线', 'baseline', '阈值', 'threshold']):
            col_map['target'] = i
        elif any(kw in h for kw in ['实测', '实际', 'actual', '当前', 'current', '结果值', '测量', 'measured', 'cwv']):
            col_map['actual'] = i
        elif any(kw in h for kw in ['风险', 'risk']):
            col_map['risk'] = i

    # 如果没有找到 result 列，但有 actual(CWV) 列，用 actual 作为 result
    # (CWV 值可能是 Pass/Fail 文本或数值，后续通过 _compare_target_actual 判断)
    if 'result' not in col_map and 'actual' in col_map:
        col_map['result'] = col_map['actual']

    if 'result' not in col_map:
        for i, h in enumerate(headers_lower):
            if any(kw in h for kw in ['pass', 'fail', '通过', '不通过', '结论', '判定']):
                col_map['result'] = i
                break

    if 'name' not in col_map:
        col_map['name'] = 0

    if 'reason' not in col_map and len(headers) > 0:
        col_map['reason'] = len(headers) - 1

    # 优先使用 SR6 Target 作为目标值（覆盖通用 target 列）
    # 当存在多个 Target 列（SR4/SR5/SR6）时，优先取 SR6 Target
    for i, h in enumerate(headers_lower):
        if 'sr6' in h and ('target' in h or '目标' in h):
            col_map['target'] = i
            break

    return col_map


def _get_cell_value(cells, idx):
    if idx < 0 or idx >= len(cells):
        return ''
    return str(cells[idx]).strip() if cells[idx] else ''


def _extract_action_items(text):
    """
    从 Key issue / Comment 文本中提取待办事项。
    识别模式：
    - 显式标记: TODO, 待办, 待跟进, action, 行动项, 下一步, follow up
    - 动作引导词: 需要, 需, 应, 应当, 建议, 请, 要求, 必须, 需跟进, 待修复, 待验证
    - 序号列表: 1. 2. 3. 或 ① ② ③ 或 - / * 开头
    - 时态标记: 将, 计划, 拟, 预计
    - 状态标记: 未关闭, 未解决, open, pending, in progress, 进行中
    """
    if not text or not text.strip():
        return []

    import re
    lines = text.replace('\r', '\n').split('\n')
    action_items = []

    # 待办关键词（行首或独立出现）
    action_keywords = [
        'todo', '待办', '待跟进', '待修复', '待验证', '待确认', '待补齐',
        'action', '行动项', '行动', 'follow up', 'follow-up', 'followup',
        '下一步', '后续', '计划', '拟', '预计', '将跟进', '将修复',
        '需要', '需跟进', '需修复', '需验证', '需确认', '需推进',
        '建议', '应', '应当', '应该', '必须', '请',
        '未关闭', '未解决', '未完成', '未达标',
        'open', 'pending', 'in progress', '进行中', '处理中',
        '跟进', '推进', '修复', '验证', '确认', '闭环', '整改',
        '负责', 'owner', 'assign',
    ]

    # 排除词（避免误判）
    exclude_keywords = ['已通过', '已关闭', '已解决', '已完成', 'pass', 'passed', 'done', 'closed', 'resolved']

    for line in lines:
        line = line.strip()
        if not line or len(line) < 3:
            continue

        line_lower = line.lower()

        # 跳过纯排除词行
        if any(line_lower == kw for kw in exclude_keywords):
            continue

        is_action = False

        # 检测序号开头: 1. / 1) / ① / - / * / •
        if re.match(r'^(\d+[.)\]]|[\u2460-\u2473]|[—\-*/•·])\s+', line):
            is_action = True

        # 检测待办关键词
        for kw in action_keywords:
            if kw in line_lower:
                # 排除"已通过"等已完成的描述
                if not any(ex in line_lower for ex in exclude_keywords):
                    is_action = True
                    break

        # 检测"XX:XXX"格式中的动作描述 (如 "责任人:张三 待修复")
        if re.search(r'[:：].*(?:待|需|应|建议|跟进|修复|验证)', line):
            is_action = True

        if is_action:
            # 清理前缀符号
            cleaned = re.sub(r'^(\d+[.)\]]|[\u2460-\u2473]|[—\-*/•·])\s*', '', line)
            cleaned = re.sub(r'^(todo|action|待办|行动项|下一步|follow\s*up)[:：\s]*', '', cleaned, flags=re.IGNORECASE)
            if cleaned and len(cleaned) >= 3:
                # 去重
                if cleaned not in action_items:
                    action_items.append(cleaned)

    return action_items


def _compare_target_actual(target_val, actual_val):
    """
    比较目标值和实测值，返回结果类别和文本描述。
    支持：百分比(1/0.95/95%/100%)、数值、天数等
    """
    def parse_number(val):
        """尝试解析为浮点数，正确处理百分比"""
        if not val:
            return None
        s = str(val).strip()
        has_percent = '%' in s
        # 去掉前缀符号和单位
        s = s.replace('%', '').replace('days', '').replace('day', '').replace('天', '')
        s = re.sub(r'[>=≤<≥~≈约]', '', s).replace('H', '').replace('h', '').replace('fps', '').strip()
        # 去掉括号及后面的内容
        s = re.sub(r'[（(].*$', '', s).strip()
        try:
            num = float(s)
            # 百分比形式：90% -> 0.90, 95% -> 0.95
            if has_percent and num > 1:
                num = num / 100.0
            return num
        except (ValueError, TypeError):
            return None

    target_num = parse_number(target_val)
    actual_num = parse_number(actual_val)

    if target_num is None or actual_num is None:
        return 'unknown', ''

    # 判断通过/不通过
    if actual_num >= target_num:
        return 'pass', f'达标 ({actual_val} >= {target_val})'
    else:
        return 'fail', f'未达标 ({actual_val} < {target_val})'


def _classify_result(text):
    """
    全面分类测试结果状态，返回标准化的状态类别。
    类别: pass / fail / blocked / delayed / n_a / unknown
    - pass:     通过、合格、达标
    - fail:     不通过、失败、不达标、未达标、NG
    - blocked:  阻塞、未执行、跳过、skip
    - delayed:  已延期、延期、推迟
    - n_a:      不适用、N/A
    - unknown:  无法识别
    """
    if not text:
        return 'unknown'

    t = text.strip().lower()

    # N/A 类
    na_words = ['n/a', 'na', '不适用', '无', 'none', 'null', '-']
    if t in na_words:
        return 'n_a'

    # 否定形式优先判断（"不通过"含"通过"，"未达标"含"达标"，须在 pass 之前判断）
    negative_fail_kws = ['不通过', '未通过', '不达标', '未达标', '不合格', '未合格',
                         '不满足', '未满足', '失败', 'fail', 'failed', 'failure',
                         'ng', 'error', 'reject', 'rejected', 'crash',
                         'abort', 'timeout', '超时', '异常', '拒绝']
    for kw in negative_fail_kws:
        if kw in t:
            return 'fail'

    # 延期类
    delayed_words = ['已延期', '延期', '推迟', 'delayed', 'postponed', 'deferred',
                     'pending', '待定', '未开始', 'not started', '暂缓']
    for kw in delayed_words:
        if kw in t:
            return 'delayed'

    # 阻塞类
    blocked_words = ['block', 'blocked', '阻塞', '阻塞中', 'skip', 'skipped',
                     '跳过', '未执行', 'not executed', 'not run', 'wip',
                     '进行中', 'in progress', 'in_progress']
    for kw in blocked_words:
        if kw in t:
            return 'blocked'

    # 通过类
    pass_words = ['pass', 'passed', '通过', '合格', '达标', 'yes', 'y', 'ok',
                  'success', '√', '✓', 'p', 'done', 'complete', 'completed',
                  'closed', 'resolved', '完成', '已关闭', '已解决']
    if t in pass_words or t.startswith('pass') or '通过' in t or '合格' in t or '达标' in t:
        return 'pass'

    # 失败类（精确匹配兜底）
    fail_exact = ['no', 'n', '×', '✗', 'f', 'bug']
    if t in fail_exact:
        return 'fail'

    return 'unknown'


def _is_pass_result(text):
    return _classify_result(text) == 'pass'


def _is_fail_result(text):
    cls = _classify_result(text)
    return cls in ('fail', 'blocked', 'delayed')


# === 智能分析引擎 ===

# 测试项分类关键词映射
_CATEGORY_KEYWORDS = {
    '功能测试': ['functional', '功能', 'cuj', 'oobe', 'ota', 'requirement', '需求', '用例', 'case', 'p0'],
    '性能专项': ['ttid', 'ttfd', 'fps', 'latency', '延迟', '响应', 'performance', '性能', '帧率',
                 '启动', 'launch', 'scroll', '滚动', 'time to', '功耗'],
    '续航DOU': ['dou', '续航', 'battery', '电池', '功耗', 'power', '耗电', '待机', 'endurance'],
    '稳定性MTTF': ['mttf', '稳定性', 'stability', 'crash', '死机', '重启', 'reboot', '挂机',
                   '内存', 'memory', 'leak', '泄漏'],
    '运动健康算法': ['心率', 'heart rate', '睡眠', 'sleep', '步数', 'step', 'gps', '轨迹',
                    '游泳', 'swim', '血氧', 'spo2', 'spO2', '血氧', '训练', 'training',
                    '运动', 'sport', '健康', 'health', '算法', 'algorithm', '配速', 'pace',
                    '划次', 'stroke', '能量', 'energy', '压力', 'stress', '指南针', 'compass'],
    '兼容性': ['compatib', '兼容', 'bluetooth', '蓝牙', 'wifi', 'pairing', '配对',
              '连接', 'connect', 'interop'],
    '音频': ['audio', '音质', '音量', '麦克风', 'mic', 'speaker', '扬声器', '降噪', 'anc'],
    'UI/UX': ['ui', 'ux', '界面', '交互', '动画', 'animation', '表盘', 'watch face',
              'theme', '主题', '壁纸', 'widget', '小组件'],
    '安全': ['security', '安全', '加密', 'encrypt', 'privacy', '隐私', '认证', 'auth'],
    '网络通信': ['network', '网络', 'sync', '同步', 'push', '通知', 'notification',
                'cellular', '蜂窝', ' esim', 'esim'],
}

# 状态中文标签
_STATUS_LABELS = {
    'pass': '通过',
    'fail': '不通过',
    'blocked': '阻塞',
    'delayed': '已延期',
    'n_a': '不适用',
    'unknown': '未识别',
}


def _detect_category(item_name, item_module=''):
    """根据测试项名称和模块识别所属分类"""
    text = f"{item_name} {item_module}".lower()
    for category, keywords in _CATEGORY_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in text:
                return category
    return '其他'


def _assess_risk_level(items):
    """
    根据一组测试项的结果评估风险等级。
    返回: '高' / '中' / '低' / '无'
    """
    if not items:
        return '无'

    total = len(items)
    fail_count = sum(1 for i in items if i['result'] in ('fail',))
    delayed_count = sum(1 for i in items if i['result'] == 'delayed')
    blocked_count = sum(1 for i in items if i['result'] == 'blocked')
    pass_count = sum(1 for i in items if i['result'] == 'pass')

    problem_count = fail_count + delayed_count + blocked_count
    problem_ratio = problem_count / total if total > 0 else 0
    fail_ratio = fail_count / total if total > 0 else 0

    # 高风险：有失败项且占比 >= 30%，或全部延期/阻塞
    if fail_ratio >= 0.3 or (problem_count == total and total > 0):
        return '高'
    # 中风险：有失败项但占比 < 30%，或有延期/阻塞项
    if fail_count > 0 or delayed_count > 0 or blocked_count > 0:
        return '中'
    # 低风险：全部通过但有未知项
    if pass_count < total:
        return '低'
    return '无'


def _generate_intelligent_analysis(test_items, project_info, stats):
    """
    生成结构化智能分析报告，对标豆包的分析质量。
    输出包含：执行摘要、分类汇总、风险等级、关键发现、改进建议。
    """
    if not test_items:
        return {
            'executive_summary': '未找到测试项数据，请检查Excel文件格式是否正确。',
            'overall_risk': '无',
            'sections': [],
            'key_findings': [],
            'recommendations': []
        }

    total = stats['total']
    pass_count = stats['pass']
    fail_count = stats['fail']
    blocked_count = stats['blocked']
    delayed_count = stats['delayed']
    unknown_count = stats['unknown']
    executed = stats['executed']
    pass_rate = stats['pass_rate']
    executed_pass_rate = stats['executed_pass_rate']

    # === 1. 按分类分组 ===
    category_groups = {}
    for item in test_items:
        cat = _detect_category(item['name'], item.get('module', ''))
        if cat not in category_groups:
            category_groups[cat] = []
        category_groups[cat].append(item)

    # === 2. 生成各分类的汇总 ===
    sections = []
    for cat, items in category_groups.items():
        cat_total = len(items)
        cat_pass = sum(1 for i in items if i['result'] == 'pass')
        cat_fail = sum(1 for i in items if i['result'] == 'fail')
        cat_blocked = sum(1 for i in items if i['result'] == 'blocked')
        cat_delayed = sum(1 for i in items if i['result'] == 'delayed')
        cat_unknown = sum(1 for i in items if i['result'] == 'unknown')
        cat_risk = _assess_risk_level(items)

        # 生成分类摘要文本
        summary_parts = []
        summary_parts.append(f"共 {cat_total} 项")
        if cat_pass > 0:
            summary_parts.append(f"通过 {cat_pass} 项")
        if cat_fail > 0:
            summary_parts.append(f"不通过 {cat_fail} 项")
        if cat_delayed > 0:
            summary_parts.append(f"延期 {cat_delayed} 项")
        if cat_blocked > 0:
            summary_parts.append(f"阻塞 {cat_blocked} 项")
        if cat_unknown > 0:
            summary_parts.append(f"未识别 {cat_unknown} 项")

        cat_pass_rate = f"{(cat_pass / cat_total * 100):.1f}%" if cat_total > 0 else "0%"
        summary_parts.append(f"通过率 {cat_pass_rate}")

        # 提取该分类下的问题项详情
        problem_items = [i for i in items if i['result'] in ('fail', 'blocked', 'delayed')]

        sections.append({
            'category': cat,
            'risk_level': cat_risk,
            'total': cat_total,
            'pass': cat_pass,
            'fail': cat_fail,
            'blocked': cat_blocked,
            'delayed': cat_delayed,
            'pass_rate': cat_pass_rate,
            'summary': '，'.join(summary_parts),
            'problem_items': [{
                'name': i['name'],
                'result': i['result'],
                'result_text': i['result_text'],
                'reason': i.get('reason', ''),
                'target': i.get('target', ''),
                'actual': i.get('actual', '')
            } for i in problem_items],
            'items': items
        })

    # 按风险等级排序：高 > 中 > 低 > 无
    risk_order = {'高': 0, '中': 1, '低': 2, '无': 3}
    sections.sort(key=lambda s: risk_order.get(s['risk_level'], 4))

    # === 3. 评估整体风险 ===
    all_risks = [s['risk_level'] for s in sections]
    if '高' in all_risks:
        overall_risk = '高'
    elif '中' in all_risks:
        overall_risk = '中'
    elif '低' in all_risks:
        overall_risk = '低'
    else:
        overall_risk = '无'

    # === 4. 生成执行摘要 ===
    summary_lines = []

    # 基本信息
    summary_lines.append(f"本次共分析 {total} 项测试")

    # 执行情况
    if delayed_count > 0 or blocked_count > 0:
        summary_lines.append(f"其中已执行 {executed} 项，延期 {delayed_count} 项，阻塞 {blocked_count} 项")
    else:
        summary_lines.append(f"已执行 {executed} 项")

    # 通过/失败情况
    result_parts = []
    if pass_count > 0:
        result_parts.append(f"通过 {pass_count} 项")
    if fail_count > 0:
        result_parts.append(f"不通过 {fail_count} 项")
    if blocked_count > 0:
        result_parts.append(f"阻塞 {blocked_count} 项")
    if delayed_count > 0:
        result_parts.append(f"延期 {delayed_count} 项")
    if unknown_count > 0:
        result_parts.append(f"未识别 {unknown_count} 项")
    summary_lines.append('，'.join(result_parts))

    # 通过率
    if executed > 0 and executed < total:
        summary_lines.append(f"整体通过率 {pass_rate}（已执行项通过率 {executed_pass_rate}）")
    else:
        summary_lines.append(f"整体通过率 {pass_rate}")

    # 整体评估
    if fail_count > 0:
        high_risk_sections = [s for s in sections if s['risk_level'] == '高']
        mid_risk_sections = [s for s in sections if s['risk_level'] == '中']
        if high_risk_sections:
            risk_names = '、'.join([s['category'] for s in high_risk_sections])
            summary_lines.append(f"整体风险等级为「高」，{risk_names} 存在高风险项，需重点关注")
        elif mid_risk_sections:
            summary_lines.append(f"整体风险等级为「中」，部分测试项未通过，需跟进闭环")
        else:
            summary_lines.append(f"整体风险等级为「低」，存在少量未通过项")
    elif delayed_count > 0 or blocked_count > 0:
        summary_lines.append(f"整体风险等级为「{'中' if (delayed_count + blocked_count) > total * 0.3 else '低'}」，"
                             f"部分测试项延期或阻塞，需推进执行")
    elif pass_count == total:
        summary_lines.append("所有测试项均已通过，整体风险等级为「无」")
    else:
        summary_lines.append("部分测试项状态未明确，建议核实数据完整性")

    executive_summary = '。'.join(summary_lines) + '。'

    # === 5. 关键发现 ===
    key_findings = []

    # 高风险分类
    for s in sections:
        if s['risk_level'] == '高':
            finding = f"【高风险】{s['category']}：{s['summary']}"
            if s['problem_items']:
                fail_names = [pi['name'] for pi in s['problem_items'][:5]]
                finding += f"。主要问题项：{', '.join(fail_names)}"
            key_findings.append(finding)

    # 中风险分类
    for s in sections:
        if s['risk_level'] == '中':
            finding = f"【中风险】{s['category']}：{s['summary']}"
            if s['problem_items']:
                fail_names = [pi['name'] for pi in s['problem_items'][:3]]
                finding += f"。关注项：{', '.join(fail_names)}"
            key_findings.append(finding)

    # 延期项汇总
    if delayed_count > 0:
        delayed_items = [i for i in test_items if i['result'] == 'delayed']
        delayed_names = [i['name'] for i in delayed_items[:5]]
        more = f" 等 {len(delayed_items)} 项" if len(delayed_items) > 5 else ""
        key_findings.append(f"【延期项】{len(delayed_items)} 项测试已延期：{', '.join(delayed_names)}{more}，需尽快安排执行")

    # 阻塞项汇总
    if blocked_count > 0:
        blocked_items = [i for i in test_items if i['result'] == 'blocked']
        blocked_names = [i['name'] for i in blocked_items[:5]]
        more = f" 等 {len(blocked_items)} 项" if len(blocked_items) > 5 else ""
        key_findings.append(f"【阻塞项】{len(blocked_items)} 项测试受阻：{', '.join(blocked_names)}{more}，需解除阻塞依赖")

    # 通过项亮点
    pass_sections = [s for s in sections if s['risk_level'] == '无' and s['pass'] == s['total']]
    if pass_sections:
        pass_names = '、'.join([s['category'] for s in pass_sections])
        key_findings.append(f"【达标项】{pass_names} 全部通过，无风险")

    # === 6. 改进建议 ===
    recommendations = []

    if fail_count > 0:
        recommendations.append(f"针对 {fail_count} 项不通过测试，建议优先修复 Fail 项并安排回归验证")
    if delayed_count > 0:
        recommendations.append(f"针对 {delayed_count} 项延期测试，建议明确执行时间节点并推进落地")
    if blocked_count > 0:
        recommendations.append(f"针对 {blocked_count} 项阻塞测试，建议排查阻塞根因并协调资源解除依赖")
    if unknown_count > 0:
        recommendations.append(f"针对 {unknown_count} 项状态未识别的测试，建议核实结果字段填写是否规范")

    high_risk_sections = [s for s in sections if s['risk_level'] == '高']
    if high_risk_sections:
        rec_cats = '、'.join([s['category'] for s in high_risk_sections])
        recommendations.append(f"「{rec_cats}」为高风险领域，建议作为下一阶段重点攻关方向")

    if not recommendations:
        recommendations.append("所有测试项均已通过，建议保持现有质量水平，持续监控")

    return {
        'executive_summary': executive_summary,
        'overall_risk': overall_risk,
        'sections': sections,
        'key_findings': key_findings,
        'recommendations': recommendations,
        'status_labels': _STATUS_LABELS
    }


# === 分块上传 API（绕过预览代理的请求体大小限制）===
# v2.0: 使用 SQLite 持久化存储分块上传元数据（替代 JSON 文件 + fcntl 文件锁）

_chunk_uploads_dir = os.path.join(app.config['UPLOAD_FOLDER'], '_chunk_meta')
os.makedirs(_chunk_uploads_dir, exist_ok=True)

# v2.0: 分块上传元数据使用 SQLite，替代 JSON 文件 + 文件锁
def _load_chunk_meta(upload_id):
    """从数据库加载分块上传元数据"""
    session = db.get_upload_session(upload_id)
    if not session:
        return None
    filename = session['filename']
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ('.xlsx', '.xls', '.csv'):
        ext = '.xlsx'
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"excel_{upload_id}{ext}")
    return {
        'upload_id': upload_id,
        'filename': filename,
        'file_path': file_path,
        'ext': ext,
        'total_chunks': session['total_chunks'],
        'chunk_size': session['chunk_size'],
        'file_size': session['file_size'],
        'total_size': session['file_size'],
        'received_chunks': session.get('received_set', set()),
    }

def _save_chunk_meta(upload_id, meta):
    """创建/更新分块上传会话到数据库"""
    db.create_upload_session(
        upload_id,
        meta.get('filename', ''),
        meta.get('total_chunks', 0),
        meta.get('chunk_size', 2 * 1024 * 1024),
        meta.get('total_size', meta.get('file_size', 0))
    )

def _add_received_chunk(upload_id, chunk_index):
    """线程安全地添加已接收分块（SQLite 事务保证原子性）"""
    return db.add_received_chunk(upload_id, chunk_index)

def _delete_chunk_meta(upload_id):
    """删除分块上传会话"""
    db.delete_upload_session(upload_id)

@app.route('/api/upload-init', methods=['POST'])
def api_upload_init():
    """初始化分块上传，返回 upload_id"""
    # 自行检查认证（因为此端点在 PUBLIC_PATHS 中）
    user = auth.get_current_user()
    if not user:
        return jsonify({'error': '请先登录', 'need_login': True}), 401

    data = request.get_json(silent=True)
    if data is None:
        data = {}

    filename = data.get('filename', '')
    total_size = data.get('total_size', 0)
    total_chunks = data.get('total_chunks', 0)

    if not filename or total_chunks == 0:
        return jsonify({'error': '缺少必要参数: filename, total_chunks'}), 400

    filename_lower = filename.lower()
    if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
        return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400

    orig_ext = os.path.splitext(filename)[1].lower()
    if orig_ext not in ('.xlsx', '.xls', '.csv'):
        orig_ext = '.xlsx'

    upload_id = secrets.token_hex(8)  # 16位不可预测的随机ID
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"excel_{upload_id}{orig_ext}")

    _chunk_uploads = {
        'filename': filename,
        'file_path': file_path,
        'ext': orig_ext,
        'total_chunks': total_chunks,
        'total_size': total_size,
        'received_chunks': set(),
        'created_at': time.time()
    }
    _save_chunk_meta(upload_id, _chunk_uploads)

    # 预分配文件大小（支持并发分块按 offset 写入）
    with open(file_path, 'wb') as f:
        if total_size > 0:
            f.truncate(total_size)

    logger.info(f"分块上传初始化: {filename}, upload_id={upload_id}, total_chunks={total_chunks}, total_size={total_size}")

    return jsonify({
        'status': 'success',
        'data': {
            'upload_id': upload_id
        }
    })

@app.route('/api/upload-chunk', methods=['POST'])
def api_upload_chunk():
    """上传单个分块"""
    # 自行检查认证
    user = auth.get_current_user()
    if not user:
        return jsonify({'error': '请先登录', 'need_login': True}), 401

    upload_id = request.form.get('upload_id', '')
    chunk_index = request.form.get('chunk_index', '')
    chunk_file = request.files.get('chunk', None)

    if not upload_id:
        return jsonify({'error': '无效的 upload_id'}), 400

    meta = _load_chunk_meta(upload_id)
    if meta is None:
        return jsonify({'error': '无效的 upload_id'}), 400

    if chunk_index == '' or chunk_file is None:
        return jsonify({'error': '缺少 chunk_index 或 chunk'}), 400

    chunk_index = int(chunk_index)

    if chunk_index in meta['received_chunks']:
        return jsonify({'status': 'success', 'data': {'chunk_index': chunk_index, 'duplicate': True, 'received': len(meta['received_chunks']), 'total': meta['total_chunks']}})

    # 读取分块数据，按 offset 写入正确位置（支持并发乱序上传）
    chunk_data = chunk_file.read()
    offset = int(request.form.get('offset', -1))
    if offset < 0:
        return jsonify({'error': '缺少 offset 参数'}), 400

    with open(meta['file_path'], 'r+b') as f:
        f.seek(offset)
        f.write(chunk_data)

    # 原子操作：加锁更新 received_chunks（防止并发覆盖导致丢块）
    updated_meta = _add_received_chunk(upload_id, chunk_index)
    if updated_meta is None:
        return jsonify({'error': '更新分块元数据失败'}), 500

    logger.info(f"分块上传: upload_id={upload_id}, chunk={chunk_index}/{updated_meta['total_chunks'] - 1}, size={len(chunk_data)}, received={len(updated_meta['received_chunks'])}")

    return jsonify({
        'status': 'success',
        'data': {
            'chunk_index': chunk_index,
            'received': len(updated_meta['received_chunks']),
            'total': updated_meta['total_chunks']
        }
    })

@app.route('/api/upload-complete', methods=['POST'])
def api_upload_complete():
    """分块上传完成，验证文件并返回 file_id"""
    # 自行检查认证
    user = auth.get_current_user()
    if not user:
        return jsonify({'error': '请先登录', 'need_login': True}), 401

    data = request.get_json(silent=True)
    if data is None:
        data = {}

    upload_id = data.get('upload_id', '')
    if not upload_id:
        return jsonify({'error': '无效的 upload_id'}), 400

    meta = _load_chunk_meta(upload_id)
    if meta is None:
        return jsonify({'error': '无效的 upload_id'}), 400

    if len(meta['received_chunks']) != meta['total_chunks']:
        # 返回缺失的分块索引，前端可自动补传
        missing = sorted(set(range(meta['total_chunks'])) - meta['received_chunks'])
        return jsonify({
            'error': f'分块不完整: 已收到 {len(meta["received_chunks"])}/{meta["total_chunks"]}',
            'missing_chunks': missing,
            'total_chunks': meta['total_chunks'],
            'received_chunks': len(meta['received_chunks'])
        }), 400

    file_path = meta['file_path']
    filename = meta['filename']

    # 从文件路径提取 file_id
    file_id = os.path.basename(file_path).replace('excel_', '').replace(meta['ext'], '')

    try:
        file_size = os.path.getsize(file_path)
        logger.info(f"========== 分块上传完成: {filename} (size={file_size / 1024 / 1024:.2f}MB) ==========")

        # 检测文件格式
        is_html = False
        with open(file_path, 'rb') as f:
            header = f.read(200)
        if b'<html' in header.lower() or b'<!doctype' in header.lower():
            is_html = True
            logger.info(f"文件格式检测: HTML 伪装的 .xls 文件")
        elif header[:4] == b'\x50\x4b\x03\x04':
            logger.info(f"文件格式检测: .xlsx (ZIP格式)")
        elif header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            logger.info(f"文件格式检测: 真正的 .xls (OLE2格式)")
        else:
            logger.info(f"文件格式检测: 未知格式, header={header[:20]}")

        reader = ExcelReader(file_path)
        reader.open()
        sheet_names = reader.get_sheet_names()
        reader.close()

        # 清理元数据
        _delete_chunk_meta(upload_id)

        logger.info(f"上传成功: file_id={file_id}, sheets={sheet_names}, is_html={is_html}")

        return jsonify({
            'status': 'success',
            'data': {
                'file_id': file_id,
                'file_name': filename,
                'sheet_names': sheet_names,
                'file_size_mb': round(file_size / 1024 / 1024, 2),
                'is_html': is_html,
            }
        })

    except Exception as e:
        logger.error(f"分块上传文件分析失败: {traceback.format_exc()}")
        # 清理元数据和文件
        try:
            os.unlink(file_path)
        except Exception:
            pass
        _delete_chunk_meta(upload_id)
        return jsonify({'error': str(e)}), 500


# === Excel问题分析 API ===
@app.route('/api/excel-analyze', methods=['POST'])
def api_excel_analyze():
    if 'file' not in request.files:
        return jsonify({'error': '请选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '请选择文件'}), 400

    filename_lower = file.filename.lower()
    if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
        return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400

    # 检查文件大小（云平台内存有限）
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    if file_size > 200 * 1024 * 1024:
        return jsonify({'error': f'文件过大({file_size // 1024 // 1024}MB)，云端最大支持200MB。超大文件建议使用分片上传或本地部署'}), 413

    orig_ext = os.path.splitext(file.filename)[1].lower()
    if orig_ext not in ('.xlsx', '.xls', '.csv'):
        orig_ext = '.xlsx'

    try:
        file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"excel_{file_id}{orig_ext}")
        file.save(file_path)

        logger.info(f"========== 收到问题分析文件: {file.filename} ==========")

        reader = ExcelReader(file_path)
        reader.open()
        sheet_names = reader.get_sheet_names()
        reader.close()

        return jsonify({
            'status': 'success',
            'data': {
                'file_id': file_id,
                'file_name': file.filename,
                'sheet_names': sheet_names
            }
        })

    except Exception as e:
        logger.error(f"文件上传失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# 缓存完整分析结果（避免重复计算）
_analysis_cache = {}

# 后台任务存储
_background_tasks = {}
_background_tasks_dir = os.path.join(app.config['UPLOAD_FOLDER'], '_task_meta')
os.makedirs(_background_tasks_dir, exist_ok=True)

# v2.0: 后台任务元数据使用 SQLite，替代 JSON 文件
def _load_task_meta(task_id):
    """从数据库加载任务元数据"""
    task = db.get_task(task_id)
    if not task:
        return None
    return {
        'task_id': task_id,
        'task_type': task.get('task_type', ''),
        'status': task.get('status', 'pending'),
        'progress': task.get('progress', 0),
        'result': task.get('result'),
        'error': task.get('error'),
    }

def _save_task_meta(task_id, task_data):
    """持久化任务元数据到数据库"""
    status = task_data.get('status', 'pending')
    progress = task_data.get('progress', 0)
    result = task_data.get('result')
    error = task_data.get('error')

    # 检查任务是否已存在
    existing = db.get_task(task_id)
    if existing:
        db.update_task(task_id, status=status, progress=progress, result=result, error=error)
    else:
        db.create_task(task_id, task_data.get('task_type', 'unknown'))
        db.update_task(task_id, status=status, progress=progress, result=result, error=error)

def _delete_task_meta(task_id):
    """删除任务元数据"""
    db.delete_task(task_id)

import threading
from excel_analyzers import _analyze_issue_sheet, _analyze_issue_sheet_fast, _detect_issue_columns, _is_valid_severity_value, _log_mem, _match_severity_level, _safe_get

@app.route('/api/excel-analyze-fields', methods=['POST'])
def api_excel_analyze_fields():
    """轻量级字段映射接口：同步执行（仅需读取表头，耗时约5秒）"""
    data = request.get_json(silent=True)
    if data is None:
        data = {}

    file_id = data.get('file_id', '')
    sheet_name = data.get('sheet_name', '')

    logger.info(f"excel-analyze-fields 请求: file_id={file_id}, sheet_name={sheet_name}")

    if not file_id or not sheet_name:
        return jsonify({'error': '缺少参数: file_id, sheet_name'}), 400
    if not _validate_file_id(file_id):
        return jsonify({'error': '无效的文件ID'}), 400

    file_path = None
    for ext in ['.xlsx', '.xls', '.csv']:
        candidate = os.path.join(app.config['UPLOAD_FOLDER'], f"excel_{file_id}{ext}")
        if os.path.exists(candidate):
            file_path = candidate
            break

    if not file_path:
        return jsonify({'error': f'文件不存在: {file_id}'}), 404

    try:
        reader = ExcelReader(file_path)
        reader.open()

        # 使用轻量级表头读取方法，避免大文件 OOM
        if reader._is_html:
            headers, first_data_row, data_row_count = reader._parse_html_headers_only()
            reader.close()
            
            if not headers:
                return jsonify({
                    'status': 'done',
                    'data': {
                        'headers': [],
                        'detected_columns': {},
                        'detected_fields_count': 0,
                        'current_sheet': sheet_name,
                        'summary': {'total_issues': 0},
                        'sample_data': [],
                    }
                })

            headers = [str(c).strip() if c else '' for c in headers]
            col_map = _detect_issue_columns(headers)

            raw_detected = {
                'issue_id': col_map.get('id', -1),
                'title': col_map.get('title', -1),
                'module': col_map.get('module', -1),
                'severity': col_map.get('severity', -1),
                'status': col_map.get('status', -1),
                'developer': col_map.get('developer', -1),
                'create_date': col_map.get('created_date', -1),
                'resolve_date': col_map.get('resolved_date', -1),
                'fixed_date': col_map.get('closed_date', -1),
                'fixed_version': col_map.get('fix_version', -1),
            }
            detected_columns = {k: v for k, v in raw_detected.items() if v >= 0}

            sample_data = [first_data_row] if first_data_row else []
            total_issues = data_row_count

            logger.info(f"字段映射完成(轻量级): detected_columns={detected_columns}, total_issues≈{total_issues}")

            result = {
                'headers': headers,
                'detected_columns': detected_columns,
                'detected_fields_count': len(detected_columns),
                'current_sheet': sheet_name,
                'summary': {'total_issues': total_issues},
                'sample_data': sample_data,
            }

            gc.collect()
            return jsonify({'status': 'done', 'data': result})
        else:
            # 非 HTML 格式：正常读取（内存安全）
            rows = reader.get_sheet_data(sheet_name)
            reader.close()

            if not rows or len(rows) < 1:
                return jsonify({
                    'status': 'done',
                    'data': {
                        'headers': [],
                        'detected_columns': {},
                        'detected_fields_count': 0,
                        'current_sheet': sheet_name,
                        'summary': {'total_issues': 0},
                        'sample_data': [],
                    }
                })

            headers = [str(c).strip() if c else '' for c in rows[0]]
            col_map = _detect_issue_columns(headers)

            raw_detected = {
                'issue_id': col_map.get('id', -1),
                'title': col_map.get('title', -1),
                'module': col_map.get('module', -1),
                'severity': col_map.get('severity', -1),
                'status': col_map.get('status', -1),
                'developer': col_map.get('developer', -1),
                'create_date': col_map.get('created_date', -1),
                'resolve_date': col_map.get('resolved_date', -1),
                'fixed_date': col_map.get('closed_date', -1),
                'fixed_version': col_map.get('fix_version', -1),
            }
            detected_columns = {k: v for k, v in raw_detected.items() if v >= 0}

            data_rows = rows[1:]
            total_issues = sum(1 for row in data_rows if any(str(c).strip() for c in row))
            sample_data = data_rows[:3] if data_rows else []

            logger.info(f"字段映射完成: detected_columns={detected_columns}, total_issues={total_issues}")

            result = {
                'headers': headers,
                'detected_columns': detected_columns,
                'detected_fields_count': len(detected_columns),
                'current_sheet': sheet_name,
                'summary': {'total_issues': total_issues},
                'sample_data': sample_data,
            }

            del rows, data_rows
            gc.collect()

            return jsonify({'status': 'done', 'data': result})

    except Exception as e:
        logger.error(f"字段映射失败: {traceback.format_exc()}")
        return jsonify({'status': 'error', 'error': str(e)}), 500


@app.route('/api/task-status', methods=['POST'])
def api_task_status():
    """查询后台任务状态"""
    data = request.get_json(silent=True)
    if data is None:
        data = {}

    task_id = data.get('task_id', '')
    if not task_id:
        return jsonify({'error': '无效的 task_id'}), 400

    # 优先从内存读取，worker 重启后从磁盘恢复
    task = _background_tasks.get(task_id)
    if task is None:
        task = _load_task_meta(task_id)
        if task is None:
            return jsonify({'status': 'error', 'error': '任务不存在或已过期，请重新上传文件'}), 400

    # 检查任务是否超时（超过 15 分钟仍在 processing 视为超时，大文件需要更长时间）
    if task['status'] == 'processing':
        created_at = task.get('created_at', 0)
        if created_at and (time.time() - created_at) > 900:
            task['status'] = 'error'
            task['error'] = '分析超时（超过15分钟），可能是文件过大或格式异常，请尝试减少数据量或转换为 .xlsx 格式'
            _save_task_meta(task_id, task)

    resp = {'status': task['status']}

    if task['status'] == 'processing':
        resp['progress'] = task.get('progress', 0)
        resp['progress_msg'] = task.get('progress_msg', '正在分析...')
    elif task['status'] == 'done':
        resp['data'] = task['result']
    elif task['status'] == 'error':
        resp['error'] = task['error']

    return jsonify(resp)


@app.route('/api/excel-analyze-sheet', methods=['POST'])
def api_excel_analyze_sheet():
    """完整分析接口：启动后台分析，立即返回 task_id"""
    data = request.get_json(silent=True)
    if data is None:
        data = {}

    file_id = data.get('file_id', '')
    sheet_name = data.get('sheet_name', '')

    logger.info(f"excel-analyze-sheet 请求: file_id={file_id}, sheet_name={sheet_name}, content_type={request.content_type}")

    if not file_id or not sheet_name:
        return jsonify({'error': f'缺少参数: file_id={repr(file_id)}, sheet_name={repr(sheet_name)}'}), 400
    if not _validate_file_id(file_id):
        return jsonify({'error': '无效的文件ID'}), 400

    file_path = None
    for ext in ['.xlsx', '.xls', '.csv']:
        candidate = os.path.join(app.config['UPLOAD_FOLDER'], f"excel_{file_id}{ext}")
        if os.path.exists(candidate):
            file_path = candidate
            break

    if not file_path:
        return jsonify({'error': f'文件不存在: {file_id}'}), 404

    # 创建后台任务
    task_id = hashlib.md5(f"sheet_{file_id}_{sheet_name}_{time.time()}".encode()).hexdigest()[:16]
    task_data = {
        'status': 'processing',
        'result': None,
        'error': None,
        'created_at': time.time()
    }
    _background_tasks[task_id] = task_data
    _save_task_meta(task_id, task_data)

    def _do_full_analysis():
        try:
            gc.collect()
            _log_mem("开始分析前")
            t0 = time.time()

            # 判断文件大小，>10MB 使用 pandas 高速分析
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            use_fast = file_size > 10 * 1024 * 1024
            logger.info(f"分析模式: {'fast(pandas)' if use_fast else 'standard(openpyxl)'}, 文件大小: {file_size/1024/1024:.1f}MB")

            if use_fast:
                def progress_cb(percent, message):
                    _background_tasks[task_id]['progress'] = percent
                    _background_tasks[task_id]['progress_msg'] = message
                    _save_task_meta(task_id, _background_tasks[task_id])
                result = _analyze_issue_sheet_fast(file_path, sheet_name, progress_cb=progress_cb)
            else:
                result = _analyze_issue_sheet(file_path, sheet_name)

            elapsed = time.time() - t0
            _log_mem(f"分析完成，耗时 {elapsed:.1f}s")
            gc.collect()

            _background_tasks[task_id]['result'] = result
            _background_tasks[task_id]['status'] = 'done'
            _background_tasks[task_id]['progress'] = 100
            _save_task_meta(task_id, _background_tasks[task_id])
        except Exception as e:
            error_detail = str(e) if str(e) else f'{type(e).__name__} (无详细错误信息)'
            logger.error(f"分析失败: {traceback.format_exc()}")
            _background_tasks[task_id]['error'] = error_detail
            _background_tasks[task_id]['status'] = 'error'
            _save_task_meta(task_id, _background_tasks[task_id])

    thread = threading.Thread(target=_do_full_analysis, daemon=True)
    thread.start()

    return jsonify({'status': 'success', 'data': {'task_id': task_id}})

_pdf_render_lock = threading.Lock()
_pw_instance = None
_pw_browser = None

def _get_pw_browser():
    """获取或创建全局 Chromium 浏览器实例（复用，避免每次冷启动）"""
    global _pw_instance, _pw_browser
    if _pw_browser is not None:
        try:
            _ = _pw_browser.version
            return _pw_browser
        except Exception:
            _pw_browser = None
            logger.warning("Chromium 浏览器已断开，正在重新创建...")
    if _pw_instance is None:
        _pw_instance = _sync_playwright().start()
    # 优先使用系统 Chrome（本地环境），失败后回退到 Playwright 内置 Chromium（Docker/Railway）
    try:
        _pw_browser = _pw_instance.chromium.launch(headless=True, channel="chrome")
    except Exception:
        _pw_browser = _pw_instance.chromium.launch(headless=True)
    logger.info("Chromium 浏览器实例已创建（全局复用）")
    return _pw_browser

def _render_pdf(html_path, pdf_path, margin=None, extra_wait_ms=0, wait_selector=None):
    """使用全局 Chromium 实例渲染 PDF（线程安全）。

    通过 _pdf_render_lock 串行化访问，避免 Playwright sync API 的线程安全问题。
    每次创建独立的 browser context，渲染完毕后关闭；浏览器进程本身保持常驻。
    """
    if margin is None:
        margin = {'top': '2cm', 'right': '2cm', 'bottom': '2cm', 'left': '2cm'}
    with _pdf_render_lock:
        browser = _get_pw_browser()
        context = browser.new_context()
        page = context.new_page()
        try:
            page.goto(f'file://{html_path}')
            page.wait_for_load_state('networkidle')
            if wait_selector:
                try:
                    page.wait_for_selector(wait_selector, timeout=5000)
                except Exception:
                    pass
            if extra_wait_ms > 0:
                page.wait_for_timeout(extra_wait_ms)
            page.pdf(
                path=pdf_path,
                format='A4',
                margin=margin,
                print_background=True
            )
        finally:
            context.close()

# === Markdown转PDF API ===
@app.route('/api/md2pdf', methods=['POST'])
def api_md2pdf():
    data = request.json or {}
    markdown_content = data.get('content', '')
    if not markdown_content:
        return jsonify({'error': '内容不能为空'}), 400

    try:
        import markdown
        html_content = markdown.markdown(
            markdown_content,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )

        import tempfile as tf

        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: -apple-system, "Segoe UI", Roboto, sans-serif; padding: 40px; line-height: 1.8; }}
        h1, h2, h3 {{ margin-top: 1.5em; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; }}
        th {{ background: #f5f5f7; }}
        pre {{ background: #f5f5f7; padding: 16px; border-radius: 8px; overflow-x: auto; }}
        code {{ font-family: "SF Mono", Monaco, Consolas, monospace; }}
        blockquote {{ border-left: 4px solid #0071e3; margin: 1em 0; padding: 0.5em 1em; background: #f9f9f9; }}
    </style>
</head>
<body>{html_content}</body>
</html>
''')
            html_path = f.name

        pdf_path = os.path.join(app.config['PDF_FOLDER'], f"md2pdf_{int(time.time())}.pdf")

        _render_pdf(html_path, pdf_path)

        return jsonify({'filename': os.path.basename(pdf_path)})
    except Exception as e:
        logger.error(f"PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# === PDF快转 - 预览API ===
@app.route('/preview', methods=['POST'])
def api_preview():
    data = request.json or {}
    markdown_content = data.get('content', '')
    watermark = data.get('watermark', '')

    if not markdown_content:
        return jsonify({'html': ''})

    try:
        import markdown
        html_content = markdown.markdown(
            markdown_content,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )
        
        # 如果有水印，添加水印样式
        if watermark:
            html_content += f'''
            <div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">
                {watermark}
            </div>
            '''
        
        return jsonify({'html': html_content})
    except Exception as e:
        logger.error(f"预览失败: {traceback.format_exc()}")
        return jsonify({'html': '', 'error': str(e)}), 500


# === PDF快转 - 转换API ===
@app.route('/convert', methods=['POST'])
def api_convert():
    data = request.json or {}
    markdown_content = data.get('content', '')
    watermark = data.get('watermark', '')
    filename = data.get('filename', '')

    if not markdown_content:
        return jsonify({'error': '内容不能为空'}), 400

    try:
        import markdown
        html_content = markdown.markdown(
            markdown_content,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )

        import tempfile as tf

        # 添加水印
        watermark_html = ''
        if watermark:
            watermark_html = f'''
            <div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">
                {watermark}
            </div>
            '''

        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: -apple-system, "Segoe UI", Roboto, "PingFang SC", "Microsoft YaHei", sans-serif; padding: 40px; line-height: 1.8; }}
        h1, h2, h3 {{ margin-top: 1.5em; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; }}
        th {{ background: #f5f5f7; }}
        pre {{ background: #f5f5f7; padding: 16px; border-radius: 8px; overflow-x: auto; }}
        code {{ font-family: "SF Mono", Monaco, Consolas, monospace; }}
        blockquote {{ border-left: 4px solid #0071e3; margin: 1em 0; padding: 0.5em 1em; background: #f9f9f9; }}
    </style>
</head>
<body>
{html_content}
{watermark_html}
</body>
</html>
''')
            html_path = f.name

        if filename:
            safe_filename = re.sub(r'[^\w\s-]', '', filename).strip() or 'document'
            pdf_filename = f"{safe_filename}_{int(time.time())}.pdf"
        else:
            pdf_filename = f"md2pdf_{int(time.time())}.pdf"

        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

        # 后台渲染 PDF（避免同步请求超过 Railway 5 分钟超时）
        task_id = hashlib.md5(f"convert_{time.time()}".encode()).hexdigest()[:16]
        task_data = {
            'status': 'processing',
            'result': None,
            'error': None,
            'created_at': time.time()
        }
        _background_tasks[task_id] = task_data
        _save_task_meta(task_id, task_data)

        def _do_convert_pdf():
            try:
                _render_pdf(html_path, pdf_path)
                _background_tasks[task_id]['result'] = {'filename': pdf_filename}
                _background_tasks[task_id]['status'] = 'done'
                _save_task_meta(task_id, _background_tasks[task_id])
            except Exception as e:
                error_detail = str(e) if str(e) else f'{type(e).__name__}'
                logger.error(f"PDF转换失败: {traceback.format_exc()}")
                _background_tasks[task_id]['error'] = error_detail
                _background_tasks[task_id]['status'] = 'error'
                _save_task_meta(task_id, _background_tasks[task_id])

        thread = threading.Thread(target=_do_convert_pdf, daemon=True)
        thread.start()

        return jsonify({'status': 'success', 'data': {'task_id': task_id}})
    except Exception as e:
        logger.error(f"PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# === PDF快转 - Word上传转换API ===
@app.route('/upload', methods=['POST'])
def api_upload():
    if 'file' not in request.files:
        return jsonify({'error': '请选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '请选择文件'}), 400

    watermark = request.form.get('watermark', '')
    orig_name = os.path.splitext(file.filename)[0]

    if not file.filename.lower().endswith('.docx'):
        return jsonify({'error': '只支持Word文件(.docx)'}), 400

    try:
        file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"word_{file_id}.docx")
        file.save(file_path)

        # 使用python-docx2pdf或其他方式转换
        from docx2pdf import convert
        pdf_filename = f"{orig_name}_{int(time.time())}.pdf"
        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

        convert(file_path, pdf_path)

        # 如果有水印，添加水印
        if watermark:
            # 对于Word转PDF，水印已在PDF快转页面的前端添加
            pass

        return jsonify({
            'filename': pdf_filename,
            'original_name': orig_name
        })
    except ImportError:
        # 如果没有docx2pdf库，尝试使用其他方法
        try:
            # 使用docx库读取并转换
            from docx import Document

            doc = Document(file_path)
            html_content = ''
            for para in doc.paragraphs:
                html_content += f'<p>{para.text}</p>'
            
            # 简单的表格处理
            for table in doc.tables:
                html_content += '<table border="1">'
                for row in table.rows:
                    html_content += '<tr>'
                    for cell in row.cells:
                        html_content += f'<td>{cell.text}</td>'
                    html_content += '</tr>'
                html_content += '</table>'
            
            watermark_html = ''
            if watermark:
                watermark_html = f'''
                <div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">
                    {watermark}
                </div>
                '''

            import tempfile as tf
            with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: -apple-system, "Segoe UI", Roboto, "PingFang SC", "Microsoft YaHei", sans-serif; padding: 40px; line-height: 1.8; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        td {{ border: 1px solid #ddd; padding: 8px 12px; }}
    </style>
</head>
<body>
{html_content}
{watermark_html}
</body>
</html>
''')
                html_path = f.name

            pdf_filename = f"{orig_name}_{int(time.time())}.pdf"
            pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

            _render_pdf(html_path, pdf_path)

            return jsonify({
                'filename': pdf_filename,
                'original_name': orig_name
            })
        except Exception as e2:
            logger.error(f"Word转PDF失败: {traceback.format_exc()}")
            return jsonify({'error': f'Word转PDF失败: {str(e2)}'}), 500
    except Exception as e:
        logger.error(f"Word上传失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# === Excel 解析路由 (for md2pdf tool) ===
@app.route('/excel-parse', methods=['POST'])
def api_excel_parse():
    if 'file' not in request.files:
        return jsonify({'error': '请选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '请选择文件'}), 400

    filename_lower = file.filename.lower()
    if not (filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls') or filename_lower.endswith('.csv')):
        return jsonify({'error': '只支持Excel/CSV文件(.xlsx, .xls, .csv)'}), 400

    try:
        file_id = hashlib.md5(f"{time.time()}_{file.filename}".encode()).hexdigest()[:16]
        orig_ext = os.path.splitext(file.filename)[1].lower()
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"excel_parse_{file_id}{orig_ext}")
        file.save(file_path)

        reader = ExcelReader(file_path)
        reader.open()
        sheet_names = reader.get_sheet_names()
        
        sheets_data = []
        for sheet_name in sheet_names:
            try:
                rows = reader.get_sheet_data(sheet_name)
                if not rows:
                    sheets_data.append({
                        'name': sheet_name,
                        'row_count': 0,
                        'column_count': 0,
                        'headers': [],
                        'data_preview': [],
                        'rows': [],
                        'summary': '空工作表',
                        'categories': {}
                    })
                    continue
                
                headers = [str(c).strip() if c else '' for c in rows[0]]
                data_rows = rows[1:] if len(rows) > 1 else []
                col_count = len(headers)
                row_count = len(data_rows)
                
                # 生成数据预览（前5行）
                data_preview = []
                for row in data_rows[:5]:
                    cells = [str(c).strip() if c else '' for c in row]
                    # 对齐到headers长度
                    while len(cells) < col_count:
                        cells.append('')
                    data_preview.append(cells[:col_count])
                
                # 生成完整数据（用于PDF生成）
                all_rows = []
                for row in data_rows:
                    cells = [str(c).strip() if c else '' for c in row]
                    while len(cells) < col_count:
                        cells.append('')
                    all_rows.append(cells[:col_count])
                
                # 分析表头并分类
                categories = _categorize_headers(headers)
                
                # 生成摘要
                non_empty_headers = [h for h in headers if h]
                summary_parts = []
                if row_count > 0:
                    summary_parts.append(f'{row_count}行数据')
                if non_empty_headers:
                    summary_parts.append(f'{len(non_empty_headers)}列字段')
                
                # 检测日期列
                date_cols = [h for h in headers if any(kw in h.lower() for kw in ['date', '日期', '时间', 'created', 'updated', 'resolved'])]
                if date_cols:
                    summary_parts.append(f'含时间字段')
                
                # 检测数值列
                numeric_count = 0
                for col_idx in range(min(col_count, 20)):
                    for row in data_rows[:10]:
                        cells = [str(c).strip() if c else '' for c in row]
                        if col_idx < len(cells) and cells[col_idx]:
                            try:
                                float(cells[col_idx].replace(',', ''))
                                numeric_count += 1
                                break
                            except (ValueError, IndexError):
                                pass
                
                if numeric_count > col_count * 0.3:
                    summary_parts.append('数值型数据')
                
                # 检测是否有状态/分类字段
                status_cols = [h for h in headers if any(kw in h.lower() for kw in ['status', '状态', 'type', '类型', 'category', '分类'])]
                if status_cols:
                    summary_parts.append(f'含状态字段')
                
                summary = ' · '.join(summary_parts) if summary_parts else f'{row_count}行 x {col_count}列'
                
                # 保留完整列数，确保headers和rows对齐
                # 如果列数超过100，限制到100列以避免性能问题
                max_cols = min(col_count, 100)
                limited_headers = headers[:max_cols]
                limited_data_preview = [row[:max_cols] for row in data_preview]
                limited_rows = [row[:max_cols] for row in all_rows]
                
                sheets_data.append({
                    'name': sheet_name,
                    'row_count': row_count,
                    'column_count': max_cols,
                    'headers': limited_headers,
                    'data_preview': limited_data_preview,
                    'rows': limited_rows,
                    'summary': summary,
                    'categories': categories
                })
            except Exception as e:
                logger.warning(f"解析sheet {sheet_name} 失败: {e}")
                sheets_data.append({
                    'name': sheet_name,
                    'row_count': 0,
                    'column_count': 0,
                    'headers': [],
                    'data_preview': [],
                    'summary': f'解析失败: {str(e)}',
                    'categories': {}
                })
        
        reader.close()
        
        return jsonify({
            'status': 'success',
            'data': {
                'file_name': file.filename,
                'total_sheets': len(sheet_names),
                'sheets': sheets_data
            }
        })
    except Exception as e:
        logger.error(f"Excel解析失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


def _categorize_headers(headers):
    """将表头分类，返回类别到列名的映射"""
    categories = {}
    
    category_rules = {
        '基本信息': ['id', 'key', '编号', '名称', 'name', 'title', '标题', 'summary', '描述', 'description'],
        '状态信息': ['status', '状态', 'state', 'resolved', '解决', 'closed', '关闭'],
        '时间信息': ['date', '日期', '时间', 'created', 'updated', 'resolved', 'date', 'due'],
        '人员信息': ['assignee', 'developer', '负责人', '研发', 'reporter', '报告人', 'creator', '创建人'],
        '优先级': ['priority', '优先级', 'severity', '严重性', 'criticality'],
        '模块/组件': ['component', 'module', '模块', '组件', 'project', '项目'],
        '版本信息': ['version', '版本', 'fix version', 'affected version'],
        '数值指标': ['count', 'total', '金额', 'cost', 'price', 'rate', '比率', 'percentage'],
    }
    
    for header in headers:
        h_lower = header.lower().strip()
        if not h_lower:
            continue
        for category, keywords in category_rules.items():
            if any(kw in h_lower for kw in keywords):
                if category not in categories:
                    categories[category] = []
                categories[category].append({'name': header})
                break
    
    return categories


# === Excel 智能整理路由 ===
@app.route('/excel-organize', methods=['POST'])
def api_excel_organize():
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': '请求数据格式错误'}), 400
    
    structured_data = data.get('structured_data', {})
    user_request = data.get('user_request', '')
    selected_sheets = data.get('selected_sheets', [])
    
    if not structured_data or not user_request:
        return jsonify({'error': '缺少必要参数'}), 400
    
    try:
        ai_config = get_ai_config()
        sheets = structured_data.get('sheets', [])
        
        # 过滤选中的sheets
        if selected_sheets:
            sheets = [s for s in sheets if s.get('name') in selected_sheets]
        
        if not sheets:
            return jsonify({'error': '没有选中的工作表数据'}), 400
        
        # 准备数据摘要供AI使用
        data_summary = _prepare_excel_summary(structured_data, sheets)
        
        if ai_config.get('enabled'):
            # 使用AI进行智能整理
            organized = _ai_organize_excel(data_summary, user_request, ai_config)
        else:
            # 本地整理
            organized = _local_organize_excel(data_summary, user_request)
        
        return jsonify({
            'status': 'success',
            'data': organized
        })
    except Exception as e:
        logger.error(f"Excel整理失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


def _prepare_excel_summary(structured_data, sheets):
    """准备Excel数据摘要"""
    summary = {
        'file_name': structured_data.get('file_name', ''),
        'sheets': []
    }
    
    for sheet in sheets:
        sheet_info = {
            'name': sheet.get('name', ''),
            'row_count': sheet.get('row_count', 0),
            'column_count': sheet.get('column_count', 0),
            'headers': sheet.get('headers', []),
            'data_preview': sheet.get('data_preview', [])[:10]
        }
        summary['sheets'].append(sheet_info)
    
    return summary


def _ai_organize_excel(data_summary, user_request, ai_config):
    """使用AI整理Excel数据"""
    try:
        import requests as req
        
        sheets_text = ''
        for s in data_summary.get('sheets', []):
            headers_str = ', '.join(s.get('headers', [])[:20])
            preview_str = ''
            for row in s.get('data_preview', [])[:3]:
                preview_str += ' | '.join([str(c)[:30] for c in row[:10]]) + '\n'
            sheets_text += f"\n工作表「{s['name']}」({s['row_count']}行x{s['column_count']}列):\n列: {headers_str}\n数据预览:\n{preview_str}"
        
        prompt = f"""你是一个数据分析师。请根据用户需求整理以下Excel数据。

文件: {data_summary.get('file_name', '')}
{sheets_text}

用户需求: {user_request}

请以JSON格式返回:
{{
    "summary": "数据总览总结（2-3句话）",
    "sections": [
        {{
            "title": "章节标题",
            "content": "详细内容，使用HTML格式，支持表格、列表等",
            "table": [["表头1", "表头2"], ["数据1", "数据2"]]  // 可选，用于生成表格
        }}
    ]
}}

要求:
1. 根据用户需求提取关键信息
2. 使用HTML格式化输出
3. 表格数据以二维数组形式提供
4. 内容要简洁明了"""
        
        response = req.post(
            f"{ai_config.get('base_url', 'https://dashscope.aliyuncs.com/api/v1')}/services/aigc/text-generation/generation",
            headers={
                'Authorization': f'Bearer {ai_config.get("api_key", "")}',
                'Content-Type': 'application/json'
            },
            json={
                'model': ai_config.get('model', 'qwen-turbo'),
                'input': {
                    'messages': [
                        {'role': 'system', 'content': '你是一个专业的数据分析师。'},
                        {'role': 'user', 'content': prompt}
                    ]
                },
                'parameters': {
                    'result_format': 'message',
                    'max_tokens': 2000
                }
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            output_text = result.get('output', {}).get('choices', [{}])[0].get('message', {}).get('content', '')
            
            # 尝试解析JSON
            try:
                # 清理可能的markdown代码块标记
                json_str = output_text
                if '```json' in json_str:
                    json_str = json_str.split('```json')[1].split('```')[0]
                elif '```' in json_str:
                    json_str = json_str.split('```')[1].split('```')[0]
                json_str = json_str.strip()
                
                parsed = json.loads(json_str)
                return _build_organized_result(parsed)
            except (json.JSONDecodeError, IndexError):
                # 如果JSON解析失败，使用文本
                return _build_text_result(output_text, user_request)
        else:
            logger.warning(f"AI请求失败: {response.status_code}")
            return _local_organize_excel(data_summary, user_request)
    except Exception as e:
        logger.warning(f"AI整理失败，使用本地整理: {e}")
        return _local_organize_excel(data_summary, user_request)


def _local_organize_excel(data_summary, user_request):
    """本地整理Excel数据"""
    sheets = data_summary.get('sheets', [])
    
    sections = []
    summary_text = f"文件「{data_summary.get('file_name', '')}」共包含 {len(sheets)} 个工作表。"
    
    for sheet in sheets:
        headers = sheet.get('headers', [])
        data_preview = sheet.get('data_preview', [])
        row_count = sheet.get('row_count', 0)
        
        # 构建表格数据
        table_data = [headers[:10]]  # 表头
        for row in data_preview[:10]:
            table_data.append([str(c)[:50] for c in row[:10]])
        
        # 分析数据特征
        numeric_cols = []
        for col_idx in range(min(len(headers), 10)):
            numeric_count = 0
            for row in data_preview:
                if col_idx < len(row) and row[col_idx]:
                    try:
                        float(str(row[col_idx]).replace(',', ''))
                        numeric_count += 1
                    except ValueError:
                        pass
            if numeric_count > len(data_preview) * 0.3:
                numeric_cols.append(headers[col_idx])
        
        # 生成章节
        section_content = f"<p>工作表「{sheet.get('name', '')}」包含 <strong>{row_count}</strong> 行数据，<strong>{len(headers)}</strong> 列字段。</p>"
        
        if numeric_cols:
            section_content += f"<p>主要数值字段: {', '.join(numeric_cols)}</p>"
        
        section_content += "<h4>数据预览</h4>"
        
        sections.append({
            'title': f"📋 {sheet.get('name', '')} ({row_count}行)",
            'content': section_content,
            'table': table_data
        })
    
    # 添加用户需求相关的总结
    sections.insert(0, {
        'title': '📊 数据总览',
        'content': f"<p>{summary_text}</p><p><strong>用户需求:</strong> {user_request}</p><p>请查看以下各工作表的详细数据预览。</p>",
        'table': []
    })
    
    return {
        'summary': summary_text,
        'sections': sections
    }


def _build_organized_result(parsed):
    """构建整理后的结果"""
    sections = parsed.get('sections', [])
    html_sections = []
    
    for section in sections:
        content = section.get('content', '')
        table = section.get('table', [])
        
        if table and len(table) > 1:
            # 生成HTML表格
            headers = table[0] if table else []
            rows = table[1:] if len(table) > 1 else []
            
            table_html = '<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px;"><thead><tr>'
            for h in headers:
                table_html += f'<th style="background:#1d1d1f;color:white;padding:10px 12px;text-align:left;font-weight:600;font-size:12px;">{h}</th>'
            table_html += '</tr></thead><tbody>'
            
            for row in rows:
                table_html += '<tr>'
                for cell in row:
                    table_html += f'<td style="padding:8px 12px;border-bottom:1px solid #e5e5ea;">{cell}</td>'
                table_html += '</tr>'
            
            table_html += '</tbody></table>'
            content += table_html
        
        html_sections.append({
            'title': section.get('title', ''),
            'content': content,
            'table': table
        })
    
    return {
        'summary': parsed.get('summary', ''),
        'sections': html_sections
    }


def _build_text_result(text, user_request):
    """从纯文本构建结果"""
    sections = [{
        'title': '📊 AI分析结果',
        'content': text.replace('\n', '<br>'),
        'table': []
    }]
    return {
        'summary': f'根据需求「{user_request}」生成的分析结果',
        'sections': sections
    }


# === Excel 整理后PDF生成路由 ===
@app.route('/excel-organize-pdf', methods=['POST'])
def api_excel_organize_pdf():
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': '请求数据格式错误'}), 400
    
    organized_data = data.get('organized_data', {})
    watermark = data.get('watermark', '')
    
    if not organized_data:
        return jsonify({'error': '缺少整理后的数据'}), 400
    
    try:
        sections = organized_data.get('sections', [])
        summary = organized_data.get('summary', '')
        user_request = organized_data.get('user_request', '')
        
        # 构建HTML内容
        html_content = _build_excel_report_html(sections, summary, user_request, watermark)

        import tempfile as tf
        
        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name
        
        pdf_filename = f"excel_report_{int(time.time())}.pdf"
        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)
        
        _render_pdf(html_path, pdf_path,
                    margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'},
                    extra_wait_ms=1000)
        
        try:
            os.unlink(html_path)
        except Exception:
            pass
        
        return jsonify({
            'status': 'success',
            'filename': pdf_filename
        })
    except Exception as e:
        logger.error(f"Excel PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


def _build_excel_report_html(sections, summary, user_request, watermark):
    """构建Excel报告HTML"""
    sections_html = ''
    for section in sections:
        title = section.get('title', '')
        content = section.get('content', '')
        
        sections_html += f'''
        <div style="margin-bottom:32px;">
            <h2 style="font-size:18px;font-weight:700;margin-bottom:16px;color:#1d1d1f;border-bottom:2px solid #0071e3;padding-bottom:8px;">{title}</h2>
            <div style="font-size:13px;line-height:1.8;color:#3c3c43;">{content}</div>
        </div>
        '''
    
    watermark_html = ''
    if watermark:
        watermark_html = f'''
        <div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.035;font-size:70px;font-weight:600;color:#0071e3;text-align:center;display:flex;align-items:center;justify-content:center;transform:rotate(-15deg);">
            {watermark}
        </div>
        '''
    
    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ 
            font-family: -apple-system, "SF Pro Display", "PingFang SC", "Microsoft YaHei", "Segoe UI", Roboto, sans-serif; 
            padding: 40px; 
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
        }}
        h1 {{ 
            font-size: 24px; 
            font-weight: 700; 
            margin-bottom: 8px;
            color: #1d1d1f;
        }}
        h2 {{ 
            margin-top: 1.5em;
        }}
        table {{ 
            border-collapse: collapse; 
            width: 100%; 
            margin: 1em 0;
            font-size: 13px;
        }}
        th {{ 
            background: #1d1d1f; 
            color: white; 
            padding: 10px 12px; 
            text-align: left;
            font-weight: 600;
            font-size: 12px;
        }}
        td {{ 
            border: 1px solid #e5e5ea; 
            padding: 8px 12px; 
        }}
        tr:nth-child(even) td {{
            background: #f5f5f7;
        }}
        .header {{
            text-align: center;
            margin-bottom: 32px;
            padding-bottom: 20px;
            border-bottom: 1px solid #e5e5ea;
        }}
        .summary-box {{
            background: linear-gradient(135deg, #f0f7ff, #e8f1ff);
            border: 1px solid #bae0ff;
            border-radius: 12px;
            padding: 16px 20px;
            margin-bottom: 24px;
            font-size: 13px;
            color: #3c3c43;
        }}
        .user-request {{
            font-size: 12px;
            color: #6e6e73;
            margin-top: 8px;
            font-style: italic;
        }}
    </style>
</head>
<body>
    {watermark_html}
    <div class="header">
        <h1>📊 数据分析报告</h1>
    </div>
    <div class="summary-box">
        <strong>📋 分析摘要:</strong> {summary}
        {f'<div class="user-request">💡 需求: {user_request}</div>' if user_request else ''}
    </div>
    {sections_html}
</body>
</html>'''


@app.route('/excel-pdf', methods=['POST'])
def api_excel_pdf():
    """PDF快转 - 根据结构化数据生成PDF"""
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': '请求数据格式错误'}), 400
    
    structured_data = data.get('structured_data', {})
    selected_sheets = data.get('selected_sheets', [])
    watermark = data.get('watermark', '')
    custom_title = data.get('custom_title', '').strip()
    
    if not structured_data:
        return jsonify({'error': '缺少结构化数据'}), 400
    
    try:
        html_content = _build_excel_structured_report_html(structured_data, selected_sheets, watermark, custom_title)

        import tempfile as tf
        
        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name
        
        # 生成PDF文件名：如果有自定义标题，使用自定义标题作为文件名
        if custom_title:
            safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
            pdf_filename = f"{safe_title}_{int(time.time())}.pdf"
            download_name = f"{safe_title}.pdf"
        else:
            pdf_filename = f"excel_pdf_{int(time.time())}.pdf"
            download_name = pdf_filename
        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

        _render_pdf(html_path, pdf_path,
                    margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'},
                    extra_wait_ms=1500)
        
        try:
            os.unlink(html_path)
        except Exception:
            pass
        
        return jsonify({
            'status': 'success',
            'filename': pdf_filename,
            'download_name': download_name
        })
    except Exception as e:
        logger.error(f"Excel PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/excel-select-pdf', methods=['POST'])
def api_excel_select_pdf():
    """PDF快转 - 根据选中数据生成PDF"""
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': '请求数据格式错误'}), 400
    
    structured_data = data.get('structured_data', {})
    selected_data = data.get('selected_data', {})  # {sheetName: [rowIndices]}
    selected_columns = data.get('selected_columns', {})  # {sheetName: [colIndices]}
    watermark = data.get('watermark', '')
    custom_title = data.get('custom_title', '').strip()
    
    if not selected_data:
        return jsonify({'error': '缺少选中数据'}), 400
    
    try:
        html_content = _build_excel_selected_report_html(structured_data, selected_data, selected_columns, watermark, custom_title)

        import tempfile as tf
        
        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name
        
        # 生成PDF文件名：如果有自定义标题，使用自定义标题作为文件名
        if custom_title:
            safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
            pdf_filename = f"{safe_title}_{int(time.time())}.pdf"
            download_name = f"{safe_title}.pdf"
        else:
            pdf_filename = f"excel_select_pdf_{int(time.time())}.pdf"
            download_name = pdf_filename
        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

        _render_pdf(html_path, pdf_path,
                    margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'},
                    extra_wait_ms=1500)
        
        try:
            os.unlink(html_path)
        except Exception:
            pass
        
        return jsonify({
            'status': 'success',
            'filename': pdf_filename,
            'download_name': download_name
        })
    except Exception as e:
        logger.error(f"Excel select PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


def _build_excel_structured_report_html(structured_data, selected_sheets, watermark, custom_title=''):
    """构建结构化数据报告HTML - 支持数组格式"""
    watermark_html = ''
    if watermark:
        watermark_items = ''.join([
            f'<div style="position:absolute;left:{x}px;top:{y}px;transform:rotate(-30deg);font-size:12px;font-weight:400;color:#0071e3;white-space:nowrap;">{_escape_html(watermark)}</div>'
            for x in range(80, 1200, 200)
            for y in range(150, 1000, 200)
        ])
        watermark_html = f'''
        <div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.08;overflow:hidden;">
            {watermark_items}
        </div>
        '''
    
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    title = _escape_html(custom_title) if custom_title else '📊 数据分析报告'
    
    # 获取工作表数据 - 支持数组格式 [{name, headers, rows}, ...]
    sheets_data = structured_data.get('sheets', []) if isinstance(structured_data, dict) else []
    
    # 构建工作表名称到数据的映射
    sheet_map = {}
    max_cols = 0
    for sheet in sheets_data:
        if isinstance(sheet, dict):
            name = sheet.get('name', '')
            if name:
                sheet_map[name] = sheet
                col_count = len(sheet.get('headers', []))
                if col_count > max_cols:
                    max_cols = col_count
    
    # 确定页面方向：列数超过15列使用横向打印
    use_landscape = max_cols > 15
    page_size = 'A4 landscape' if use_landscape else 'A4'
    
    # 确定要显示的工作表
    sheets_to_show = selected_sheets if selected_sheets else list(sheet_map.keys())
    
    # 处理选中的工作表数据
    sheets_html = ''
    for sheet_name in sheets_to_show:
        sheet_data = sheet_map.get(sheet_name, {})
        if not sheet_data:
            continue
        
        headers = sheet_data.get('headers', [])
        rows = sheet_data.get('rows', [])
        
        if not headers:
            continue
        
        # 如果没有 rows，使用 data_preview
        if not rows:
            rows = sheet_data.get('data_preview', [])
        
        if not rows:
            continue
        
        # 确保 headers 和 rows 对齐
        num_cols = len(headers)
        aligned_headers = [_escape_html(h) for h in headers]
        
        # 根据列数调整字体大小
        if num_cols > 50:
            font_size = '6px'
            padding = '2px 3px'
        elif num_cols > 30:
            font_size = '8px'
            padding = '3px 4px'
        elif num_cols > 15:
            font_size = '9px'
            padding = '4px 5px'
        else:
            font_size = '10px'
            padding = '5px 8px'
        
        # 生成表头
        header_html = ''.join([f'<th style="padding:{padding};font-size:{font_size};">{h}</th>' for h in aligned_headers])
        
        # 生成数据行（确保每行的列数与headers对齐）
        rows_html = ''
        display_rows = rows[:50]  # 最多显示50行
        for row in display_rows:
            # 对齐行数据到headers长度
            aligned_row = list(row[:num_cols]) + [''] * max(0, num_cols - len(row))
            cells = ''.join([f'<td style="border:1px solid #e5e5ea;padding:{padding};font-size:{font_size};">{_escape_html(cell)}</td>' for cell in aligned_row])
            rows_html += f'<tr>{cells}</tr>'
        
        total_rows = len(rows)
        sheets_html += f'''
        <div style="margin-bottom:20px;break-inside:avoid;">
            <h2 style="font-size:14px;font-weight:600;margin-bottom:10px;color:#1d1d1f;padding:6px 10px;background:#f5f5f7;border-radius:6px;border-left:3px solid #0071e3;">📋 {_escape_html(sheet_name)} ({num_cols}列)</h2>
            <div style="overflow-x:auto;border:1px solid #e5e5ea;border-radius:6px;">
                <table style="border-collapse:collapse;width:100%;">
                    <thead><tr style="background:#1d1d1f;color:white;">{header_html}</tr></thead>
                    <tbody>{rows_html}</tbody>
                </table>
            </div>
            {f'<p style="font-size:9px;color:#6e6e73;margin-top:4px;text-align:right;">共 {total_rows} 行数据，仅显示前50行</p>' if total_rows > 50 else ''}
        </div>
        '''
    
    if not sheets_html:
        sheets_html = '<p style="text-align:center;color:#6e6e73;padding:40px;">暂无数据</p>'
    
    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: {page_size}; margin: 10mm 8mm; }}
        body {{
            font-family: -apple-system, "SF Pro Display", "PingFang SC", "Microsoft YaHei", "Segoe UI", Roboto, sans-serif;
            padding: 5px;
            line-height: 1.3;
            max-width: 100%;
            margin: 0 auto;
            color: #1d1d1f;
        }}
        table {{ border-collapse: collapse; width: 100%; margin: 0; }}
        th {{ background: #1d1d1f; color: white; text-align: left; font-weight: 600; }}
        td {{ border: 1px solid #e5e5ea; }}
        h1 {{ font-size: 20px; font-weight: 700; margin: 0 0 6px 0; color: #1d1d1f; }}
        h2 {{ font-size: 14px; font-weight: 600; margin-bottom: 10px; }}
        .header-box {{
            text-align:center;
            margin-bottom:16px;
            padding-bottom:12px;
            border-bottom:2px solid #e5e5ea;
        }}
        .meta-info {{
            font-size: 9px;
            color:#6e6e73;
            margin-top:4px;
        }}
    </style>
</head>
<body>
    {watermark_html}
    <div class="header-box">
        <h1>{title}</h1>
        <div class="meta-info">生成时间: {now}</div>
    </div>
    {sheets_html}
</body>
</html>'''


def _build_excel_selected_report_html(structured_data, selected_data, selected_columns, watermark, custom_title=''):
    """构建选中数据报告HTML - 根据选中的行列索引从原始数据中提取"""
    watermark_html = ''
    if watermark:
        watermark_items = ''.join([
            f'<div style="position:absolute;left:{x}px;top:{y}px;transform:rotate(-30deg);font-size:12px;font-weight:400;color:#0071e3;white-space:nowrap;">{_escape_html(watermark)}</div>'
            for x in range(80, 1200, 200)
            for y in range(150, 1000, 200)
        ])
        watermark_html = f'''
        <div style="position:fixed;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:9999;opacity:0.08;overflow:hidden;">
            {watermark_items}
        </div>
        '''
    
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    title = _escape_html(custom_title) if custom_title else '📊 数据分析报告'
    
    # 获取原始工作表数据
    sheets_data = structured_data.get('sheets', []) if isinstance(structured_data, dict) else []
    
    # 构建工作表名称到数据的映射
    sheet_map = {}
    max_cols = 0
    for sheet in sheets_data:
        if isinstance(sheet, dict):
            name = sheet.get('name', '')
            if name:
                sheet_map[name] = sheet
    
    # 计算最大列数（用于决定页面方向）
    for sheet_name, row_indices in selected_data.items():
        col_indices = selected_columns.get(sheet_name, [])
        if len(col_indices) > max_cols:
            max_cols = len(col_indices)
    
    # 确定页面方向：列数超过15列使用横向打印
    use_landscape = max_cols > 15
    page_size = 'A4 landscape' if use_landscape else 'A4'
    
    # selected_data: {sheetName: [rowIndices]} 行索引
    # selected_columns: {sheetName: [colIndices]} 列索引
    
    tables_html = ''
    for sheet_name, row_indices in selected_data.items():
        sheet_data = sheet_map.get(sheet_name, {})
        if not sheet_data:
            continue
        
        headers = sheet_data.get('headers', [])
        rows = sheet_data.get('rows', [])
        
        # 如果没有 rows，使用 data_preview
        if not rows:
            rows = sheet_data.get('data_preview', [])
        
        if not headers or not rows:
            continue
        
        # 获取选中的列索引
        col_indices = selected_columns.get(sheet_name, list(range(len(headers))))
        
        # 过滤列（带HTML转义）
        filtered_headers = [_escape_html(headers[i]) if i < len(headers) else '' for i in col_indices]
        
        # 过滤行（确保对齐）
        num_filtered_cols = len(filtered_headers)
        filtered_rows = []
        for row_idx in row_indices:
            if row_idx < len(rows):
                row = rows[row_idx]
                filtered_row = [_escape_html(row[i]) if i < len(row) else '' for i in col_indices]
                # 对齐到列数
                filtered_row = filtered_row[:num_filtered_cols] + [''] * max(0, num_filtered_cols - len(filtered_row))
                filtered_rows.append(filtered_row)
        
        if not filtered_rows:
            continue
        
        # 根据列数调整字体大小
        if num_filtered_cols > 50:
            font_size = '6px'
            padding = '2px 3px'
        elif num_filtered_cols > 30:
            font_size = '8px'
            padding = '3px 4px'
        elif num_filtered_cols > 15:
            font_size = '9px'
            padding = '4px 5px'
        else:
            font_size = '10px'
            padding = '5px 8px'
        
        # 生成表格
        header_html = ''.join([f'<th style="padding:{padding};font-size:{font_size};">{h}</th>' for h in filtered_headers])
        rows_html = ''
        display_rows = filtered_rows[:50]
        for row in display_rows:
            cells = ''.join([f'<td style="border:1px solid #e5e5ea;padding:{padding};font-size:{font_size};">{c}</td>' for c in row])
            rows_html += f'<tr>{cells}</tr>'
        
        total_filtered = len(filtered_rows)
        tables_html += f'''
        <div style="margin-bottom:20px;break-inside:avoid;">
            <h2 style="font-size:14px;font-weight:600;margin-bottom:10px;color:#1d1d1f;padding:6px 10px;background:#f5f5f7;border-radius:6px;border-left:3px solid #0071e3;">📋 {_escape_html(sheet_name)} ({num_filtered_cols}列)</h2>
            <div style="overflow-x:auto;border:1px solid #e5e5ea;border-radius:6px;">
                <table style="border-collapse:collapse;width:100%;">
                    <thead><tr style="background:#1d1d1f;color:white;">{header_html}</tr></thead>
                    <tbody>{rows_html}</tbody>
                </table>
            </div>
            {f'<p style="font-size:9px;color:#6e6e73;margin-top:4px;text-align:right;">共 {total_filtered} 行数据，仅显示前50行</p>' if total_filtered > 50 else ''}
        </div>
        '''
    
    if not tables_html:
        tables_html = '<p style="text-align:center;color:#6e6e73;padding:40px;">暂无选中数据</p>'
    
    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: {page_size}; margin: 10mm 8mm; }}
        body {{
            font-family: -apple-system, "SF Pro Display", "PingFang SC", "Microsoft YaHei", "Segoe UI", Roboto, sans-serif;
            padding: 5px;
            line-height: 1.3;
            max-width: 100%;
            margin: 0 auto;
            color: #1d1d1f;
        }}
        table {{ border-collapse: collapse; width: 100%; margin: 0; }}
        th {{ background: #1d1d1f; color: white; text-align: left; font-weight: 600; }}
        td {{ border: 1px solid #e5e5ea; }}
        h1 {{ font-size: 20px; font-weight: 700; margin: 0 0 6px 0; color: #1d1d1f; }}
        h2 {{ font-size: 14px; font-weight: 600; margin-bottom: 10px; }}
        .header-box {{
            text-align:center;
            margin-bottom:16px;
            padding-bottom:12px;
            border-bottom:2px solid #e5e5ea;
        }}
        .meta-info {{
            font-size: 9px;
            color:#6e6e73;
            margin-top:4px;
        }}
    </style>
</head>
<body>
    {watermark_html}
    <div class="header-box">
        <h1>{title}</h1>
        <div class="meta-info">生成时间: {now}</div>
    </div>
    {tables_html}
</body>
</html>'''


# === v3.0: AI 增强 CR 问题分析 ===
@app.route('/api/excel-analyze-ai', methods=['POST'])
@login_required_or_guest
def api_excel_analyze_ai():
    """v3.0: AI 增强 CR 问题分析 — 根因分析和改进建议"""

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请在设置页面配置 API Key'}), 503

    data = request.get_json(silent=True) or {}
    analysis = data.get('analysis', {})
    if not analysis:
        return jsonify({'error': '缺少分析数据'}), 400

    summary = analysis.get('summary', {})
    modules = analysis.get('modules', [])
    developers = analysis.get('developers', [])

    # 模块统计 Top 10
    modules_text = ''
    for m in (modules[:10] if isinstance(modules, list) else []):
        if isinstance(m, dict):
            modules_text += f"- {m.get('name', 'N/A')}: {m.get('count', 0)}个问题"
            if m.get('unresolved'):
                modules_text += f" (未解决: {m['unresolved']})"
            modules_text += "\n"

    # 研发统计 Top 10
    devs_text = ''
    for d in (developers[:10] if isinstance(developers, list) else []):
        if isinstance(d, dict):
            devs_text += f"- {d.get('name', 'N/A')}: {d.get('count', 0)}个问题"
            if d.get('unresolved'):
                devs_text += f" (未解决: {d['unresolved']})"
            devs_text += "\n"

    prompt = f"""你是一位资深质量管理专家，请基于以下CR问题分析数据生成专业的AI分析报告。

## 问题统计
- 总问题数: {summary.get('total_issues', 0)}
- 已解决: {summary.get('resolved', 0)}
- 未解决: {summary.get('unresolved', 0)}
- 解决率: {summary.get('resolution_rate', 'N/A')}
- 阻塞问题: {summary.get('blocker', 0)}
- 严重问题: {summary.get('critical', 0)}
- 主要问题: {summary.get('major', 0)}
- 次要问题: {summary.get('minor', 0)}
- 建议问题: {summary.get('trivial', 0)}

## 模块统计 (Top 10)
{modules_text or 'N/A'}

## 研发人员统计 (Top 10)
{devs_text or 'N/A'}

请按以下格式输出分析：

### 📊 总体评估
（1-2段话总结问题整体状况，包括解决率评价和质量风险等级）

### 🔍 根因分析
（分析问题集中的模块和人员，指出可能的根因，3-5条）

### ⚠️ 高风险领域
（识别需要重点关注的模块或人员，3-5条）

### 📈 趋势预测
（基于当前数据预测可能的发展趋势，2-3条）

### 💡 改进建议
（列出3-5条可操作的流程改进建议）

请使用简洁专业的中文，避免空话套话。"""

    try:
        messages = [{'role': 'user', 'content': prompt}]
        reply = _call_ai(messages, max_tokens=1500, temperature=0.3, timeout=60)
        return jsonify({'status': 'success', 'analysis': reply})
    except Exception as e:
        logger.error(f'AI CR分析失败: {e}')
        return jsonify({'error': f'AI分析失败: {str(e)}'}), 502


@app.route('/api/excel-analyze-ai-stream', methods=['POST'])
@login_required_or_guest
def api_excel_analyze_ai_stream():
    """SSE 流式版：AI 增强 CR 问题分析 — 边生成边输出"""

    ai_config = get_ai_config()
    if not ai_config.get('enabled'):
        return jsonify({'error': 'AI功能未配置，请在设置页面配置 API Key'}), 503

    data = request.get_json(silent=True) or {}
    analysis = data.get('analysis', {})
    if not analysis:
        return jsonify({'error': '缺少分析数据'}), 400

    summary = analysis.get('summary', {})
    modules = analysis.get('modules', [])
    developers = analysis.get('developers', [])

    modules_text = ''
    for m in (modules[:10] if isinstance(modules, list) else []):
        if isinstance(m, dict):
            modules_text += f"- {m.get('name', 'N/A')}: {m.get('count', 0)}个问题"
            if m.get('unresolved'):
                modules_text += f" (未解决: {m['unresolved']})"
            modules_text += "\n"

    devs_text = ''
    for d in (developers[:10] if isinstance(developers, list) else []):
        if isinstance(d, dict):
            devs_text += f"- {d.get('name', 'N/A')}: {d.get('count', 0)}个问题"
            if d.get('unresolved'):
                devs_text += f" (未解决: {d['unresolved']})"
            devs_text += "\n"

    prompt = f"""你是一位资深质量管理专家，请基于以下CR问题分析数据生成专业的AI分析报告。

## 问题统计
- 总问题数: {summary.get('total_issues', 0)}
- 已解决: {summary.get('resolved', 0)}
- 未解决: {summary.get('unresolved', 0)}
- 解决率: {summary.get('resolution_rate', 'N/A')}
- 阻塞问题: {summary.get('blocker', 0)}
- 严重问题: {summary.get('critical', 0)}
- 主要问题: {summary.get('major', 0)}
- 次要问题: {summary.get('minor', 0)}
- 建议问题: {summary.get('trivial', 0)}

## 模块统计 (Top 10)
{modules_text or 'N/A'}

## 研发人员统计 (Top 10)
{devs_text or 'N/A'}

请按以下格式输出分析：

### 📊 总体评估
（1-2段话总结问题整体状况，包括解决率评价和质量风险等级）

### 🔍 根因分析
（分析问题集中的模块和人员，指出可能的根因，3-5条）

### ⚠️ 高风险领域
（识别需要重点关注的模块或人员，3-5条）

### 📈 趋势预测
（基于当前数据预测可能的发展趋势，2-3条）

### 💡 改进建议
（列出3-5条可操作的流程改进建议）

请使用简洁专业的中文，避免空话套话。"""

    messages = [{'role': 'user', 'content': prompt}]
    return Response(
        stream_with_context(_call_ai_stream(messages, max_tokens=1500, temperature=0.3)),
        content_type='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no', 'Connection': 'keep-alive'}
    )


# === Excel分析PDF生成 API ===
@app.route('/api/excel-analyze-pdf', methods=['POST'])
def api_excel_analyze_pdf():
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': '请求数据格式错误'}), 400

    analysis_data = data.get('analysis_data', {})
    watermark = data.get('watermark', '')
    custom_title = data.get('custom_title', '').strip()
    file_name = data.get('file_name', '')

    if not analysis_data:
        return jsonify({'error': '缺少分析数据'}), 400

    try:
        html_content = _build_cr_analysis_report_html(analysis_data, watermark, file_name, custom_title)

        import tempfile as tf

        with tf.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name

        # 生成PDF文件名：如果有自定义标题，使用自定义标题作为文件名
        if custom_title:
            # 清理文件名中的非法字符
            safe_title = re.sub(r'[\\/:*?"<>|]', '_', custom_title)
            pdf_filename = f"{safe_title}_{datetime.now(_CST).strftime('%Y%m%d_%H%M%S')}.pdf"
            download_name = f"{safe_title}.pdf"
        else:
            pdf_filename = f"cr_analysis_{int(time.time())}.pdf"
            download_name = pdf_filename
        
        pdf_path = os.path.join(app.config['PDF_FOLDER'], pdf_filename)

        _render_pdf(html_path, pdf_path,
                    margin={'top': '20mm', 'bottom': '20mm', 'left': '15mm', 'right': '15mm'},
                    extra_wait_ms=3000, wait_selector='canvas')

        try:
            os.unlink(html_path)
        except Exception:
            pass

        return jsonify({
            'status': 'success',
            'filename': pdf_filename,
            'download_name': download_name
        })
    except Exception as e:
        logger.error(f"CR分析PDF生成失败: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500
@app.route('/download/<filename>')
def download_file(filename):
    # 安全：使用 secure_filename 防止路径遍历攻击
    from werkzeug.utils import secure_filename
    safe_name = secure_filename(filename)
    if not safe_name or safe_name != filename:
        # 如果文件名被修改（含路径遍历字符），拒绝请求
        return jsonify({'error': '无效的文件名'}), 400

    # 先检查UPLOAD_FOLDER
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)

    # 再检查PDF_FOLDER
    filepath = os.path.join(app.config['PDF_FOLDER'], safe_name)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)

    return jsonify({'error': '文件不存在'}), 404
# === NoteNB 笔记应用路由 ===
NOTENB_DIST = os.path.join(static_dir, 'noteNB')

@app.route('/noteNB')
def notenb_redirect():
    return redirect('/noteNB/')

@app.route('/noteNB/')
def notenb_index():
    # 服务端注入用户 ID，实现同步数据隔离（避免异步 fetch 时序问题）
    user = auth.get_current_user()
    uid = user['id'] if user else ''
    user_name = user.get('name', '') if user else ''
    user_avatar = user.get('avatar', '') if user else ''
    html_path = os.path.join(NOTENB_DIST, 'index.html')
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()
    # 安全：使用 json.dumps 转义用户信息，防止 XSS
    import json as _json
    user_info_script = '<script>window._SERVER_USER_ID=' + _json.dumps(str(uid)) + ';window._SERVER_USER_NAME=' + _json.dumps(str(user_name)) + ';window._SERVER_USER_AVATAR=' + _json.dumps(str(user_avatar)) + ';</script>'
    html = html.replace('</head>', user_info_script + '\n</head>')
    resp = make_response(html)
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    return resp

@app.route('/noteNB/assets/<path:filename>')
def notenb_assets(filename):
    notenb_assets = os.path.join(NOTENB_DIST, 'assets')
    return send_from_directory(notenb_assets, filename)

@app.route('/noteNB/<path:path>')
def notenb_catch_all(path):
    file_path = os.path.join(NOTENB_DIST, path)
    if os.path.isfile(file_path):
        return send_file(file_path)
    # 对于 SPA 路由回退，也需要注入用户信息
    return notenb_index()


# === 静态资源路由 ===
@app.route('/assets/<path:filename>')
def serve_assets(filename):
    assets_dir = os.path.join(base_dir, 'assets')
    return send_from_directory(assets_dir, filename)


if __name__ == '__main__':
    try:
        port = int(os.environ.get('PORT', 8080))
        print(f"GGB 1.0 启动中... (port={port})")
        logger.info(f"GGB 1.0 启动中... (port={port})")
        app.run(host='0.0.0.0', port=port, debug=False)
    except Exception as e:
        logger.error(f"启动失败: {e}")
        logger.error(traceback.format_exc())
        print(f"启动失败: {str(e)}")

